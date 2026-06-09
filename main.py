from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any, Tuple
import json
import math
import re
from collections import defaultdict
import pandas as pd
import numpy as np
from scipy import stats
from scipy.stats import (
    shapiro, spearmanr, mannwhitneyu, kruskal, ttest_rel, ttest_ind,
    linregress, kstest, levene, tukey_hsd, skew, kurtosis,
)
import anthropic
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config import ANTHROPIC_API_KEY, ANTHROPIC_MODEL

app = FastAPI(title="StatAI - Akademik Analiz API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Modeller ──────────────────────────────────────────────────────────────────

class Variable(BaseModel):
    name: str
    label: str
    type: str  # "continuous" | "categorical"
    role: str = "grouping"  # "grouping" | "outcome"
    categories: Optional[List[str]] = None
    included: bool = True
    scale_min: Optional[float] = None
    scale_max: Optional[float] = None

class DataRow(BaseModel):
    values: dict

class AnalysisRequest(BaseModel):
    variables: List[Variable]
    data: List[DataRow]
    active_types: Optional[List[str]] = None

class PlanRequest(BaseModel):
    variables: List[Variable]
    data: List[DataRow]

class BulguRequest(BaseModel):
    result: Any

class WordExportRequest(BaseModel):
    results: List[dict]
    bulgular: Optional[Dict[str, str]] = None
    intro: Optional[str] = None

class CronbachRequest(BaseModel):
    columns: List[str]
    data: List[DataRow]

class PairedRequest(BaseModel):
    col1: str
    col2: str
    data: List[DataRow]

class ClassifyRequest(BaseModel):
    columns: List[str]
    samples: Dict[str, List[Any]]

class RecommendRequest(BaseModel):
    columns: List[str]
    samples: Dict[str, List[Any]]
    research_topic: str

class DetectScalesRequest(BaseModel):
    columns: List[str]

class CronbachBatchRequest(BaseModel):
    scales: List[dict]
    data: List[DataRow]

class SpssTableRequest(BaseModel):
    content: str

# ── SPSS Tablo → Markdown ─────────────────────────────────────────────────────

STAT_COL_RE = re.compile(r"χ|chi|anova|\bf\s*\(|^p$|^p\s", re.I)


def _flatten_columns(columns) -> List[str]:
    if isinstance(columns, pd.MultiIndex):
        out = []
        for i, col in enumerate(columns):
            parts = [str(p).strip() for p in col if str(p).strip() not in ("", "nan", "None")]
            out.append(" ".join(parts) if parts else f"col_{i}")
        return out
    return [str(c).strip() if str(c).strip() else f"col_{i}" for i, c in enumerate(columns)]


def _is_blank(val) -> bool:
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return True
    s = str(val).strip()
    return s in ("", "nan", "None", "NaN", ",")


def _clean_stat_value(val) -> str:
    if _is_blank(val):
        return ""
    s = str(val).strip()
    m = re.search(r"=\s*(.+)$", s)
    return m.group(1).strip() if m else s


def _is_stat_column(col_name: str) -> bool:
    return bool(STAT_COL_RE.search(str(col_name)))


def _clean_spss_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = _flatten_columns(df.columns)

    # Yinelenen sütun başlıklarını benzersizleştir
    seen: Dict[str, int] = {}
    new_cols = []
    for col in df.columns:
        base = str(col)
        if base not in seen:
            seen[base] = 0
            new_cols.append(base)
        else:
            seen[base] += 1
            new_cols.append(f"{base}_{seen[base]}")
    df.columns = new_cols

    # Hayalet / tamamen boş sütunları kaldır
    keep = []
    for col in df.columns:
        if any(not _is_blank(v) for v in df[col]):
            keep.append(col)
    if keep:
        df = df[keep]

    # Birleştirilmiş hücreleri dağıt (forward fill)
    df = df.ffill(axis=0)

    # Test istatistikleri (χ², p, F) — genelde Toplam satırında; tüm satırlara yay
    for col in df.columns:
        if _is_stat_column(col):
            series = df[col].apply(_clean_stat_value)
            series = series.replace("", np.nan)
            df[col] = series.ffill().bfill()

    # İlk sütun: boş kategori → Kayıp Veri
    if len(df.columns):
        first = df.columns[0]
        for idx in df.index:
            if _is_blank(df.at[idx, first]):
                rest = df.loc[idx, df.columns[1:]]
                if any(not _is_blank(v) for v in rest):
                    df.at[idx, first] = "Kayıp Veri"

    # Post-hoc (I)/(J) grupları — (I) birleştirilmişse forward fill
    for col in df.columns:
        cl = str(col).lower()
        if "(i)" in cl or cl.startswith("i grup") or cl == "i":
            df[col] = df[col].ffill()

    # Tamamen boş satırları kaldır
    mask = df.apply(lambda row: all(_is_blank(v) for v in row), axis=1)
    df = df[~mask].reset_index(drop=True)

    return df


def _dataframe_to_markdown(df: pd.DataFrame) -> str:
    cols = [str(c) for c in df.columns]
    lines = [
        "| " + " | ".join(cols) + " |",
        "| " + " | ".join([":---"] * len(cols)) + " |",
    ]
    for _, row in df.iterrows():
        cells = [str(v).strip() if not _is_blank(v) else "" for v in row]
        if len(cells) != len(cols):
            continue
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def convert_spss_table(content: str) -> str:
    content = content.strip()
    if not content:
        raise ValueError("Boş içerik")

    if "<table" in content.lower():
        dfs = pd.read_html(io.StringIO(content))
    else:
        sep = "\t" if content.count("\t") > content.count(",") else ","
        try:
            dfs = [pd.read_csv(io.StringIO(content), sep=sep, engine="python", on_bad_lines="skip")]
        except Exception:
            dfs = pd.read_html(io.StringIO(content))

    parts = []
    for df in dfs:
        if df is None or df.empty:
            continue
        cleaned = _clean_spss_dataframe(df)
        if not cleaned.empty:
            parts.append(_dataframe_to_markdown(cleaned))

    if not parts:
        raise ValueError("Tablo ayrıştırılamadı")
    return "\n\n".join(parts)


SPSS_CONVERT_SYSTEM = """Sen SPSS çıktı dönüştürme uzmanısın. Ham metin, HTML veya kopyala-yapıştır formatındaki SPSS analiz çıktılarını (çapraz tablolar, ANOVA, Ki-Kare, t-Testi vb.) sıfır veri kaybıyla kusursuz Markdown tablosuna dönüştür.

KATI KURALLAR:
1. BİRLEŞTİRİLMİŞ HÜCRELER: χ², p, F değerlerini TÜM satırlara forward fill ile çoğalt.
2. HAYALET SÜTUNLAR: Boş sütunları ve sütun kaymalarını düzelt; header ile satır sütun sayısı eşit olsun.
3. KAYIP VERİ: İsimsiz ama frekanslı satırları "Kayıp Veri" olarak adlandır.
4. POST-HOC: (I) grubu birleştirilmişse altındaki (J) satırlarına (I) adını yaz.

SADECE temizlenmiş Markdown tablolarını döndür. Yorum yapma."""


# ── APA Yardımcıları ─────────────────────────────────────────────────────────

class TableCounter:
    def __init__(self):
        self.n = 0

    def next(self, title_suffix: str) -> Tuple[int, str]:
        self.n += 1
        return self.n, f"Tablo {self.n}. {title_suffix}"

def fmt_num(x: float, decimals: int = 2) -> str:
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return "—"
    return f"{float(x):.{decimals}f}"

def fmt_p(p: float) -> str:
    if p is None:
        return "—"
    p = float(p)
    if p < 0.001:
        return "< .001"
    return f"{p:.3f}".lstrip("0")

def p_stars(p: float) -> str:
    if p is None:
        return ""
    p = float(p)
    if p < 0.001:
        return "***"
    if p < 0.01:
        return "**"
    if p < 0.05:
        return "*"
    return ""

def fmt_p_display(p: float) -> str:
    return f"{fmt_p(p)}{p_stars(p)}"

def cohens_d_interpretation(d: float) -> str:
    ad = abs(d)
    if ad < 0.20:
        return "önemsiz"
    if ad < 0.50:
        return "küçük"
    if ad < 0.80:
        return "orta"
    return "büyük"

def eta_interpretation(eta2: float) -> str:
    if eta2 < 0.01:
        return "küçük"
    if eta2 <= 0.06:
        return "orta"
    if eta2 <= 0.14:
        return "orta-üst"
    return "büyük"

def cohens_d(g1: list, g2: list) -> float:
    n1, n2 = len(g1), len(g2)
    if n1 < 2 or n2 < 2:
        return 0.0
    v1, v2 = np.var(g1, ddof=1), np.var(g2, ddof=1)
    pooled = np.sqrt(((n1 - 1) * v1 + (n2 - 1) * v2) / (n1 + n2 - 2))
    if pooled == 0:
        return 0.0
    return float((np.mean(g1) - np.mean(g2)) / pooled)

def welch_df(g1: list, g2: list) -> float:
    n1, n2 = len(g1), len(g2)
    v1, v2 = np.var(g1, ddof=1), np.var(g2, ddof=1)
    num = (v1 / n1 + v2 / n2) ** 2
    den = (v1 / n1) ** 2 / (n1 - 1) + (v2 / n2) ** 2 / (n2 - 1)
    return float(num / den) if den else n1 + n2 - 2

def eta_squared(group_lists: list) -> float:
    all_data = np.concatenate([np.array(g) for g in group_lists])
    grand = np.mean(all_data)
    ss_total = np.sum((all_data - grand) ** 2)
    if ss_total == 0:
        return 0.0
    ss_between = sum(len(g) * (np.mean(g) - grand) ** 2 for g in group_lists)
    return float(ss_between / ss_total)

def make_result(
    rtype: str, table_no: int, title: str, headers: list, rows: list, note: str, **extra
) -> dict:
    return {
        "type": rtype,
        "table_number": table_no,
        "title": title,
        "headers": headers,
        "rows": rows,
        "note": note,
        **extra,
    }

def normalize_continuous_columns(df: pd.DataFrame, variables: List[Variable]) -> pd.DataFrame:
    df = df.copy()
    for v in variables:
        if v.type != "continuous" or v.name not in df.columns:
            continue
        df[v.name] = (
            df[v.name].astype(str)
            .str.strip()
            .str.replace(",", ".", regex=False)
        )
        df[v.name] = pd.to_numeric(df[v.name], errors="coerce")
    return df


def is_numeric_continuous(df: pd.DataFrame, v: Variable, norm_map: dict) -> bool:
    """Gerçekten sayısal sürekli mi? Kategorik kodlanmış değilse."""
    if v.name not in df.columns:
        return False
    series = df[v.name].dropna()
    if len(series) < 3:
        return False
    unique_count = series.nunique()
    if unique_count < 10:
        return False
    return True

def missing_data_report(df: pd.DataFrame, variables: List[Variable]) -> list:
    total = len(df)
    report = []
    for v in variables:
        if v.name not in df.columns:
            continue
        col = df[v.name]
        missing = col.isna() | (col.astype(str).str.strip() == "")
        pct = round(float(missing.sum()) / total * 100, 1) if total else 0.0
        warning = "none"
        if pct > 30:
            warning = "high"
        elif pct > 10:
            warning = "medium"
        report.append({
            "column": v.label, "name": v.name,
            "missing_pct": pct, "missing_n": int(missing.sum()), "warning": warning,
        })
    return report

# ── Normallik ─────────────────────────────────────────────────────────────────

def assess_normality(series: pd.Series) -> dict:
    s = series.dropna().astype(float)
    n = len(s)
    if n < 3:
        return {
            "n": n, "statistic": None, "stat_label": "W", "df": None, "p": None,
            "skewness": None, "skew_se": None, "kurtosis": None, "kurt_se": None,
            "is_parametric": True, "method": "yetersiz_veri",
        }

    sk = float(skew(s))
    ku = float(kurtosis(s, fisher=True))
    se_sk = float(np.sqrt(6 / n))
    se_ku = float(np.sqrt(24 / n))
    skew_ok = abs(sk) <= 2.0
    kurt_ok = abs(ku) <= 2.0

    if n <= 50:
        stat, p = shapiro(s)
        stat_label, method = "W", "Shapiro-Wilk"
    else:
        stat, p = kstest(s, "norm", args=(s.mean(), s.std(ddof=1)))
        stat_label, method = "D", "Kolmogorov-Smirnov"

    p = float(p)
    if p >= 0.05:
        is_parametric = True
    elif skew_ok and kurt_ok:
        is_parametric = True
    else:
        is_parametric = False

    return {
        "n": n,
        "statistic": round(float(stat), 3),
        "stat_label": stat_label,
        "df": n,
        "p": round(p, 3),
        "skewness": round(sk, 2),
        "skew_se": round(se_sk, 2),
        "kurtosis": round(ku, 2),
        "kurt_se": round(se_ku, 2),
        "is_parametric": bool(is_parametric),
        "method": method,
    }

def build_intro(n_total: int, parametric_vars: list, nonparametric_vars: list) -> str:
    parts = [f"Çalışmaya katılan toplam örneklem sayısı N = {n_total}'dir."]
    if parametric_vars:
        parts.append(
            f"{', '.join(parametric_vars)} değişkenleri için normallik varsayımları karşılandığından "
            "parametrik testler uygulanmıştır."
        )
    if nonparametric_vars:
        parts.append(
            f"{', '.join(nonparametric_vars)} değişkenleri için normallik varsayımları karşılanmadığından "
            "non-parametrik testler tercih edilmiştir."
        )
    return " ".join(parts)

# ── Tablo Üreticileri ─────────────────────────────────────────────────────────

def table_descriptive(tc: TableCounter, variables: List[Variable], df: pd.DataFrame) -> Optional[dict]:
    rows = []
    for v in variables:
        if v.name not in df.columns:
            continue
        s = pd.to_numeric(df[v.name], errors="coerce").dropna()
        if len(s) == 0:
            continue
        theory = "—"
        if v.scale_min is not None and v.scale_max is not None:
            theory = f"{fmt_num(v.scale_min)} – {fmt_num(v.scale_max)}"
        rows.append([
            v.label, str(len(s)),
            f"{fmt_num(s.mean())} ± {fmt_num(s.std(ddof=1))}",
            fmt_num(s.median()), f"{fmt_num(s.min())} – {fmt_num(s.max())}", theory,
        ])
    if not rows:
        return None
    no, title = tc.next("Tanımlayıcı İstatistikler")
    return make_result(
        "descriptive", no, title,
        ["Ölçek", "n", "x̄ ± SS", "Medyan", "Min – Maks", "Teorik Aralık"],
        rows, "Not. SS = Standart Sapma.",
    )

def table_normality(tc: TableCounter, variables: List[Variable], df: pd.DataFrame, norm_map: dict) -> Optional[dict]:
    rows = []
    for v in variables:
        if v.name not in norm_map:
            continue
        nm = norm_map[v.name]
        if nm["statistic"] is None:
            continue
        rows.append([
            v.label, f"{nm['stat_label']} = {nm['statistic']}", str(nm["df"]),
            fmt_p_display(nm["p"]),
            f"{nm['skewness']} / {nm['skew_se']}",
            f"{nm['kurtosis']} / {nm['kurt_se']}",
        ])
    if not rows:
        return None
    no, title = tc.next("Normallik Testi Sonuçları")
    return make_result(
        "normality", no, title,
        ["Değişken", "İstatistik", "df", "p", "Çarpıklık / Std. Hata", "Basıklık / Std. Hata"],
        rows,
        "Not. * p < .05; ** p < .01; *** p < .001. Çarpıklık ve basıklık ±2.0 içindeyse normal dağılım varsayımı karşılanmış kabul edilir.",
        norm_map={v.name: norm_map[v.name] for v in variables if v.name in norm_map},
    )

def table_frequency(tc: TableCounter, series: pd.Series, label: str) -> dict:
    counts = series.dropna().value_counts()
    total = int(counts.sum())
    rows = []
    for val, cnt in counts.items():
        pct = round(cnt / total * 100, 1) if total else 0
        rows.append([label, str(val), str(int(cnt)), fmt_num(pct, 1)])
    rows.append([label, "Toplam", str(total), "100.0"])
    no, title = tc.next(f"{label} Dağılımı")
    return make_result(
        "frequency", no, title,
        ["Değişken", "Kategori", "n", "%"],
        rows, "Not. Değerler frekans (n) ve yüzde (%) olarak verilmiştir.",
        variable=label, n=total,
    )

def table_chi_square(tc: TableCounter, df: pd.DataFrame, v1: Variable, v2: Variable) -> dict:
    ct = pd.crosstab(df[v1.name], df[v2.name])
    chi2, p, dof, _ = stats.chi2_contingency(ct)
    col_headers = [str(c) for c in ct.columns]
    rows = []
    for idx in ct.index:
        row_total = int(ct.loc[idx].sum())
        cells = [str(idx)]
        for col in ct.columns:
            n = int(ct.loc[idx, col])
            pct = round(n / row_total * 100, 1) if row_total else 0
            cells.append(f"{n} ({fmt_num(pct, 1)})")
        cells.append(str(row_total))
        cells.extend(["", ""])
        rows.append(cells)
    summary_row = ["Toplam"] + [str(int(ct[c].sum())) for c in ct.columns] + [str(int(ct.values.sum()))]
    summary_row += [f"χ² = {fmt_num(chi2)}", fmt_p_display(p)]
    rows.append(summary_row)
    headers = [v1.label] + [f"{c} n (%)" for c in col_headers] + ["Toplam", "χ²", "p"]
    no, title = tc.next(f"{v1.label} × {v2.label} Dağılımı (Ki-Kare Testi)")
    return make_result(
        "chi_square", no, title, headers, rows,
        "Not. * p < .05. Değerler n (kişi sayısı) ve satır yüzdesi (%) olarak verilmiştir.",
        chi2=round(float(chi2), 3), p=round(float(p), 3), dof=int(dof),
        significant=bool(p < 0.05), var1=v1.label, var2=v2.label,
    )

def table_ttest(tc: TableCounter, df: pd.DataFrame, cv: Variable, sv: Variable) -> dict:
    groups = df.groupby(cv.name)[sv.name].apply(lambda x: pd.to_numeric(x, errors="coerce").dropna().tolist())
    g1, g2 = list(groups.iloc[0]), list(groups.iloc[1])
    n1, n2 = len(g1), len(g2)
    lev_f, lev_p = levene(g1, g2)
    welch = float(lev_p) < 0.05
    t, p = ttest_ind(g1, g2, equal_var=not welch)
    df_t = int(welch_df(g1, g2)) if welch else n1 + n2 - 2
    d = cohens_d(g1, g2)
    rows = []
    for name, g in zip(groups.index, [g1, g2]):
        rows.append([
            sv.label, str(name), str(len(g)),
            f"{fmt_num(np.mean(g))} ± {fmt_num(np.std(g, ddof=1))}",
            "", "", "", fmt_num(d) if name == groups.index[0] else "",
        ])
    rows[0][4] = fmt_num(t)
    rows[0][5] = str(df_t)
    rows[0][6] = fmt_p_display(p)
    lev_note = (
        f"Welch düzeltmeli t değeri raporlanmıştır. Levene varyans homojenliği testi: F({n1-1}, {n2-1}) = {fmt_num(lev_f)}; p = {fmt_p(lev_p)}."
        if welch else
        f"Eşit varyans varsayımı karşılanmıştır. Levene varyans homojenliği testi: F({n1-1}, {n2-1}) = {fmt_num(lev_f)}; p = {fmt_p(lev_p)}."
    )
    no, title = tc.next(f"{cv.label}'e Göre {sv.label} Karşılaştırması (Bağımsız Örneklem t-Testi)")
    return make_result(
        "ttest", no, title,
        ["Değişken", "Grup", "n", "x̄ ± SS", "t", "df", "p", "Cohen's d"],
        rows, f"Not. {lev_note} * p < .05",
        t=round(float(t), 3), p=round(float(p), 3), cohens_d=round(d, 3),
        cohens_d_interp=cohens_d_interpretation(d), significant=bool(p < 0.05),
        welch=welch, levene_f=round(float(lev_f), 3), levene_p=round(float(lev_p), 3),
    )

def table_mann_whitney(tc: TableCounter, df: pd.DataFrame, cv: Variable, sv: Variable) -> dict:
    groups = df.groupby(cv.name)[sv.name].apply(lambda x: pd.to_numeric(x, errors="coerce").dropna().tolist())
    g1, g2 = list(groups.iloc[0]), list(groups.iloc[1])
    u, p = mannwhitneyu(g1, g2, alternative="two-sided")
    rows = []
    for name, g in zip(groups.index, [g1, g2]):
        rows.append([sv.label, str(name), str(len(g)), fmt_num(np.median(g)), "", ""])
    rows[0][4] = f"U = {fmt_num(u)}"
    rows[0][5] = fmt_p_display(p)
    no, title = tc.next(f"{cv.label}'e Göre {sv.label} Karşılaştırması (Mann-Whitney U)")
    return make_result(
        "mann_whitney", no, title,
        ["Değişken", "Grup", "n", "Medyan", "U", "p"],
        rows,
        "Not. * p < .05. Non-parametrik test uygulanmıştır.",
        U=round(float(u), 3), p=round(float(p), 3), significant=bool(p < 0.05),
    )

def table_anova(tc: TableCounter, df: pd.DataFrame, cv: Variable, sv: Variable) -> dict:
    groups = df.groupby(cv.name)[sv.name].apply(lambda x: pd.to_numeric(x, errors="coerce").dropna().tolist())
    group_lists = [g for g in groups]
    k = len(group_lists)
    lev_f, lev_p = levene(*group_lists)
    f, p = stats.f_oneway(*group_lists)
    df1, df2 = k - 1, sum(len(g) for g in group_lists) - k
    eta2 = eta_squared(group_lists)
    rows = []
    for name, g in zip(groups.index, group_lists):
        m, sd = np.mean(g), np.std(g, ddof=1)
        ci_lo, ci_hi = m - 1.96 * sd / np.sqrt(len(g)), m + 1.96 * sd / np.sqrt(len(g))
        rows.append([
            str(name), str(len(g)), f"{fmt_num(m)} ± {fmt_num(sd)}",
            f"{fmt_num(ci_lo)}–{fmt_num(ci_hi)}", "", "", "",
        ])
    rows[0][4] = fmt_num(f)
    rows[0][5] = fmt_p_display(p)
    rows[0][6] = fmt_num(eta2)
    lev_note = (
        f"Levene testi: F({k-1}, {sum(len(g) for g in group_lists)-k}) = {fmt_num(lev_f)}; p = {fmt_p(lev_p)}."
    )
    no, title = tc.next(f"{sv.label}'in {cv.label}'e Göre Karşılaştırması (Tek Yönlü ANOVA)")
    return make_result(
        "anova", no, title,
        ["Grup", "n", "x̄ ± SS", "Alt–Üst", "F", "p", "η²"],
        rows, f"Not. {lev_note} Post-hoc: Tukey HSD (anlamlıysa). ** p < .01",
        f=round(float(f), 3), p=round(float(p), 3), eta_squared=round(eta2, 3),
        eta_interp=eta_interpretation(eta2), significant=bool(p < 0.05),
        df1=df1, df2=df2,
    )

def table_tukey(tc: TableCounter, df: pd.DataFrame, cv: Variable, sv: Variable) -> Optional[dict]:
    groups = df.groupby(cv.name)[sv.name].apply(lambda x: pd.to_numeric(x, errors="coerce").dropna().tolist())
    group_lists = [np.array(g) for g in groups]
    names = [str(x) for x in groups.index]
    if len(group_lists) < 2:
        return None
    res = tukey_hsd(*group_lists)
    rows = []
    for i in range(len(names)):
        for j in range(i + 1, len(names)):
            mean_diff = float(np.mean(group_lists[i]) - np.mean(group_lists[j]))
            p = float(res.pvalue[i, j])
            se_val = float(np.std(np.concatenate(group_lists), ddof=1) * np.sqrt(1/len(group_lists[i]) + 1/len(group_lists[j])))
            try:
                ci = res.confidence_interval(alpha=0.05)
                lo, hi = float(ci.low[i, j]), float(ci.high[i, j])
            except Exception:
                lo, hi = mean_diff - 1.96 * se_val, mean_diff + 1.96 * se_val
            rows.append([
                names[i], names[j], fmt_num(mean_diff), fmt_num(se_val),
                fmt_p_display(p), f"{fmt_num(lo)}–{fmt_num(hi)}",
            ])
    if not rows:
        return None
    no, title = tc.next(f"{sv.label} Post-Hoc Tukey HSD Çoklu Karşılaştırma")
    return make_result(
        "tukey", no, title,
        ["(I) Grup", "(J) Grup", "Ort. Fark (I–J)", "Std. Hata", "p", "95% GA (Alt–Üst)"],
        rows, "Not. * p < .05. GA = Güven Aralığı.",
    )

def table_kruskal(tc: TableCounter, df: pd.DataFrame, cv: Variable, sv: Variable) -> dict:
    groups = df.groupby(cv.name)[sv.name].apply(lambda x: pd.to_numeric(x, errors="coerce").dropna().tolist())
    group_lists = [g for g in groups]
    h, p = kruskal(*group_lists)
    rows = []
    for name, g in zip(groups.index, group_lists):
        rows.append([str(name), str(len(g)), fmt_num(np.median(g))])
    no, title = tc.next(f"{sv.label}'in {cv.label}'e Göre Karşılaştırması (Kruskal-Wallis)")
    return make_result(
        "kruskal_wallis", no, title,
        ["Grup", "n", "Medyan"],
        rows + [["H", fmt_num(h), fmt_p_display(p)]],
        "Not. * p < .05. Non-parametrik test uygulanmıştır.",
        H=round(float(h), 3), p=round(float(p), 3), significant=bool(p < 0.05),
    )

def table_correlation_matrix(
    tc: TableCounter, variables: List[Variable], df: pd.DataFrame, norm_map: dict, use_pearson: bool
) -> Optional[dict]:
    names = [v.name for v in variables if v.name in df.columns]
    labels = [v.label for v in variables if v.name in df.columns]
    n_vars = len(names)
    if n_vars < 2:
        return None

    method = "Pearson" if use_pearson else "Spearman"
    corr_fn = stats.pearsonr if use_pearson else spearmanr
    matrix = np.eye(n_vars)
    p_matrix = np.zeros((n_vars, n_vars))
    n_matrix = np.zeros((n_vars, n_vars), dtype=int)

    for i in range(n_vars):
        for j in range(i + 1, n_vars):
            s1 = pd.to_numeric(df[names[i]], errors="coerce").dropna()
            s2 = pd.to_numeric(df[names[j]], errors="coerce").dropna()
            common = s1.index.intersection(s2.index)
            if len(common) < 3:
                continue
            r, p = corr_fn(s1[common], s2[common])
            matrix[i, j] = matrix[j, i] = r
            p_matrix[i, j] = p_matrix[j, i] = p
            n_matrix[i, j] = n_matrix[j, i] = len(common)

    headers = ["Değişken"] + [str(i + 1) for i in range(n_vars)] + ["n"]
    rows = []
    for i, label in enumerate(labels):
        row = [f"{i+1}. {label}"]
        for j in range(n_vars):
            if i == j:
                row.append("—")
            else:
                r, p = matrix[i, j], p_matrix[i, j]
                sym = "ρ" if not use_pearson else "r"
                sign = "−" if r < 0 else ""
                row.append(f"{sign}{fmt_num(abs(r))}{p_stars(p)}")
        row.append(str(int(np.max(n_matrix[i])) if np.max(n_matrix[i]) else len(df)))
        rows.append(row)

    no, title = tc.next(f"Değişkenler Arası {method} Korelasyon Katsayıları")
    star_note = "* p < .05; ** p < .01; *** p < .001"
    return make_result(
        "correlation_matrix", no, title, headers, rows,
        f"Not. {star_note} (çift kuyruklu).",
        method=method, variables=labels,
    )

def table_regression(tc: TableCounter, df: pd.DataFrame, v1: Variable, v2: Variable) -> dict:
    x = pd.to_numeric(df[v1.name], errors="coerce").dropna()
    y = pd.to_numeric(df[v2.name], errors="coerce").dropna()
    common = x.index.intersection(y.index)
    slope, intercept, r, p, se = linregress(x[common], y[common])
    no, title = tc.next(f"{v1.label} → {v2.label} Basit Doğrusal Regresyon")
    rows = [
        ["β (eğim)", fmt_num(slope), "SE", fmt_num(se)],
        ["Kesme (intercept)", fmt_num(intercept), "R²", fmt_num(r ** 2)],
        ["p", fmt_p_display(p), "n", str(len(common))],
    ]
    return make_result(
        "regression", no, title,
        ["Parametre", "Değer", "Parametre", "Değer"],
        rows, "Not. * p < .05.",
        slope=round(float(slope), 3), intercept=round(float(intercept), 3),
        r_squared=round(float(r ** 2), 3), p=round(float(p), 3), se=round(float(se), 3),
        significant=bool(p < 0.05), predictor=v1.label, outcome=v2.label,
    )

def cronbach_analysis(df: pd.DataFrame, columns: List[str], table_no: Optional[int] = None) -> Optional[dict]:
    items_df = df[columns].apply(pd.to_numeric, errors="coerce").dropna()
    k = len(columns)
    if k < 2 or len(items_df) < 3:
        return None
    item_var = items_df.var(axis=0, ddof=1).sum()
    total_var = items_df.sum(axis=1).var(ddof=1)
    if total_var == 0:
        return None
    alpha = float((k / (k - 1)) * (1 - item_var / total_var))

    if alpha >= 0.90:
        interp = "Yüksek güvenilirlik"
    elif alpha >= 0.70:
        interp = "İyi güvenilirlik"
    elif alpha >= 0.60:
        interp = "Kabul edilebilir"
    else:
        interp = "Düşük güvenilirlik"

    if table_no is None:
        tc = TableCounter()
        table_no, title = tc.next("Ölçek Güvenilirlik Analizi (Cronbach α)")
    else:
        title = f"Tablo {table_no}. Ölçek Güvenilirlik Analizi (Cronbach α)"

    return make_result(
        "cronbach", table_no, title,
        ["Madde Sayısı", "Geçerli n", "Cronbach α", "Değerlendirme"],
        [[k, len(items_df), f"{alpha:.3f}", interp]],
        "Not. α = Cronbach alfa iç tutarlılık katsayısı. Kabul edilebilir sınır: α ≥ .70.",
        items=columns, n_items=k, n=int(len(items_df)),
        alpha=round(alpha, 3), interpretation=interp,
    )

def paired_ttest(df: pd.DataFrame, col1: str, col2: str) -> dict:
    s1 = pd.to_numeric(df[col1], errors="coerce")
    s2 = pd.to_numeric(df[col2], errors="coerce")
    mask = s1.notna() & s2.notna()
    s1, s2 = s1[mask], s2[mask]
    diff = (s1 - s2).tolist()
    t, p = ttest_rel(s1, s2)
    sd_d = float(np.std(diff, ddof=1))
    d = float(np.mean(diff) / sd_d) if sd_d > 0 else 0.0
    n = len(s1)
    return {
        "type": "paired_ttest",
        "title": f"{col1} × {col2} - Bağımlı Örneklem t-Testi",
        "headers": ["Değişken", "n", "Ort.1", "Ort.2", "Ort. Fark", "SS Fark", "t", "df", "p", "Cohen's d"],
        "rows": [[f"{col1}–{col2}", str(n), fmt_num(s1.mean()), fmt_num(s2.mean()),
                  fmt_num(np.mean(diff)), fmt_num(np.std(diff, ddof=1)),
                  fmt_num(t), str(n - 1), fmt_p_display(p), fmt_num(d)]],
        "note": "Not. * p < .05.",
        "var1": col1, "var2": col2, "n": n,
        "mean1": round(float(s1.mean()), 2), "mean2": round(float(s2.mean()), 2),
        "mean_diff": round(float(np.mean(diff)), 2),
        "sd_diff": round(float(np.std(diff, ddof=1)), 2),
        "t": round(float(t), 3), "p": round(float(p), 3),
        "cohens_d": round(d, 3), "significant": bool(p < 0.05),
    }

def detect_scale_groups(columns: List[str]) -> Dict[str, List[str]]:
    pattern = re.compile(r"^([a-zA-Z]+)(\d+)$", re.IGNORECASE)
    groups: Dict[str, List[str]] = {}
    for col in columns:
        m = pattern.match(col)
        if m:
            groups.setdefault(m.group(1).lower(), []).append(col)
    return {
        p: sorted(c, key=lambda x: int(re.search(r"\d+", x).group()))
        for p, c in groups.items() if len(c) >= 2
    }

# ── Ana Analiz Akışı ──────────────────────────────────────────────────────────

PRIMARY_GROUPING_KEYS = ("bolum", "cinsiyet", "yas")


def is_primary_grouping(var: Variable) -> bool:
    name = var.name.lower()
    return any(key in name for key in PRIMARY_GROUPING_KEYS)


def generate_plan(df: pd.DataFrame, variables: List[Variable]) -> List[dict]:
    active = [v for v in variables if v.included]
    cat_vars = [v for v in active if v.type == "categorical"]
    cont_vars = [v for v in active if v.type == "continuous"]
    grouping_cat = [v for v in cat_vars if v.role == "grouping"]
    outcome_cat = [v for v in cat_vars if v.role == "outcome"]
    grouping_cont = [v for v in cont_vars if v.role == "grouping"]
    outcome_cont = [v for v in cont_vars if v.role == "outcome"]
    all_cont = outcome_cont + grouping_cont

    tests: List[dict] = []

    if all_cont:
        tests.append({
            "id": "descriptive",
            "type": "descriptive",
            "label": "Tanımlayıcı İstatistikler",
            "detail": f"{len(all_cont)} sürekli değişken: {', '.join(v.label for v in all_cont[:3])}{'...' if len(all_cont) > 3 else ''}",
            "recommended": True,
            "count": 1,
        })

    if all_cont:
        tests.append({
            "id": "normality",
            "type": "normality",
            "label": "Normallik Testi",
            "detail": f"{len(all_cont)} değişken için Shapiro-Wilk / Kolmogorov-Smirnov",
            "recommended": True,
            "count": 1,
        })

    freq_vars = grouping_cat + outcome_cat
    if freq_vars:
        tests.append({
            "id": "frequency",
            "type": "frequency",
            "label": "Frekans Tabloları",
            "detail": f"{len(freq_vars)} kategorik değişken: {', '.join(v.label for v in freq_vars[:3])}{'...' if len(freq_vars) > 3 else ''}",
            "recommended": True,
            "count": len(freq_vars),
        })

    for cv in grouping_cat:
        if cv.name not in df.columns:
            continue
        kare_pairs = [
            f"{cv.label} × {ov.label}"
            for ov in outcome_cat
            if ov.name in df.columns
        ]
        if not kare_pairs:
            continue
        tests.append({
            "id": f"chi_square_{cv.name}",
            "type": "chi_square",
            "label": f"Ki-Kare — {cv.label}",
            "detail": "; ".join(kare_pairs[:3]) + ("..." if len(kare_pairs) > 3 else ""),
            "recommended": is_primary_grouping(cv),
            "count": len(kare_pairs),
        })

    cont_targets = [
        v for v in outcome_cont + grouping_cont
        if v.name in df.columns and is_numeric_continuous(df, v, {})
    ]
    for cv in grouping_cat:
        if cv.name not in df.columns:
            continue
        n_groups = df[cv.name].dropna().nunique()
        if n_groups < 2:
            continue
        test_name = "t-Testi" if n_groups == 2 else "ANOVA"
        pairs = [f"{cv.label} × {sv.label}" for sv in cont_targets]
        if pairs:
            tests.append({
                "id": f"ttest_anova_{cv.name}",
                "type": "ttest_anova",
                "label": f"{cv.label} için {test_name}",
                "detail": "; ".join(pairs[:3]) + ("..." if len(pairs) > 3 else ""),
                "recommended": is_primary_grouping(cv),
                "count": len(pairs),
            })

    corr_vars = [v for v in outcome_cont if is_numeric_continuous(df, v, {})]
    if len(corr_vars) >= 2:
        tests.append({
            "id": "correlation",
            "type": "correlation",
            "label": "Korelasyon Matrisi",
            "detail": f"{len(corr_vars)} değişken: {', '.join(v.label for v in corr_vars)}",
            "recommended": True,
            "count": 1,
        })

    if len(corr_vars) >= 2:
        reg_count = len(corr_vars) * (len(corr_vars) - 1) // 2
        tests.append({
            "id": "regression",
            "type": "regression",
            "label": "Basit Doğrusal Regresyon",
            "detail": f"{reg_count} çift kombinasyon",
            "recommended": False,
            "count": reg_count,
        })

    return tests


def run_analyze(
    df: pd.DataFrame,
    variables: List[Variable],
    active_types: Optional[List[str]] = None,
) -> Tuple[List[dict], dict]:
    def enabled(test_key: str) -> bool:
        if active_types is None:
            return True
        return test_key in active_types

    def enabled_group(group_type: str, item_key: str) -> bool:
        if active_types is None:
            return True
        return group_type in active_types or item_key in active_types
    tc = TableCounter()
    results: List[dict] = []
    active = [v for v in variables if v.included]

    cat_vars = [v for v in active if v.type == "categorical"]
    cont_vars = [v for v in active if v.type == "continuous"]

    grouping_cat = [v for v in cat_vars if v.role == "grouping"]
    outcome_cat = [v for v in cat_vars if v.role == "outcome"]
    grouping_cont = [v for v in cont_vars if v.role == "grouping"]
    outcome_cont = [v for v in cont_vars if v.role == "outcome"]
    all_cont = cont_vars

    norm_map: Dict[str, dict] = {}
    for v in all_cont:
        if v.name in df.columns:
            try:
                norm_map[v.name] = assess_normality(df[v.name])
            except Exception:
                norm_map[v.name] = {"is_parametric": True}

    parametric_labels = [v.label for v in outcome_cont if norm_map.get(v.name, {}).get("is_parametric", True)]
    nonparametric_labels = [v.label for v in outcome_cont if not norm_map.get(v.name, {}).get("is_parametric", True)]

    meta = {
        "n_total": len(df),
        "intro": build_intro(len(df), parametric_labels, nonparametric_labels),
        "norm_map": {v.name: norm_map[v.name] for v in outcome_cont if v.name in norm_map},
    }

    # 1. Tanımlayıcı
    if enabled("descriptive"):
        try:
            r = table_descriptive(tc, outcome_cont, df)
            if r:
                results.append(r)
        except Exception:
            pass

    # 2. Normallik tablosu
    if enabled("normality"):
        try:
            r = table_normality(tc, outcome_cont, df, norm_map)
            if r:
                results.append(r)
        except Exception:
            pass

    # 3. Cronbach
    if enabled("cronbach"):
        for cols in detect_scale_groups(list(df.columns)).values():
            try:
                if all(c in df.columns for c in cols):
                    no, _ = tc.next("Ölçek Güvenilirlik Analizi (Cronbach α)")
                    cb = cronbach_analysis(df, cols, table_no=no)
                    if cb:
                        results.append(cb)
            except Exception:
                pass

    # 4. Frekans
    if not enabled("frequency"):
        pass
    for v in (grouping_cat + outcome_cat if enabled("frequency") else []):
        if v.name in df.columns:
            try:
                results.append(table_frequency(tc, df[v.name], v.label))
            except Exception:
                pass

    # 5. Ki-kare: gruplandırma × sonuç (kategorik)
    for cv in grouping_cat:
        chi_key = f"chi_square_{cv.name}"
        if not enabled_group("chi_square", chi_key):
            continue
        for ov in outcome_cat:
            if cv.name in df.columns and ov.name in df.columns:
                try:
                    results.append(table_chi_square(tc, df, cv, ov))
                except Exception:
                    pass

    # 6. t-test / ANOVA / Mann-Whitney / Kruskal: gruplandırma × (sonuç + ölçüm)
    for cv in grouping_cat:
        ttest_key = f"ttest_anova_{cv.name}"
        if not enabled_group("ttest_anova", ttest_key):
            continue
        for sv in outcome_cont + grouping_cont:
            if cv.name not in df.columns or sv.name not in df.columns:
                continue
            if not is_numeric_continuous(df, sv, norm_map):
                continue
            try:
                n_groups = df[cv.name].dropna().nunique()
                if n_groups == 2:
                    if norm_map.get(sv.name, {}).get("is_parametric", True):
                        results.append(table_ttest(tc, df, cv, sv))
                    else:
                        results.append(table_mann_whitney(tc, df, cv, sv))
                elif n_groups > 2:
                    if norm_map.get(sv.name, {}).get("is_parametric", True):
                        anova_res = table_anova(tc, df, cv, sv)
                        results.append(anova_res)
                        if anova_res.get("significant"):
                            tukey = table_tukey(tc, df, cv, sv)
                            if tukey:
                                results.append(tukey)
                    else:
                        results.append(table_kruskal(tc, df, cv, sv))
            except Exception:
                pass

    # 7. Korelasyon matrisi (kodlanmış kategorikler hariç)
    corr_vars = [v for v in outcome_cont if is_numeric_continuous(df, v, norm_map)]
    if enabled("correlation") and len(corr_vars) >= 2:
        try:
            all_parametric = all(norm_map.get(v.name, {}).get("is_parametric", True) for v in corr_vars)
            r = table_correlation_matrix(tc, corr_vars, df, norm_map, all_parametric)
            if r:
                results.append(r)
        except Exception:
            pass

    # 8. Regresyon: gerçek sürekli sonuç × sonuç (parametrik)
    if enabled("regression"):
        reg_cont = [v for v in outcome_cont if is_numeric_continuous(df, v, norm_map)]
        for i, v1 in enumerate(reg_cont):
            for v2 in reg_cont[i + 1:]:
                if v1.name not in df.columns or v2.name not in df.columns:
                    continue
                if not (norm_map.get(v1.name, {}).get("is_parametric", True) and
                        norm_map.get(v2.name, {}).get("is_parametric", True)):
                    continue
                try:
                    results.append(table_regression(tc, df, v1, v2))
                except Exception:
                    pass

    results = [
        r for r in results
        if r is not None
        and r.get("rows")
        and len(r["rows"]) > 0
        and not all(
            str(cell) in ("—", "nan", "None", "")
            for row in r.get("rows", [])
            for cell in (row[:2] if len(row) >= 2 else row)
        )
    ]

    return results, meta

# ── Word Export ───────────────────────────────────────────────────────────────

def _shade_row(row, color="F2F2F2"):
    for cell in row.cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color)
        cell._tc.get_or_add_tcPr().append(shading)

def add_apa_table(doc: Document, result: dict):
    p = doc.add_paragraph()
    run = p.add_run(result.get("title", ""))
    run.bold = True

    headers = result.get("headers", [])
    rows = result.get("rows", [])
    if not headers or not rows:
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = str(h)
        for r in cell.paragraphs[0].runs:
            r.bold = True

    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            if j < len(table.rows[i + 1].cells):
                table.rows[i + 1].cells[j].text = str(val)
        if i % 2 == 1:
            _shade_row(table.rows[i + 1])

    note_p = doc.add_paragraph()
    note_text = result.get("note", "")
    if note_text.startswith("Not."):
        r1 = note_p.add_run("Not.")
        r1.bold = True
        r1.italic = True
        r2 = note_p.add_run(note_text[4:])
        r2.italic = True
    else:
        r = note_p.add_run(note_text)
        r.italic = True

def build_word_document(results: List[dict], bulgular: Optional[Dict[str, str]] = None, intro: str = "") -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    doc.add_heading("BULGULAR", level=1)
    if intro:
        doc.add_paragraph(intro)
    doc.add_paragraph()

    for i, result in enumerate(results):
        add_apa_table(doc, result)
        doc.add_paragraph()
        key = str(i)
        if bulgular and key in bulgular and bulgular[key]:
            bulgu_p = doc.add_paragraph(bulgular[key])
            bulgu_p.paragraph_format.first_line_indent = Inches(0.5)
        doc.add_paragraph()
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

BULGU_SYSTEM = (
    "AKADEMİK BULGULAR (RESULTS) PRENSİBİ:\n"
    "- Bu program sadece ve sadece bir \"Bulgular (Results)\" metni üretmektedir; "
    "\"Tartışma (Discussion)\" metni YAZMAMAKTADIR.\n"
    "- Bulgular bölümünün görevi sadece tablodaki matematiksel verileri objektif olarak yazıya dökmektir.\n"
    "- Değişken isimlerini (OYS_TOPLAM, SBITO_TOPLAM vb.) tablodaki orijinal kodlarıyla aynen bırakın. "
    "Bu kısaltmaların ne anlama geldiğine dair hiçbir yorumsal tahmin, adlandırma veya genişletme yapmayın.\n"
    "- Tabloda olmayan hiçbir kelimeyi, kavramı veya teorik açıklamayı metne dahil etmeyin. "
    "Sadece sayılar ve kodlar konuşsun.\n\n"
    "Sen akademik bir istatistik uzmanısın. Verilen analiz tablosunu APA 7 formatında, Türkçe, "
    "2-4 cümlelik bulgular paragrafı olarak yaz.\n\n"
    "Kurallar:\n"
    "- p değerlerini APA 7 ile yaz: p = .023, p < .001 (sıfır olmadan)\n"
    "- İstatistikleri parantez içinde ver: [F(3, 296) = 3.090; p = .027]\n"
    "- Anlamlı ve anlamsız sonuçları tablodaki değerlere göre nesnel bildir\n"
    "- Cohen's d veya η² tabloda varsa yalnızca sayısal değeri ve APA etiketini yaz (küçük/orta/büyük)\n"
    "- Türkçe akademik dil kullan\n"
    "- Sadece bulguyu yaz, başlık veya açıklama ekleme"
)

# ── Endpoints ─────────────────────────────────────────────────────────────────

def _prepare_analysis_df(df: pd.DataFrame, variables: List[Variable]) -> pd.DataFrame:
    df = normalize_continuous_columns(df, variables)
    for v in variables:
        if re.match(r"^vki$", v.name, re.I) and v.name in df.columns:
            max_val = df[v.name].max()
            if pd.notna(max_val) and max_val > 1000:
                df[v.name] = pd.to_numeric(df[v.name], errors="coerce")
    return df


@app.post("/convert-spss-table")
async def convert_spss_table_endpoint(req: SpssTableRequest):
    try:
        markdown = convert_spss_table(req.content)
        return {"markdown": markdown, "method": "pandas"}
    except Exception as pandas_err:
        if not ANTHROPIC_API_KEY:
            raise HTTPException(status_code=400, detail=str(pandas_err))
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        msg = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=4000,
            system=SPSS_CONVERT_SYSTEM,
            messages=[{"role": "user", "content": req.content}],
        )
        return {"markdown": msg.content[0].text.strip(), "method": "ai", "pandas_error": str(pandas_err)}


@app.post("/plan")
async def generate_analysis_plan(req: PlanRequest):
    rows = [r.values for r in req.data]
    df = pd.DataFrame(rows)
    df = _prepare_analysis_df(df, req.variables)
    return {"tests": generate_plan(df, req.variables)}


@app.post("/analyze")
async def analyze(req: AnalysisRequest):
    rows = [r.values for r in req.data]
    df = pd.DataFrame(rows)
    df = _prepare_analysis_df(df, req.variables)
    missing_data = missing_data_report(df, req.variables)
    results, meta = run_analyze(df, req.variables, req.active_types)
    return sanitize({"results": results, "missing_data": missing_data, "meta": meta})


@app.post("/analyze/cronbach")
async def analyze_cronbach(req: CronbachRequest):
    import traceback
    try:
        if len(req.columns) < 2:
            raise HTTPException(status_code=400, detail="En az 2 sütun gerekli")
        df = pd.DataFrame([r.values for r in req.data])
        missing = [c for c in req.columns if c not in df.columns]
        if missing:
            raise HTTPException(status_code=400, detail=f"Sütunlar bulunamadı: {missing}")
        for col in req.columns:
            df[col] = pd.to_numeric(df[col].astype(str).str.replace(",", ".", regex=False), errors="coerce")
        result = cronbach_analysis(df, req.columns)
        if result is None:
            raise HTTPException(status_code=400, detail="Cronbach alfa hesaplanamadı (yetersiz veri veya varyans)")
        return sanitize({"result": result})
    except HTTPException:
        raise
    except Exception as e:
        print(traceback.format_exc())
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/detect-scales")
async def detect_scales(req: DetectScalesRequest):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY ayarlanmamış")

    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

    ITEM_PATTERN = re.compile(r"^[a-zA-Z]+_\d+(_ters|_T)?$", re.IGNORECASE)
    item_cols = [c for c in req.columns if ITEM_PATTERN.match(c)]

    if len(item_cols) < 2:
        return {"scales": []}

    prefix_groups: Dict[str, list] = defaultdict(list)
    for col in item_cols:
        prefix = re.match(r"^([a-zA-Z]+)_", col)
        if prefix:
            prefix_groups[prefix.group(1)].append(col)

    valid_groups = {k: v for k, v in prefix_groups.items() if len(v) >= 3}

    if not valid_groups:
        return {"scales": []}

    group_info = "\n".join([
        f"- {prefix}: {', '.join(sorted(cols))}"
        for prefix, cols in valid_groups.items()
    ])

    msg = client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=800,
        system="""Sen akademik ölçek analizi uzmanısın. Verilen madde gruplarını ölçeklere dönüştür.

KURALLAR:
- Her prefix grubu ayrı bir ölçektir
- _ters veya _T ile biten madde varsa o maddenin ters versiyonunu kullan, orijinalini KULLANMA
- Orijinal madde ile _ters versiyonu AYNI ANDA listede olmamalı
- Her ölçeğe Türkçe anlamlı isim ver (oys → OYŞTÖ, neq → GYA veya NEQ, sbito → SBİTO)
- Tüm gruplar için ölçek oluştur, hiçbirini atlama

SADECE JSON döndür:
{
  "scales": [
    {"name": "OYŞTÖ", "items": ["oys_1", "oys_2", "oys_3", "oys_4_ters", "oys_5"]},
    {"name": "GYA", "items": ["neq_1_ters", "neq_2", "neq_3", "neq_4_ters"]}
  ]
}""",
        messages=[{
            "role": "user",
            "content": f"Şu madde gruplarını ölçeklere dönüştür:\n{group_info}",
        }],
    )

    text = msg.content[0].text.strip()
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        try:
            result = json.loads(match.group())
            if not result.get("scales"):
                raise ValueError("Boş sonuç")
            return result
        except Exception:
            pass

    scales = []
    scale_names = {"oys": "OYŞTÖ", "neq": "GYA", "sbito": "SBİTO"}
    for prefix, cols in valid_groups.items():
        ters_cols = {
            c.replace("_ters", "").replace("_T", "")
            for c in cols
            if "_ters" in c.lower() or c.endswith("_T")
        }
        final_items = [c for c in cols if c not in ters_cols]
        name = scale_names.get(prefix.lower(), prefix.upper())
        scales.append({"name": name, "items": final_items})

    return {"scales": scales}


@app.post("/analyze/cronbach-batch")
async def analyze_cronbach_batch(req: CronbachBatchRequest):
    df = pd.DataFrame([r.values for r in req.data])
    results = []
    tc = TableCounter()

    for scale in req.scales:
        name = scale.get("name", "Ölçek")
        items = scale.get("items", [])

        valid_items = []
        for col in items:
            if col in df.columns:
                df[col] = pd.to_numeric(
                    df[col].astype(str).str.replace(",", ".", regex=False),
                    errors="coerce",
                )
                valid_items.append(col)

        if len(valid_items) < 2:
            continue

        try:
            items_df = df[valid_items].dropna()
            k = len(valid_items)
            if len(items_df) < 3:
                continue

            item_vars = items_df.var(axis=0, ddof=1).sum()
            total_var = items_df.sum(axis=1).var(ddof=1)

            if total_var == 0:
                continue

            alpha = float((k / (k - 1)) * (1 - item_vars / total_var))

            if alpha >= 0.90:
                interp = "Yüksek"
            elif alpha >= 0.70:
                interp = "İyi"
            elif alpha >= 0.60:
                interp = "Kabul Edilebilir"
            else:
                interp = "Düşük"

            tno, title = tc.next(f"Ölçek Güvenilirlik Analizi — {name}")

            results.append({
                "type": "cronbach",
                "table_number": tno,
                "title": title,
                "headers": ["Ölçek", "Madde Sayısı", "Geçerli n", "Cronbach α", "Değerlendirme"],
                "rows": [[name, k, len(items_df), f"{alpha:.3f}", interp]],
                "note": "Not. α = Cronbach alfa iç tutarlılık katsayısı. Kabul edilebilir sınır: α ≥ .70.",
                "significant": None,
            })
        except Exception:
            continue

    return sanitize({"results": results})


@app.post("/analyze/paired")
async def analyze_paired(req: PairedRequest):
    df = pd.DataFrame([r.values for r in req.data])
    for col in (req.col1, req.col2):
        if col not in df.columns:
            raise HTTPException(status_code=400, detail=f"Sütun bulunamadı: {col}")
        df[col] = pd.to_numeric(df[col].astype(str).str.replace(",", ".", regex=False), errors="coerce")
    try:
        result = paired_ttest(df, req.col1, req.col2)
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))
    return {"result": result}


CLASSIFY_SYSTEM = """Sen bir veri analisti yardımcısısın. Verilen sütun isimlerini ve örnek değerleri inceleyerek her sütunu sınıflandır.

ÖNEMLİ: Sadece açıkça madde olan sütunları exclude et. Şüphe durumunda categorical veya continuous seç.

EXCLUDE — sadece bunlar:
- Tam olarak: anket_no, id, no, sira, num, serial
- Madde pattern: oys_1, oys_2, neq_1, sbito_6, sbito_6_ters, SBITO_6_T gibi (harf_rakam formatı)

CATEGORICAL — bunlar:
- Metin değerli sütunlar: Kadın/Erkek, Evet/Hayır, bölüm adları
- Sayısal ama az benzersiz değer (≤6): cinsiyet, bolum, md, ed, gd, kh, ik, sk, ak
- _grup, _binary, _kategori, YAS_GRUBU, GYA_RISK_GRUBU, VKI_Kategori gibi

CONTINUOUS — bunlar:
- _TOPLAM, _toplam ile bitenler: OYS_TOPLAM, NEQ_TOPLAM, SBITO_TOPLAM
- Sayısal, geniş aralık: dbf_yas, dbf_boy, dbf_kilo, vki, VKI
- Ölçüm değerleri

SADECE JSON döndür:
{"categorical": ["sütun1"], "continuous": ["sütun2"], "exclude": ["sütun3"]}"""


@app.post("/classify")
async def classify_columns(req: ClassifyRequest):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY ortam değişkeni ayarlanmamış")

    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    col_info = "\n".join([
        f"- {col}: örnek değerler = {req.samples.get(col, [])}"
        for col in req.columns
    ])

    msg = client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=800,
        system=CLASSIFY_SYSTEM,
        messages=[{"role": "user", "content": f"Şu sütunları sınıflandır:\n{col_info}"}],
    )

    text = msg.content[0].text.strip()
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except Exception:
            pass
    return {"categorical": [], "continuous": [], "exclude": []}


RECOMMEND_SYSTEM = """Sen akademik araştırma metodolojisi uzmanısın. Verilen araştırma konusu ve sütun listesine göre hangi değişkenlerin analiz için önemli olduğuna karar ver.

Her sütun için üç kategoriden birini seç:
- "recommended": Bu araştırma için mutlaka analiz edilmeli
- "optional": Analiz edilebilir ama öncelikli değil
- "skip": Bu araştırma için gereksiz veya anlamsız

Gruplandırma değişkenleri (kategorik demografik) için tavsiye ver.
Analiz değişkenleri (ölçek puanları, risk grupları) hepsini "recommended" olarak işaretle.

SADECE JSON döndür:
{
  "recommendations": {
    "sütun_adı": {
      "status": "recommended|optional|skip",
      "reason": "Kısa Türkçe açıklama"
    }
  }
}"""


@app.post("/recommend")
async def recommend_variables(req: RecommendRequest):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY ayarlanmamış")

    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    col_info = "\n".join([
        f"- {col}: örnek değerler = {req.samples.get(col, [])}"
        for col in req.columns
    ])

    msg = client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=1000,
        system=RECOMMEND_SYSTEM,
        messages=[{
            "role": "user",
            "content": f"Araştırma konusu: {req.research_topic}\n\nSütunlar:\n{col_info}",
        }],
    )

    text = msg.content[0].text.strip()
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except Exception:
            pass
    return {"recommendations": {}}


@app.post("/ai/bulgu")
async def ai_bulgu(req: BulguRequest):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY ortam değişkeni ayarlanmamış")
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    msg = client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=500,
        system=BULGU_SYSTEM,
        messages=[{"role": "user", "content": f"Şu analiz tablosunu bulgular bölümü için yaz:\n{req.result}"}],
    )
    return {"bulgu": msg.content[0].text}


@app.post("/export/word")
async def export_word(req: WordExportRequest):
    try:
        doc_bytes = build_word_document(req.results, req.bulgular, req.intro or "")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Word dosyası oluşturulamadı: {str(e)}")
    return StreamingResponse(
        io.BytesIO(doc_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=statai_bulgular.docx"},
    )


@app.get("/")
def root():
    return {"status": "ok", "app": "StatAI"}


def sanitize(obj):
    if isinstance(obj, (float, np.floating)):
        if math.isnan(obj) or math.isinf(obj):
            return None
    if isinstance(obj, dict):
        return {k: sanitize(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [sanitize(i) for i in obj]
    return obj
