"""document_parser — .docx ayrıştırma testleri."""
import io
import sys
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "backend"))

from document_parser import (
    anket_text_from_parse,
    etik_text_from_parse,
    parse_anket_docx,
    parse_etik_kurul_docx,
)


def _docx_bytes(paragraphs: list) -> bytes:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_bytes_with_styles(entries: list) -> bytes:
    """entries: [(text, style_name|None), ...]"""
    doc = Document()
    for entry in entries:
        if isinstance(entry, tuple):
            text, style = entry
            doc.add_paragraph(text, style=style) if style else doc.add_paragraph(text)
        else:
            doc.add_paragraph(entry)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_anket_table_row_items():
    raw = _docx_bytes([
        "Ölçek",
        "1 | Kendimi iyi hissediyorum | 1 2 3 4 5",
        "2 | Enerjim yüksek (R) | 1 2 3 4 5",
    ])
    result = parse_anket_docx(raw)
    items = result["sections"][0]["items"]
    assert len(items) == 2
    assert items[0]["no"] == 1


def test_anket_numbered_without_punctuation():
    raw = _docx_bytes([
        "Yaşam Kalitesi",
        "1 Kendimi iyi hissediyorum",
        "2 Enerjim yüksek",
    ])
    result = parse_anket_docx(raw)
    items = result["sections"][0]["items"]
    assert len(items) == 2
    assert items[0]["text"].startswith("Kendimi")


def test_anket_includes_raw_text():
    raw = _docx_bytes([
        "Bölüm 1: Test",
        "1. Madde bir",
        "2. Madde iki",
    ])
    result = parse_anket_docx(raw)
    assert result.get("raw_text")
    assert "Madde bir" in result["raw_text"]


def test_anket_text_from_parse_uses_raw_when_no_items():
    text = anket_text_from_parse({
        "sections": [{"title": "Kapak", "items": []}],
        "raw_text": "Tam anket metni satir 1\nSatir 2",
    })
    assert "Tam anket metni" in text


def test_etik_text_from_parse_includes_raw():
    text = etik_text_from_parse({
        "hypotheses": ["H1: Fark vardir."],
        "raw_text": "Ek etik kurul metni burada.",
    })
    assert "Fark vardir" in text


def test_anket_numbered_items_parsed():
    raw = _docx_bytes([
        "Bölüm 1: Yaşam Kalitesi",
        "1. Kendimi iyi hissediyorum",
        "2. Enerjim yüksek (R)",
        "Hiç katılmıyorum ... Tamamen katılıyorum",
    ])
    result = parse_anket_docx(raw)
    assert not result.get("parse_error")
    items = result["sections"][0]["items"]
    assert len(items) >= 2
    assert items[0]["no"] == 1
    assert "hissediyorum" in items[0]["text"]
    assert items[1]["reverse_hint"] is True
    assert result["sections"][0]["scale_type"] in ("likert_5", "likert_7", "other")


def test_anket_reverse_hint_from_r_marker():
    raw = _docx_bytes(["1) Bu madde ters (R)"])
    result = parse_anket_docx(raw)
    item = result["sections"][0]["items"][0]
    assert item["reverse_hint"] is True
    assert "(R)" not in item["text"]


def test_anket_heading2_section_title():
    raw = _docx_bytes_with_styles([
        ("Gece Yeme Anketi", "Heading 2"),
        "1. Akşam yemeğinden sonra ne kadar iştahlısınız?",
        "2. Gece uyanıp bir şeyler yer misiniz? (T)",
    ])
    result = parse_anket_docx(raw)
    assert not result.get("parse_error")
    assert result["sections"][0]["title"] == "Gece Yeme Anketi"
    assert len(result["sections"][0]["items"]) == 2
    assert result["sections"][0]["items"][1]["reverse_hint"] is True


def test_etik_kurul_institution_not_aim():
    raw = _docx_bytes([
        "Araştırmanın amacı: Üniversite öğrencilerinde online yemek sipariş tutumlarını incelemektir.",
        "Hacettepe Üniversitesi Etik Kurulu",
        "H1: Cinsiyet ile tutum arasında fark vardır.",
    ])
    result = parse_etik_kurul_docx(raw)
    assert not result.get("parse_error")
    assert result["institution"] == "Hacettepe Üniversitesi Etik Kurulu"
    assert result["aim"]
    assert "online yemek" in result["aim"].lower()
    assert result["institution"] != result["aim"]


def test_etik_kurul_institution_null_when_only_aim():
    raw = _docx_bytes([
        "Araştırmanın amacı: Üniversite öğrencilerinde gece yeme davranışını incelemektir.",
    ])
    result = parse_etik_kurul_docx(raw)
    assert not result.get("parse_error")
    assert result.get("institution") is None


def test_etik_kurul_aim_extracted():
    raw = _docx_bytes([
        "Araştırmanın amacı:",
        "Bu çalışmanın amacı üniversite öğrencilerinde gece yeme davranışını incelemektir.",
    ])
    result = parse_etik_kurul_docx(raw)
    assert not result.get("parse_error")
    assert result["aim"]
    assert "gece yeme" in result["aim"].lower()


def test_etik_kurul_hypothesis_pattern():
    raw = _docx_bytes([
        "H1: Cinsiyet ile OYS puanı arasında anlamlı fark vardır.",
        "H2: Yaş ile NEQ puanı arasında anlamlı ilişki vardır.",
    ])
    result = parse_etik_kurul_docx(raw)
    assert result["hypotheses"]
    assert any("Cinsiyet" in h for h in result["hypotheses"])


def test_etik_ignores_methodology_bullets():
    raw = _docx_bytes([
        "Araştırma Soruları",
        "- Evren ve örneklem seçimi yapılacaktır",
        "- Veri toplama formu uygulanacaktır",
        "H1: Cinsiyet ile OYS puanı arasında anlamlı fark vardır.",
    ])
    result = parse_etik_kurul_docx(raw)
    assert len(result["hypotheses"]) == 1
    assert "Cinsiyet" in result["hypotheses"][0]


def test_empty_docx_returns_parse_error():
    assert parse_anket_docx(b"")["parse_error"] is True
    assert parse_etik_kurul_docx(b"not a zip")["parse_error"] is True


def test_blank_docx_parse_error():
    raw = _docx_bytes([])
    assert parse_anket_docx(raw)["parse_error"] is True
    assert parse_etik_kurul_docx(raw)["parse_error"] is True
