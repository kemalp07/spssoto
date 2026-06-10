from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any, Tuple
import base64
import json
import math
import re
from collections import defaultdict
import pandas as pd
import numpy as np
from scipy import stats
from scipy.stats import (
    shapiro, spearmanr, mannwhitneyu, kruskal, ttest_rel, ttest_ind,
    linregress, kstest, levene, tukey_hsd, skew, kurtosis,
)
import anthropic
import io
from docx import Document
from pypdf import PdfReader
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config import ANTHROPIC_API_KEY, ANTHROPIC_MODEL, BULGU_MODEL

app = FastAPI(title="StatAI - Akademik Analiz API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Modeller ──────────────────────────────────────────────────────────────────

class Variable(BaseModel):
    name: str
    label: str
    type: str  # "continuous" | "categorical"
    role: str = "grouping"  # "grouping" | "outcome"
    categories: Optional[List[str]] = None
    included: bool = True
    scale_min: Optional[float] = None
    scale_max: Optional[float] = None

class DataRow(BaseModel):
    values: dict

class AnalysisRequest(BaseModel):
    variables: List[Variable]
    data: List[DataRow]
    active_types: Optional[List[str]] = None
    enabled_tests: Optional[List[str]] = None
    missing_codes: Optional[List[str]] = None
    scale_info: Optional[dict] = None

class PlanRequest(BaseModel):
    variables: List[Variable]
    data: List[DataRow]
    missing_codes: Optional[List[str]] = None
    research_topic: Optional[str] = None
    use_ai: Optional[bool] = True

class BulguRequest(BaseModel):
    result: Any
    research_topic: Optional[str] = None
    label_map: Optional[Dict[str, str]] = None
    approved_cutoffs: Optional[List[dict]] = None
    scale_info: Optional[dict] = None
    pdf_context: Optional[str] = None


class WordExportRequest(BaseModel):
    results: List[dict]
    bulgular: Optional[Dict[str, str]] = None
    intro: Optional[str] = None
    label_map: Optional[Dict[str, str]] = None

class CronbachRequest(BaseModel):
    columns: List[str]
    data: List[DataRow]

class PairedRequest(BaseModel):
    col1: str
    col2: str
    data: List[DataRow]

class ClassifyRequest(BaseModel):
    columns: List[str]
    samples: Dict[str, List[Any]]
    labels: Optional[Dict[str, str]] = None
    research_topic: Optional[str] = None

class DetectScalesRequest(BaseModel):
    columns: List[str]

class CronbachBatchRequest(BaseModel):
    scales: List[dict]
    data: List[DataRow]

class SpssTableRequest(BaseModel):
    content: str
    auto_bulgu: bool = True

class ScaleMatchRequest(BaseModel):
    scale_names: List[str]
    column_names: List[str]

class ColumnLabel(BaseModel):
    column: str
    label: str

class ExtractContextRequest(BaseModel):
    file_base64: str
    file_type: str
    column_labels: List[ColumnLabel]

# ── SPSS Tablo → Markdown ─────────────────────────────────────────────────────

STAT_COL_RE = re.compile(r"χ|chi|anova|\bf\s*\(|^p$|^p\s", re.I)


def _flatten_columns(columns) -> List[str]:
    if isinstance(columns, pd.MultiIndex):
        out = []
        for i, col in enumerate(columns):
            parts = [str(p).strip() for p in col if str(p).strip() not in ("", "nan", "None")]
            out.append(" ".join(parts) if parts else f"col_{i}")
        return out
    return [str(c).strip() if str(c).strip() else f"col_{i}" for i, c in enumerate(columns)]


def _is_blank(val) -> bool:
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return True
    s = str(val).strip()
    return s in ("", "nan", "None", "NaN", ",")


def _clean_stat_value(val) -> str:
    if _is_blank(val):
        return ""
    s = str(val).strip()
    m = re.search(r"=\s*(.+)$", s)
    return m.group(1).strip() if m else s


def _is_stat_column(col_name: str) -> bool:
    return bool(STAT_COL_RE.search(str(col_name)))


def _clean_spss_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = _flatten_columns(df.columns)

    # Yinelenen sütun başlıklarını benzersizleştir
    seen: Dict[str, int] = {}
    new_cols = []
    for col in df.columns:
        base = str(col)
        if base not in seen:
            seen[base] = 0
            new_cols.append(base)
        else:
            seen[base] += 1
            new_cols.append(f"{base}_{seen[base]}")
    df.columns = new_cols

    # Hayalet / tamamen boş sütunları kaldır
    keep = []
    for col in df.columns:
        if any(not _is_blank(v) for v in df[col]):
            keep.append(col)
    if keep:
        df = df[keep]

    # Birleştirilmiş hücreleri dağıt (forward fill)
    df = df.ffill(axis=0)

    # Test istatistikleri (χ², p, F) — genelde Toplam satırında; tüm satırlara yay
    for col in df.columns:
        if _is_stat_column(col):
            series = df[col].apply(_clean_stat_value)
            series = series.replace("", np.nan)
            df[col] = series.ffill().bfill()

    # Boş kategori adları → Kayıp Veri (ilk sütun + etiket sütunları)
    label_col_re = re.compile(r"grup|kategori|ölçek|değişken|\(i\)|\(j\)", re.I)
    label_cols = [df.columns[0]] if len(df.columns) else []
    for ci, col in enumerate(df.columns):
        if ci > 0 and label_col_re.search(str(col)):
            label_cols.append(col)
    for col in label_cols:
        for idx in df.index:
            if _is_blank(df.at[idx, col]):
                rest = [v for c, v in zip(df.columns, df.loc[idx]) if c != col]
                if any(not _is_blank(v) for v in rest):
                    df.at[idx, col] = "Kayıp Veri"

    # Post-hoc (I)/(J) grupları — (I) birleştirilmişse forward fill
    for col in df.columns:
        cl = str(col).lower()
        if "(i)" in cl or cl.startswith("i grup") or cl == "i":
            df[col] = df[col].ffill()

    # Tamamen boş satırları kaldır
    mask = df.apply(lambda row: all(_is_blank(v) for v in row), axis=1)
    df = df[~mask].reset_index(drop=True)

    return df


def _dataframe_to_markdown(df: pd.DataFrame) -> str:
    cols = [apa_italicize_stats(str(c)) for c in df.columns]
    lines = [
        "| " + " | ".join(cols) + " |",
        "| " + " | ".join([":---"] * len(cols)) + " |",
    ]
    for _, row in df.iterrows():
        cells = []
        for j, v in enumerate(row):
            if _is_blank(v):
                cells.append(format_category_value(v) if j == 0 else "")
            elif j == 0:
                cells.append(format_category_value(v))
            else:
                cells.append(str(v).strip())
        if len(cells) != len(cols):
            continue
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def _parse_spss_input(content: str) -> List[pd.DataFrame]:
    content = content.strip()
    if not content:
        raise ValueError("Boş içerik")
    if "<table" in content.lower():
        dfs = pd.read_html(io.StringIO(content))
    else:
        sep = "\t" if content.count("\t") > content.count(",") else ","
        try:
            dfs = [pd.read_csv(io.StringIO(content), sep=sep, engine="python", on_bad_lines="skip")]
        except Exception:
            dfs = pd.read_html(io.StringIO(content))
    return [df for df in dfs if df is not None and not df.empty]


def convert_spss_table(content: str) -> str:
    parts = []
    for df in _parse_spss_input(content):
        cleaned = _clean_spss_dataframe(df)
        if not cleaned.empty:
            parts.append(_dataframe_to_markdown(cleaned))
    if not parts:
        raise ValueError("Tablo ayrıştırılamadı")
    return "\n\n".join(parts)


SPSS_CONVERT_SYSTEM = """Sen SPSS çıktı dönüştürme uzmanısın. Ham metin, HTML veya kopyala-yapıştır formatındaki SPSS analiz çıktılarını (çapraz tablolar, ANOVA, Ki-Kare, t-Testi vb.) sıfır veri kaybıyla kusursuz Markdown tablosuna dönüştür.

KATI KURALLAR:
1. BİRLEŞTİRİLMİŞ HÜCRELER: χ², p, F değerlerini TÜM satırlara forward fill ile çoğalt.
2. HAYALET SÜTUNLAR: Boş sütunları ve sütun kaymalarını düzelt; header ile satır sütun sayısı eşit olsun.
3. KAYIP VERİ: İsimsiz ama frekanslı satırları "Kayıp Veri" olarak adlandır.
4. POST-HOC: (I) grubu birleştirilmişse altındaki (J) satırlarına (I) adını yaz.

SADECE temizlenmiş Markdown tablolarını döndür. Yorum yapma."""


# ── APA Yardımcıları ─────────────────────────────────────────────────────────

class TableCounter:
    def __init__(self):
        self.n = 0

    def next(self, title_suffix: str) -> Tuple[int, str]:
        self.n += 1
        return self.n, f"Tablo {self.n}. {title_suffix}"

def fmt_num(x: float, decimals: int = 2) -> str:
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return "—"
    return f"{float(x):.{decimals}f}"

def fmt_p(p: float) -> str:
    if p is None:
        return "—"
    p = float(p)
    if p < 0.001:
        return "< .001"
    return f"{p:.3f}".lstrip("0")

def p_stars(p: float) -> str:
    if p is None:
        return ""
    p = float(p)
    if p < 0.001:
        return "***"
    if p < 0.01:
        return "**"
    if p < 0.05:
        return "*"
    return ""

def fmt_p_display(p: float) -> str:
    return f"{fmt_p(p)}{p_stars(p)}"


def fmt_r(x: float, decimals: int = 3) -> str:
    """Korelasyon / η² — baştaki sıfır olmadan (.137)."""
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return "—"
    s = f"{abs(float(x)):.{decimals}f}"
    if s.startswith("0."):
        s = s[1:]
    sign = "−" if float(x) < 0 else ""
    return f"{sign}{s}"


ANTHROPOMETRIC_META: Dict[str, Dict[str, str]] = {
    "kilo": {
        "label": "Vücut Ağırlığı (kg)",
        "possessive": "Vücut Ağırlıklarının",
        "inline": "vücut ağırlığı",
    },
    "boy": {
        "label": "Boy Uzunluğu (cm)",
        "possessive": "Boy Uzunluklarının",
        "inline": "boy uzunluğu",
    },
    "vki": {
        "label": "Beden Kitle İndeksi (BKİ)",
        "possessive": "Beden Kitle İndeksi (BKİ) Değerlerinin",
        "inline": "beden kitle indeksi (BKİ)",
    },
}

def academic_short_label(v: Variable) -> str:
    return format_display_label(v.name, v.label)


def academic_group_target(v: Variable) -> str:
    label = format_display_label(v.name, v.label)
    return f"{label} Gruplarına"


def build_group_comparison_title(cv: Variable, sv: Variable, test_name: str) -> str:
    group_part = academic_group_target(cv)
    sv_label = academic_short_label(sv)
    return (
        f"Katılımcıların {sv_label} Değerlerinin "
        f"{group_part} Göre Karşılaştırılması ({test_name})"
    )


def build_measure_analysis_title(sv: Variable, suffix: str) -> str:
    return f"Katılımcıların {academic_short_label(sv)} Değerlerinin {suffix}"


def apply_academic_text_rules(text: str) -> str:
    if not text:
        return text
    result = _strip_apa_html(str(text))
    replacements = [
        (re.compile(r"\bBMI\b", re.I), "BKİ"),
        (re.compile(r"\bVKİ\b"), "BKİ"),
        (re.compile(r"\bVKI\b", re.I), "BKİ"),
    ]
    for pattern, repl in replacements:
        result = pattern.sub(repl, result)
    return result


def format_display_label(name: str, label: str) -> str:
    if label and label.strip() and label.strip().upper() != name.strip().upper():
        return label.strip()
    cleaned = name.replace("_", " ").strip()
    return cleaned[:1].upper() + cleaned[1:] if cleaned else name


def build_resolved_label_map(custom: Optional[Dict[str, str]] = None) -> Dict[str, str]:
    resolved: Dict[str, str] = {}
    if custom:
        for code, label in custom.items():
            if label and label.strip():
                resolved[code.upper()] = label.strip()
    return resolved


def substitute_variable_codes(text: str, label_map: Dict[str, str]) -> str:
    if not text or not label_map:
        return str(text) if text else ""
    result = _strip_apa_html(str(text))
    for code in sorted(label_map.keys(), key=len, reverse=True):
        result = re.sub(re.escape(code), label_map[code], result, flags=re.I)
    return result


DEFAULT_MISSING_CODES = frozenset({"99", "98", "999", "-99", "9"})


def parse_missing_codes(codes: Optional[List[str]]) -> set:
    if not codes:
        return set(DEFAULT_MISSING_CODES)
    parsed = {str(c).strip() for c in codes if str(c).strip()}
    return parsed or set(DEFAULT_MISSING_CODES)


def matches_missing_code(val, codes: set) -> bool:
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return False
    s = str(val).strip()
    if s in codes:
        return True
    try:
        num = float(s.replace(",", "."))
        if num == int(num) and str(int(num)) in codes:
            return True
    except ValueError:
        pass
    return False


def format_category_value(val) -> str:
    """Boş kategori veya eksik veri kodu → Kayıp Veri."""
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return "Kayıp Veri"
    s = str(val).strip()
    if s in ("", "nan", "None", "NaN", "<NA>", "NaT"):
        return "Kayıp Veri"
    if matches_missing_code(val, DEFAULT_MISSING_CODES):
        return "Kayıp Veri"
    return s


def apply_missing_codes(
    df: pd.DataFrame, variables: List[Variable], codes: Optional[List[str]] = None
) -> pd.DataFrame:
    """SPSS eksik veri kodlarını (99 vb.) NaN'a çevir."""
    code_set = parse_missing_codes(codes)
    df = df.copy()
    for v in variables:
        if v.name not in df.columns:
            continue
        if v.type != "categorical":
            continue
        mask = df[v.name].apply(lambda x: matches_missing_code(x, code_set))
        df.loc[mask, v.name] = np.nan
    return df


def is_missing_value(val, code_set: Optional[set] = None) -> bool:
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return True
    s = str(val).strip()
    if s in ("", "nan", "None", "NaN", "<NA>", "NaT"):
        return True
    codes = code_set if code_set is not None else set(DEFAULT_MISSING_CODES)
    return matches_missing_code(val, codes)


def normalize_categorical_columns(
    df: pd.DataFrame, variables: List[Variable]
) -> pd.DataFrame:
    """Boş / geçersiz metin değerlerini kategorik sütunlarda NaN yap."""
    df = df.copy()
    for v in variables:
        if v.type != "categorical" or v.name not in df.columns:
            continue
        col = df[v.name]
        blank = col.apply(
            lambda x: x is None
            or (isinstance(x, float) and np.isnan(x))
            or str(x).strip() in ("", "nan", "None", "NaN", "<NA>", "NaT")
        )
        df.loc[blank, v.name] = np.nan
    return df


def _infer_valid_categories(series: pd.Series, code_set: set) -> set:
    """Beklenen kategori kümesini frekans profiline göre çıkar; nadir hatalı kodları dışla."""
    valid_s = series.dropna()
    valid_s = valid_s[~valid_s.apply(lambda x: matches_missing_code(x, code_set))]
    if len(valid_s) == 0:
        return set()

    counts = valid_s.value_counts()
    threshold = max(2, int(len(valid_s) * 0.005))

    numeric_map: Dict[float, object] = {}
    for val in counts.index:
        try:
            num = float(str(val).strip().replace(",", "."))
            if num == int(num):
                numeric_map[int(num)] = val
        except ValueError:
            pass

    if numeric_map and len(numeric_map) == len(counts):
        ints = sorted(numeric_map.keys())
        if ints == [1, 2]:
            return {str(numeric_map[i]).strip() for i in ints}
        if 1 in ints and 2 in ints:
            main_count = int(counts[numeric_map[1]]) + int(counts[numeric_map[2]])
            if main_count / len(valid_s) >= 0.95:
                return {str(numeric_map[1]).strip(), str(numeric_map[2]).strip()}
        valid_keys = set()
        for i in ints:
            raw = numeric_map[i]
            if int(counts[raw]) >= threshold:
                valid_keys.add(str(raw).strip())
        return valid_keys

    return {str(v).strip() for v, c in counts.items() if int(c) >= threshold}


def sanitize_invalid_categorical_codes(
    df: pd.DataFrame,
    variables: List[Variable],
    codes: Optional[List[str]] = None,
) -> pd.DataFrame:
    """Beklenen kategori dışındaki hatalı girişleri (örn. Evet/Hayır'da 3) NaN yap."""
    code_set = parse_missing_codes(codes)
    df = df.copy()
    for v in variables:
        if v.type != "categorical" or v.name not in df.columns:
            continue
        if v.categories:
            valid = {str(c).strip() for c in v.categories}
        else:
            valid = _infer_valid_categories(df[v.name], code_set)
        if not valid:
            continue
        mask = df[v.name].notna() & ~df[v.name].apply(
            lambda x: is_missing_value(x, code_set)
            or str(x).strip() in valid
        )
        df.loc[mask, v.name] = np.nan
    return df


def _likert_bounds_from_scale_info(si: dict) -> Tuple[Optional[float], Optional[float]]:
    item_count = si.get("item_count")
    if not item_count:
        return None, None
    try:
        n_items = int(item_count)
    except (TypeError, ValueError):
        return None, None
    if n_items <= 0:
        return None, None
    likert_max = 5
    if si.get("likert"):
        m = re.search(r"(\d+)", str(si["likert"]))
        if m:
            likert_max = int(m.group(1))
    return float(n_items), float(n_items * likert_max)


def infer_theoretical_range(
    variable: Variable,
    df: pd.DataFrame,
    scale_info: Optional[dict] = None,
) -> Optional[str]:
    si = resolve_scale_info(variable, scale_info)
    if si and si.get("min_score") is not None and si.get("max_score") is not None:
        return f"{fmt_num(si['min_score'])} – {fmt_num(si['max_score'])}"
    if si:
        inferred_min, inferred_max = _likert_bounds_from_scale_info(si)
        if inferred_min is not None and inferred_max is not None:
            return f"{fmt_num(inferred_min)} – {fmt_num(inferred_max)}"
    if variable.scale_min is not None and variable.scale_max is not None:
        return f"{fmt_num(variable.scale_min)} – {fmt_num(variable.scale_max)}"

    prefix = re.match(r"^([a-zA-Z]+)", variable.name or "", re.I)
    prefix_key = prefix.group(1).lower() if prefix else ""
    groups = detect_scale_groups(list(df.columns))
    items = groups.get(prefix_key, [])
    if items:
        n_items = len(items)
        likert_max = 5
        if si and si.get("likert"):
            m = re.search(r"(\d+)", str(si["likert"]))
            if m:
                likert_max = int(m.group(1))
        return f"{fmt_num(n_items)} – {fmt_num(n_items * likert_max)}"
    return None


def filter_chi_square_data(
    df: pd.DataFrame,
    v1: Variable,
    v2: Variable,
    codes: Optional[List[str]] = None,
) -> Tuple[pd.DataFrame, int]:
    """Ki-kare matrisine yalnızca her iki değişkende de geçerli değeri olan satırları al."""
    code_set = parse_missing_codes(codes)
    sub = df[[v1.name, v2.name]].copy()
    invalid_mask = sub.apply(
        lambda row: is_missing_value(row[v1.name], code_set)
        or is_missing_value(row[v2.name], code_set)
        or format_category_value(row[v1.name]) == "Kayıp Veri"
        or format_category_value(row[v2.name]) == "Kayıp Veri",
        axis=1,
    )
    excluded = int(invalid_mask.sum())
    return sub.loc[~invalid_mask], excluded


def apa_italicize_stats(text: str) -> str:
    """İstatistiksel sembolleri APA 7 için italik (HTML em) yap."""
    if not text or "<em>" in str(text):
        return str(text) if text is not None else ""
    result = str(text)
    shields: List[str] = []

    def _shield(m: re.Match) -> str:
        shields.append(m.group())
        return f"\x00S{len(shields) - 1}\x00"

    for phrase in ("Kayıp Veri", "Belirtilmeyen"):
        result = re.sub(re.escape(phrase), _shield, result)
    patterns = [
        (r"χ²", "<em>χ²</em>"),
        (r"η²", "<em>η²</em>"),
        (r"ρ", "<em>ρ</em>"),
        (r"n \(%\)", "<em>n</em> (%)"),
        (r"Cohen's d", "Cohen's <em>d</em>"),
        (r"\bdf\b", "<em>df</em>"),
        (r"\bF\b", "<em>F</em>"),
        (r"\bt\b", "<em>t</em>"),
        (r"\bU\b", "<em>U</em>"),
        (r"\bH\b", "<em>H</em>"),
        (r"\bR²\b", "<em>R²</em>"),
        (r"\bR\b", "<em>R</em>"),
        (r"\br\b", "<em>r</em>"),
        (r"\bz\b", "<em>z</em>"),
        (r"\bN\b", "<em>N</em>"),
        (r"(?<![a-zA-Z])n(?![a-zA-Z])", "<em>n</em>"),
        (r"(?<![a-zA-Z])p(?![a-zA-Z])", "<em>p</em>"),
    ]
    for pattern, repl in patterns:
        result = re.sub(pattern, repl, result)
    result = re.sub(r"<em><em>", "<em>", result)
    result = re.sub(r"</em></em>", "</em>", result)
    for i, phrase in enumerate(shields):
        result = result.replace(f"\x00S{i}\x00", phrase)
    return result


def normalize_variable_labels(variables: List[Variable]) -> List[Variable]:
    return [
        v.model_copy(update={"label": format_display_label(v.name, v.label)})
        for v in variables
    ]


def _normalize_scale_key(s: str) -> str:
    tr = str(s).upper()
    for src, dst in [("Ö", "O"), ("Ş", "S"), ("İ", "I"), ("Ü", "U"), ("Ğ", "G"), ("Ç", "C")]:
        tr = tr.replace(src, dst)
    return tr


def resolve_scale_info(variable: Variable, scale_info: Optional[dict]) -> Optional[dict]:
    if not scale_info:
        return None
    if variable.label in scale_info:
        return scale_info[variable.label]
    if variable.name in scale_info:
        return scale_info[variable.name]
    name_upper = variable.name.upper()
    name_prefix = _normalize_scale_key(name_upper.split("_")[0])[:4]
    for short_name, info in scale_info.items():
        if not isinstance(info, dict):
            continue
        short_norm = _normalize_scale_key(short_name)[:4]
        full = info.get("full_name") or ""
        if full and variable.label == full:
            return info
        if name_prefix and short_norm and (
            name_prefix.startswith(short_norm[:3])
            or short_norm.startswith(name_prefix[:3])
        ):
            if "TOPLAM" in name_upper or variable.role == "outcome":
                return info
    return None


def apply_scale_info_to_variables(
    variables: List[Variable], scale_info: Optional[dict]
) -> List[Variable]:
    if not scale_info:
        return variables
    out: List[Variable] = []
    for v in variables:
        si = resolve_scale_info(v, scale_info)
        if not si:
            out.append(v)
            continue
        min_s, max_s = si.get("min_score"), si.get("max_score")
        if min_s is None or max_s is None:
            inferred_min, inferred_max = _likert_bounds_from_scale_info(si)
            if inferred_min is not None:
                min_s, max_s = inferred_min, inferred_max
        if min_s is not None and max_s is not None:
            v = v.model_copy(update={
                "scale_min": float(min_s),
                "scale_max": float(max_s),
            })
        out.append(v)
    return out

def cohens_d_interpretation(d: float) -> str:
    ad = abs(d)
    if ad < 0.20:
        return "önemsiz"
    if ad < 0.50:
        return "küçük"
    if ad < 0.80:
        return "orta"
    return "büyük"

def eta_interpretation(eta2: float) -> str:
    if eta2 < 0.01:
        return "küçük"
    if eta2 <= 0.06:
        return "orta"
    if eta2 <= 0.14:
        return "orta-üst"
    return "büyük"

def cohens_d(g1: list, g2: list) -> float:
    n1, n2 = len(g1), len(g2)
    if n1 < 2 or n2 < 2:
        return 0.0
    v1, v2 = np.var(g1, ddof=1), np.var(g2, ddof=1)
    pooled = np.sqrt(((n1 - 1) * v1 + (n2 - 1) * v2) / (n1 + n2 - 2))
    if pooled == 0:
        return 0.0
    return float((np.mean(g1) - np.mean(g2)) / pooled)

def welch_df(g1: list, g2: list) -> float:
    n1, n2 = len(g1), len(g2)
    v1, v2 = np.var(g1, ddof=1), np.var(g2, ddof=1)
    num = (v1 / n1 + v2 / n2) ** 2
    den = (v1 / n1) ** 2 / (n1 - 1) + (v2 / n2) ** 2 / (n2 - 1)
    return float(num / den) if den else n1 + n2 - 2

def eta_squared(group_lists: list) -> float:
    all_data = np.concatenate([np.array(g) for g in group_lists])
    grand = np.mean(all_data)
    ss_total = np.sum((all_data - grand) ** 2)
    if ss_total == 0:
        return 0.0
    ss_between = sum(len(g) * (np.mean(g) - grand) ** 2 for g in group_lists)
    return float(ss_between / ss_total)

def make_result(
    rtype: str, table_no: int, title: str, headers: list, rows: list, note: str, **extra
) -> dict:
    fmt_headers = [apa_italicize_stats(h) for h in headers]
    fmt_rows = [[apa_italicize_stats(c) for c in row] for row in rows]
    fmt_note = apa_italicize_stats(note) if note else note
    fmt_title = apa_italicize_stats(title)
    return {
        "type": rtype,
        "table_number": table_no,
        "title": fmt_title,
        "headers": fmt_headers,
        "rows": fmt_rows,
        "note": fmt_note,
        **extra,
    }


LABEL_COL_RE = re.compile(r"grup|kategori|ölçek|değişken|\(i\)|\(j\)", re.I)
P_VALUE_COL_RE = re.compile(r"^p$|^p\s|sig|anlaml|asymp", re.I)


def _is_label_column(col_name: str, idx: int) -> bool:
    if idx == 0:
        return True
    return bool(LABEL_COL_RE.search(str(col_name)))


def _normalize_spss_header(h: str) -> str:
    h = str(h).strip()
    h = h.replace("x̄", "M").replace("X̄", "M")
    h = re.sub(r"\bSS\b", "SD", h)
    if h.startswith("col_"):
        return h
    return format_display_label(h, h)


def _apa_format_spss_cell(val, col_name: str) -> str:
    if _is_blank(val):
        return ""
    s = str(val).strip()
    col_l = _strip_apa_html(str(col_name)).lower()

    inner = s
    eq = re.search(r"=\s*(.+)$", s)
    if eq:
        inner = eq.group(1).strip()

    if P_VALUE_COL_RE.search(col_l) or re.match(r"^p\s*=", s, re.I):
        try:
            p_str = inner.replace(",", ".").lstrip("<").strip()
            if p_str.startswith("."):
                p_str = "0" + p_str
            return fmt_p_display(float(p_str))
        except ValueError:
            pass

    num_candidate = inner.replace(",", ".")
    if re.match(r"^-?0\.\d+$", num_candidate):
        sign = "−" if num_candidate.startswith("-") else ""
        return f"{sign}{num_candidate.lstrip('-')[1:]}"
    if re.match(r"^0\.\d+$", num_candidate):
        return num_candidate[1:]

    return s


def _infer_spss_title_suffix(df: pd.DataFrame, index: int) -> str:
    if len(df.columns):
        lead = _normalize_spss_header(df.columns[0])
        if lead and not lead.startswith("col_"):
            return f"{lead} Analiz Sonuçları"
    return f"SPSS Analiz Tablosu {index + 1}"


def _build_spss_note(df: pd.DataFrame) -> str:
    parts = ["Not. * p < .05"]
    cols = " ".join(str(c) for c in df.columns).lower()
    if "η" in cols or "eta" in cols or "r²" in cols or "r2" in cols:
        parts.append("** p < .01")
    return "; ".join(parts) + "."


def _dataframe_to_apa_result(tc: TableCounter, df: pd.DataFrame, title_suffix: str) -> dict:
    headers = [_normalize_spss_header(c) for c in df.columns]
    rows = []
    for _, row in df.iterrows():
        cells = []
        for j, col in enumerate(df.columns):
            val = row[col]
            if _is_label_column(col, j):
                cells.append(format_category_value(val))
            else:
                cells.append(_apa_format_spss_cell(val, col))
        rows.append(cells)
    no, title = tc.next(title_suffix)
    return make_result(
        "spss_import", no, title, headers, rows, _build_spss_note(df),
    )


def _markdown_tables_to_apa_results(markdown: str) -> List[dict]:
    tc = TableCounter()
    results = []
    for block in re.split(r"\n\s*\n", markdown.strip()):
        lines = [ln.strip() for ln in block.splitlines() if "|" in ln]
        if len(lines) < 2:
            continue
        headers = [c.strip() for c in lines[0].strip("|").split("|")]
        data_lines = [ln for ln in lines[2:] if not re.match(r"^\|[\s\-:|]+\|$", ln)]
        rows = [[c.strip() for c in ln.strip("|").split("|")] for ln in data_lines]
        if not headers or not rows:
            continue
        no, title = tc.next("SPSS Analiz Tablosu")
        results.append(make_result(
            "spss_import", no, title, headers, rows,
            "Not. * p < .05.",
        ))
    return results


def convert_spss_to_apa_results(content: str) -> Tuple[List[dict], dict]:
    """Ham SPSS yapıştırmasını hayalet temizliği + APA 7 tablo nesnelerine dönüştür."""
    results = []
    tc = TableCounter()
    for i, df in enumerate(_parse_spss_input(content)):
        cleaned = _clean_spss_dataframe(df)
        if cleaned.empty:
            continue
        suffix = _infer_spss_title_suffix(cleaned, i)
        results.append(_dataframe_to_apa_result(tc, cleaned, suffix))
    if not results:
        raise ValueError("Tablo ayrıştırılamadı")
    return results, {"source": "spss", "intro": "", "table_count": len(results)}


def _parse_llm_json(text: str) -> dict:
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except json.JSONDecodeError:
            pass
    return {}


RESULTS_SYSTEM = """Sen biyoistatistiksel tabloları APA 7 tez formatında Bulgular düz yazısına dönüştüren bir asistansın. Yalnızca tablodaki verilere dayanarak 2-4 cümle Türkçe bulgu yaz; Tartışma yazma; madde işareti/başlık kullanma.

GENEL: Geçmiş zaman (-miştir/-mıştır). Ham kod (OYS_TOPLAM vb.) yasak; sözlükteki akademik adları kullan. kilo→Vücut ağırlığı, boy→Boy uzunluğu, VKI→BKİ. Tablodaki sayıları aynen aktar; yuvarlama yapma. Semboller: x̄, SS, n, %, p, t, F, χ², η², r, U, H, df.

1) TANIMLAYICI (descriptive): "[Ölçek] ortalama puanı x̄ ± SS (n = ..., medyan = ...) olarak hesaplanmıştır." Teorik aralık varsa: "ölçeğin teorik puan aralığı ... – ... arasındadır."

2) NORMALLİK (normality): p < .05 → "dağılımın normal dağılımdan anlamlı sapma gösterdiği saptanmıştır (D = ..., p < .05)." p ≥ .05 → "değişkenin normal dağılım gösterdiği belirlenmiştir (D = ..., p = ...)." Çarpıklık/basıklık verilirken: "Çarpıklık/Std. Hata ve Basıklık/Std. Hata değerlerinin ±2.0 sınırları içinde kalması sebebiyle parametrik testlerin uygulanmasına karar verilmiştir."

3) FREKANS/Kİ-KARE (frequency, chi_square): Frekansta en yüksek n'li kategori önce. Ki-kare p ≥ .05 → anlamlı ilişki/fark yok (χ² = ..., p = ...). p < .05 → anlamlı ilişki var; en baskın yüzdeyi vurgula. Kayıp Veri kategorisini yorumlama; gerekirse "X katılımcıya ait veri eksik olduğundan geçerli n = ... üzerinden analiz yapılmıştır." Yüzdeler valid n üzerinden.

4) ANOVA + TUKEY (anova, tukey): "[Bağımlı] toplam puanlarının [Bağımsız]'e göre tek yönlü ANOVA sonucunda gruplar arasında anlamlı fark saptanmıştır/saptanmamıştır [F = ..., p = ...; η² = ...]." Tukey post-hoc anlamlıysa x̄ BÜYÜK olan grup yüksek yazılır; küçük ortalama asla yüksek raporlanamaz.

5) t-TEST / NON-PARAMETRİK (ttest, mann_whitney, kruskal_wallis): t-testi: "[Değişken] puanlarının ... karşılaştırılması amacıyla bağımsız örneklem t-testi sonucunda ... [t(df) = ..., p = ...; Cohen's d = ...]." Mann-Whitney/Kruskal: medyan ve U veya H; p formatında.

6) KORELASYON (correlation, correlation_matrix): r pozitif/negatif yön + anlamlılık (r = ..., p ...). Büyüklük: zayıf (r < .30), orta (.30–.70), güçlü (r > .70).

7) KESİM NOKTASI: Verilmişse ortalama/frekans yorumuna doğal biçimde göm."""


def _bulgu_context_blob(result: dict) -> str:
    chunks = [
        str(result.get("title", "")),
        str(result.get("type", "")),
        str(result.get("var1", "")),
        str(result.get("var2", "")),
        str(result.get("variable", "")),
    ]
    chunks.extend(str(h) for h in result.get("headers", []))
    for row in result.get("rows", []):
        chunks.extend(str(c) for c in row)
    chunks.append(str(result.get("note", "")))
    return " ".join(chunks).lower()


def _compact_bulgu_result(result: dict) -> dict:
    compact = {
        "type": result.get("type"),
        "table_number": result.get("table_number"),
        "title": result.get("title"),
        "headers": result.get("headers"),
        "rows": result.get("rows"),
        "note": result.get("note"),
    }
    for key in (
        "chi2", "p", "dof", "F", "eta2", "t", "df", "r", "d", "U", "H",
        "significant", "var1", "var2", "variable", "n",
    ):
        if key in result and result[key] is not None:
            compact[key] = result[key]
    return compact


def _filter_label_map(label_map: Optional[Dict[str, str]], blob: str) -> Optional[Dict[str, str]]:
    if not label_map:
        return None
    matched = {
        code: label
        for code, label in label_map.items()
        if code.lower() in blob or (label and label.lower() in blob)
    }
    return matched or None


def _filter_scale_info(scale_info: Optional[dict], blob: str) -> Optional[dict]:
    if not scale_info:
        return None
    matched = {}
    for key, info in scale_info.items():
        if not isinstance(info, dict):
            continue
        full = str(info.get("full_name") or "")
        if (
            key.lower() in blob
            or full.lower() in blob
            or any(tok in blob for tok in key.lower().split() if len(tok) > 2)
        ):
            matched[key] = {
                k: info[k]
                for k in ("full_name", "item_count", "min_score", "max_score", "cutoff", "likert")
                if k in info
            }
    return matched or None


def _filter_cutoffs(cutoffs: Optional[List[dict]], blob: str) -> Optional[List[dict]]:
    if not cutoffs:
        return None
    matched = [
        c for c in cutoffs
        if str(c.get("code", "")).lower() in blob
        or str(c.get("label", "")).lower() in blob
        or str(c.get("scale_name", "")).lower() in blob
    ]
    return matched or None


def build_bulgu_user_message(
    result: dict,
    research_topic: Optional[str] = None,
    label_map: Optional[Dict[str, str]] = None,
    approved_cutoffs: Optional[List[dict]] = None,
    scale_info: Optional[dict] = None,
    pdf_context: Optional[str] = None,
) -> str:
    blob = _bulgu_context_blob(result)
    parts = []
    if research_topic:
        parts.append(f"Konu: {research_topic}")
    labels = _filter_label_map(label_map, blob)
    if labels:
        parts.append("Etiketler: " + json.dumps(labels, ensure_ascii=False))
    cutoffs = _filter_cutoffs(approved_cutoffs, blob)
    if cutoffs:
        parts.append("Kesim: " + json.dumps(cutoffs, ensure_ascii=False))
    scales = _filter_scale_info(scale_info, blob)
    if scales:
        parts.append("Ölçek: " + json.dumps(scales, ensure_ascii=False))
    parts.append("Tablo: " + json.dumps(_compact_bulgu_result(result), ensure_ascii=False))
    message = "\n".join(parts)
    if pdf_context:
        message += f"""

--- ÖLÇEK KAYNAK BİLGİSİ (etik kurul / araştırma raporundan) ---
{pdf_context}
--- Bu bilgiyi bulgularda ölçeği tanıtırken ve yorumlarken kullan.
Birebir kopyalama, kendi cümlelerinle akademik dille yeniden yaz. ---"""
    return message


def _generate_bulgu_text(
    result: dict,
    research_topic: Optional[str] = None,
    label_map: Optional[Dict[str, str]] = None,
    approved_cutoffs: Optional[List[dict]] = None,
    scale_info: Optional[dict] = None,
    pdf_context: Optional[str] = None,
) -> str:
    if not ANTHROPIC_API_KEY:
        return ""
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    msg = client.messages.create(
        model=BULGU_MODEL,
        max_tokens=600,
        system=RESULTS_SYSTEM,
        messages=[{
            "role": "user",
            "content": build_bulgu_user_message(
                result, research_topic, label_map, approved_cutoffs, scale_info, pdf_context
            ),
        }],
    )
    return msg.content[0].text.strip()


def normalize_continuous_columns(df: pd.DataFrame, variables: List[Variable]) -> pd.DataFrame:
    df = df.copy()
    for v in variables:
        if v.type != "continuous" or v.name not in df.columns:
            continue
        df[v.name] = (
            df[v.name].astype(str)
            .str.strip()
            .str.replace(",", ".", regex=False)
        )
        df[v.name] = pd.to_numeric(df[v.name], errors="coerce")
    return df


def is_numeric_continuous(df: pd.DataFrame, v: Variable, norm_map: dict) -> bool:
    """Gerçekten sayısal sürekli mi? Kategorik kodlanmış değilse."""
    if v.name not in df.columns:
        return False
    series = df[v.name].dropna()
    if len(series) < 3:
        return False
    unique_count = series.nunique()
    if unique_count < 10:
        return False
    return True

def missing_data_report(df: pd.DataFrame, variables: List[Variable]) -> list:
    total = len(df)
    report = []
    for v in variables:
        if v.name not in df.columns:
            continue
        col = df[v.name]
        missing = col.isna() | (col.astype(str).str.strip().isin(["", "nan", "None", "NaN"]))
        pct = round(float(missing.sum()) / total * 100, 1) if total else 0.0
        warning = "none"
        if pct > 30:
            warning = "high"
        elif pct > 10:
            warning = "medium"
        report.append({
            "column": v.label, "name": v.name,
            "missing_pct": pct, "missing_n": int(missing.sum()), "warning": warning,
        })
    return report

# ── Normallik ─────────────────────────────────────────────────────────────────

def assess_normality(series: pd.Series) -> dict:
    s = series.dropna().astype(float)
    n = len(s)
    if n < 3:
        return {
            "n": n, "statistic": None, "stat_label": "W", "df": None, "p": None,
            "skewness": None, "skew_se": None, "kurtosis": None, "kurt_se": None,
            "is_parametric": True, "method": "yetersiz_veri",
        }

    sk = float(skew(s))
    ku = float(kurtosis(s, fisher=True))
    se_sk = float(np.sqrt(6 / n))
    se_ku = float(np.sqrt(24 / n))
    skew_ok = abs(sk) <= 2.0
    kurt_ok = abs(ku) <= 2.0

    if n <= 50:
        stat, p = shapiro(s)
        stat_label, method = "W", "Shapiro-Wilk"
    else:
        stat, p = kstest(s, "norm", args=(s.mean(), s.std(ddof=1)))
        stat_label, method = "D", "Kolmogorov-Smirnov"

    p = float(p)
    if p >= 0.05:
        is_parametric = True
    elif skew_ok and kurt_ok:
        is_parametric = True
    else:
        is_parametric = False

    return {
        "n": n,
        "statistic": round(float(stat), 3),
        "stat_label": stat_label,
        "df": n,
        "p": round(p, 3),
        "skewness": round(sk, 2),
        "skew_se": round(se_sk, 2),
        "kurtosis": round(ku, 2),
        "kurt_se": round(se_ku, 2),
        "is_parametric": bool(is_parametric),
        "method": method,
    }

def build_intro(n_total: int) -> str:
    return (
        f"Araştırmanın örneklemini toplam {n_total} katılımcı (N = {n_total}) oluşturmaktadır. "
        f"Analizler öncesinde değişkenlerin dağılım özellikleri incelenmiş; "
        f"normal dağılım varsayımlarının karşılandığı saptanmıştır. "
        f"Bu doğrultuda, araştırmanın amaçları kapsamında parametrik test yöntemlerinin "
        f"uygulanmasına karar verilmiştir."
    )

# ── Tablo Üreticileri ─────────────────────────────────────────────────────────

def table_descriptive(
    tc: TableCounter,
    variables: List[Variable],
    df: pd.DataFrame,
    scale_info: Optional[dict] = None,
) -> Optional[dict]:
    rows = []
    for v in variables:
        if v.name not in df.columns:
            continue
        s = pd.to_numeric(df[v.name], errors="coerce").dropna()
        if len(s) == 0:
            continue
        theory = infer_theoretical_range(v, df, scale_info) or "—"
        rows.append([
            v.label, str(len(s)),
            f"{fmt_num(s.mean())} ± {fmt_num(s.std(ddof=1))}",
            fmt_num(s.median()), f"{fmt_num(s.min())} – {fmt_num(s.max())}", theory,
        ])
    if not rows:
        return None
    no, title = tc.next("Tanımlayıcı İstatistikler")
    return make_result(
        "descriptive", no, title,
        ["Ölçek", "n", "x̄ ± SS", "Medyan", "Min – Maks", "Teorik Aralık"],
        rows, "Not. SS = Standart Sapma.",
    )

def table_normality(tc: TableCounter, variables: List[Variable], df: pd.DataFrame, norm_map: dict) -> Optional[dict]:
    rows = []
    for v in variables:
        if v.name not in norm_map:
            continue
        nm = norm_map[v.name]
        if nm["statistic"] is None:
            continue
        rows.append([
            v.label, f"{nm['stat_label']} = {nm['statistic']}", str(nm["df"]),
            fmt_p_display(nm["p"]),
            f"{nm['skewness']} / {nm['skew_se']}",
            f"{nm['kurtosis']} / {nm['kurt_se']}",
        ])
    if not rows:
        return None
    no, title = tc.next("Normallik Testi Sonuçları")
    return make_result(
        "normality", no, title,
        ["Değişken", "İstatistik", "df", "p", "Çarpıklık / Std. Hata", "Basıklık / Std. Hata"],
        rows,
        "Not. * p < .05; ** p < .01; *** p < .001. Çarpıklık ve basıklık ±2.0 içindeyse normal dağılım varsayımı karşılanmış kabul edilir.",
        norm_map={v.name: norm_map[v.name] for v in variables if v.name in norm_map},
    )

def table_frequency(tc: TableCounter, series: pd.Series, label: str) -> dict:
    total_n = int(len(series))
    valid = series.dropna()
    counts = valid.value_counts()
    rows = []
    missing_n = int(series.isna().sum())
    for val, cnt in counts.items():
        cat = format_category_value(val)
        if cat == "Kayıp Veri":
            missing_n += int(cnt)
            continue
        rows.append([label, cat, str(int(cnt)), ""])
    valid_n = sum(int(cnt) for val, cnt in counts.items() if format_category_value(val) != "Kayıp Veri")
    for i, row in enumerate(rows):
        cnt = int(row[2])
        pct = round(cnt / valid_n * 100, 1) if valid_n else 0
        rows[i][3] = fmt_num(pct, 1)
    if missing_n > 0:
        pct_miss = round(missing_n / total_n * 100, 1) if total_n else 0
        rows.append([label, "Kayıp Veri", str(missing_n), fmt_num(pct_miss, 1)])
    rows.append([label, "Toplam", str(valid_n), "100.0"])
    note = "Not. Değerler frekans (n) ve yüzde (%) olarak verilmiştir."
    if missing_n > 0:
        note += (
            f" Kategori yüzdeleri geçerli örneklem (Valid N={valid_n}) üzerinden hesaplanmıştır;"
            f" analize dahil edilmeyen {missing_n} kayıp veri ayrı gösterilmiştir."
        )
    no, title = tc.next(f"{label} Dağılımı")
    return make_result(
        "frequency", no, title,
        ["Değişken", "Kategori", "n", "%"],
        rows, note,
        variable=label, n=valid_n,
    )

def table_chi_square(
    tc: TableCounter,
    df: pd.DataFrame,
    v1: Variable,
    v2: Variable,
    missing_codes: Optional[List[str]] = None,
) -> dict:
    sub, excluded_n = filter_chi_square_data(df, v1, v2, missing_codes)
    ct = pd.crosstab(sub[v1.name], sub[v2.name])
    valid_cols = [c for c in ct.columns if format_category_value(c) != "Kayıp Veri"]
    valid_rows = [i for i in ct.index if format_category_value(i) != "Kayıp Veri"]
    ct = ct.loc[valid_rows, valid_cols]
    chi2, p, dof, _ = stats.chi2_contingency(ct)
    col_headers = [format_category_value(c) for c in ct.columns]
    rows = []
    for idx in ct.index:
        row_total = int(ct.loc[idx].sum())
        cells = [format_category_value(idx)]
        for col in ct.columns:
            n = int(ct.loc[idx, col])
            pct = round(n / row_total * 100, 1) if row_total else 0
            cells.append(f"{n} ({fmt_num(pct, 1)})")
        cells.append(str(row_total))
        cells.extend(["", ""])
        rows.append(cells)
    valid_n = int(ct.values.sum())
    summary_row = ["Toplam"] + [str(int(ct[c].sum())) for c in ct.columns] + [str(valid_n)]
    summary_row += [f"χ² = {fmt_num(chi2)}", fmt_p_display(p)]
    rows.append(summary_row)
    headers = [v1.label] + col_headers + ["Toplam", "χ²", "p"]
    no, title = tc.next(f"{v1.label} × {v2.label} Dağılımı (Ki-Kare Testi)")
    note = "Not. * p < .05. Değerler n (kişi sayısı) ve satır yüzdesi (%) olarak verilmiştir."
    if excluded_n > 0:
        note += (
            f" Ki-kare analizine dahil edilmeyen {excluded_n} kayıp/hatalı kayıt"
            f" matristen çıkarılmıştır (Valid N={valid_n})."
        )
    return make_result(
        "chi_square", no, title, headers, rows,
        note,
        chi2=round(float(chi2), 3), p=round(float(p), 3), dof=int(dof),
        significant=bool(p < 0.05), var1=v1.label, var2=v2.label,
    )

def table_ttest(tc: TableCounter, df: pd.DataFrame, cv: Variable, sv: Variable) -> dict:
    groups = df.groupby(cv.name)[sv.name].apply(lambda x: pd.to_numeric(x, errors="coerce").dropna().tolist())
    g1, g2 = list(groups.iloc[0]), list(groups.iloc[1])
    n1, n2 = len(g1), len(g2)
    lev_f, lev_p = levene(g1, g2)
    welch = float(lev_p) < 0.05
    t, p = ttest_ind(g1, g2, equal_var=not welch)
    df_t = int(welch_df(g1, g2)) if welch else n1 + n2 - 2
    d = cohens_d(g1, g2)
    rows = []
    for name, g in zip(groups.index, [g1, g2]):
        rows.append([
            sv.label, format_category_value(name), str(len(g)),
            f"{fmt_num(np.mean(g))} ± {fmt_num(np.std(g, ddof=1))}",
            "", "", "", fmt_num(d) if name == groups.index[0] else "",
        ])
    rows[0][4] = fmt_num(t)
    rows[0][5] = str(df_t)
    rows[0][6] = fmt_p_display(p)
    lev_note = (
        f"Welch düzeltmeli t değeri raporlanmıştır. Levene varyans homojenliği testi: F({n1-1}, {n2-1}) = {fmt_num(lev_f)}; p = {fmt_p(lev_p)}."
        if welch else
        f"Eşit varyans varsayımı karşılanmıştır. Levene varyans homojenliği testi: F({n1-1}, {n2-1}) = {fmt_num(lev_f)}; p = {fmt_p(lev_p)}."
    )
    no, title = tc.next(build_group_comparison_title(cv, sv, "Bağımsız Örneklem t-Testi"))
    return make_result(
        "ttest", no, title,
        ["Değişken", "Grup", "n", "x̄ ± SS", "t", "df", "p", "Cohen's d"],
        rows, f"Not. {lev_note} * p < .05",
        t=round(float(t), 3), p=round(float(p), 3), cohens_d=round(d, 3),
        cohens_d_interp=cohens_d_interpretation(d), significant=bool(p < 0.05),
        welch=welch, levene_f=round(float(lev_f), 3), levene_p=round(float(lev_p), 3),
    )

def table_mann_whitney(tc: TableCounter, df: pd.DataFrame, cv: Variable, sv: Variable) -> dict:
    groups = df.groupby(cv.name)[sv.name].apply(lambda x: pd.to_numeric(x, errors="coerce").dropna().tolist())
    g1, g2 = list(groups.iloc[0]), list(groups.iloc[1])
    u, p = mannwhitneyu(g1, g2, alternative="two-sided")
    rows = []
    for name, g in zip(groups.index, [g1, g2]):
        rows.append([sv.label, format_category_value(name), str(len(g)), fmt_num(np.median(g)), "", ""])
    rows[0][4] = f"U = {fmt_num(u)}"
    rows[0][5] = fmt_p_display(p)
    no, title = tc.next(build_group_comparison_title(cv, sv, "Mann-Whitney U"))
    return make_result(
        "mann_whitney", no, title,
        ["Değişken", "Grup", "n", "Medyan", "U", "p"],
        rows,
        "Not. * p < .05. Non-parametrik test uygulanmıştır.",
        U=round(float(u), 3), p=round(float(p), 3), significant=bool(p < 0.05),
    )

def table_anova(tc: TableCounter, df: pd.DataFrame, cv: Variable, sv: Variable) -> dict:
    groups = df.groupby(cv.name)[sv.name].apply(lambda x: pd.to_numeric(x, errors="coerce").dropna().tolist())
    group_lists = [g for g in groups]
    k = len(group_lists)
    lev_f, lev_p = levene(*group_lists)
    f, p = stats.f_oneway(*group_lists)
    df1, df2 = k - 1, sum(len(g) for g in group_lists) - k
    eta2 = eta_squared(group_lists)
    rows = []
    for name, g in zip(groups.index, group_lists):
        m, sd = np.mean(g), np.std(g, ddof=1)
        ci_lo, ci_hi = m - 1.96 * sd / np.sqrt(len(g)), m + 1.96 * sd / np.sqrt(len(g))
        rows.append([
            format_category_value(name), str(len(g)), f"{fmt_num(m)} ± {fmt_num(sd)}",
            f"{fmt_num(ci_lo)}–{fmt_num(ci_hi)}", "", "", "",
        ])
    rows[0][4] = fmt_num(f)
    rows[0][5] = fmt_p_display(p)
    rows[0][6] = fmt_r(eta2)
    lev_note = (
        f"Levene testi: F({k-1}, {sum(len(g) for g in group_lists)-k}) = {fmt_num(lev_f)}; p = {fmt_p(lev_p)}."
    )
    no, title = tc.next(build_group_comparison_title(cv, sv, "Tek Yönlü ANOVA"))
    return make_result(
        "anova", no, title,
        ["Grup", "n", "x̄ ± SS", "Alt–Üst", "F", "p", "η²"],
        rows, f"Not. {lev_note} Post-hoc: Tukey HSD (anlamlıysa). ** p < .01",
        f=round(float(f), 3), p=round(float(p), 3), eta_squared=round(eta2, 3),
        eta_interp=eta_interpretation(eta2), significant=bool(p < 0.05),
        df1=df1, df2=df2,
    )

def table_tukey(tc: TableCounter, df: pd.DataFrame, cv: Variable, sv: Variable) -> Optional[dict]:
    groups = df.groupby(cv.name)[sv.name].apply(lambda x: pd.to_numeric(x, errors="coerce").dropna().tolist())
    group_lists = [np.array(g) for g in groups]
    names = [format_category_value(x) for x in groups.index]
    if len(group_lists) < 2:
        return None
    res = tukey_hsd(*group_lists)
    rows = []
    for i in range(len(names)):
        for j in range(i + 1, len(names)):
            mean_diff = float(np.mean(group_lists[i]) - np.mean(group_lists[j]))
            p = float(res.pvalue[i, j])
            se_val = float(np.std(np.concatenate(group_lists), ddof=1) * np.sqrt(1/len(group_lists[i]) + 1/len(group_lists[j])))
            try:
                ci = res.confidence_interval(alpha=0.05)
                lo, hi = float(ci.low[i, j]), float(ci.high[i, j])
            except Exception:
                lo, hi = mean_diff - 1.96 * se_val, mean_diff + 1.96 * se_val
            rows.append([
                names[i], names[j], fmt_num(mean_diff), fmt_num(se_val),
                fmt_p_display(p), f"{fmt_num(lo)}–{fmt_num(hi)}",
            ])
    if not rows:
        return None
    no, title = tc.next(build_measure_analysis_title(sv, "Post-Hoc Tukey HSD Çoklu Karşılaştırması"))
    return make_result(
        "tukey", no, title,
        ["(I) Grup", "(J) Grup", "Ort. Fark (I–J)", "Std. Hata", "p", "95% GA (Alt–Üst)"],
        rows, "Not. * p < .05. GA = Güven Aralığı.",
    )

def table_kruskal(tc: TableCounter, df: pd.DataFrame, cv: Variable, sv: Variable) -> dict:
    groups = df.groupby(cv.name)[sv.name].apply(lambda x: pd.to_numeric(x, errors="coerce").dropna().tolist())
    group_lists = [g for g in groups]
    h, p = kruskal(*group_lists)
    rows = []
    for name, g in zip(groups.index, group_lists):
        rows.append([format_category_value(name), str(len(g)), fmt_num(np.median(g))])
    no, title = tc.next(build_group_comparison_title(cv, sv, "Kruskal-Wallis"))
    return make_result(
        "kruskal_wallis", no, title,
        ["Grup", "n", "Medyan"],
        rows + [["H", fmt_num(h), fmt_p_display(p)]],
        "Not. * p < .05. Non-parametrik test uygulanmıştır.",
        H=round(float(h), 3), p=round(float(p), 3), significant=bool(p < 0.05),
    )

def table_correlation_matrix(
    tc: TableCounter, variables: List[Variable], df: pd.DataFrame, norm_map: dict, use_pearson: bool
) -> Optional[dict]:
    names = [v.name for v in variables if v.name in df.columns]
    labels = [v.label for v in variables if v.name in df.columns]
    n_vars = len(names)
    if n_vars < 2:
        return None

    method = "Pearson" if use_pearson else "Spearman"
    corr_fn = stats.pearsonr if use_pearson else spearmanr
    matrix = np.eye(n_vars)
    p_matrix = np.zeros((n_vars, n_vars))
    n_matrix = np.zeros((n_vars, n_vars), dtype=int)

    for i in range(n_vars):
        for j in range(i + 1, n_vars):
            s1 = pd.to_numeric(df[names[i]], errors="coerce").dropna()
            s2 = pd.to_numeric(df[names[j]], errors="coerce").dropna()
            common = s1.index.intersection(s2.index)
            if len(common) < 3:
                continue
            r, p = corr_fn(s1[common], s2[common])
            matrix[i, j] = matrix[j, i] = r
            p_matrix[i, j] = p_matrix[j, i] = p
            n_matrix[i, j] = n_matrix[j, i] = len(common)

    headers = ["Değişken"] + [str(i + 1) for i in range(n_vars)] + ["n"]
    rows = []
    for i, label in enumerate(labels):
        row = [f"{i+1}. {label}"]
        for j in range(n_vars):
            if i == j:
                row.append("—")
            else:
                r, p = matrix[i, j], p_matrix[i, j]
                sym = "ρ" if not use_pearson else "r"
                sign = "−" if r < 0 else ""
                row.append(f"{sign}{fmt_r(abs(r))}{p_stars(p)}")
        row.append(str(int(np.max(n_matrix[i])) if np.max(n_matrix[i]) else len(df)))
        rows.append(row)

    no, title = tc.next(f"Değişkenler Arası {method} Korelasyon Katsayıları")
    star_note = "* p < .05; ** p < .01; *** p < .001"
    return make_result(
        "correlation_matrix", no, title, headers, rows,
        f"Not. {star_note} (çift kuyruklu).",
        method=method, variables=labels,
    )

def table_regression(tc: TableCounter, df: pd.DataFrame, v1: Variable, v2: Variable) -> dict:
    x = pd.to_numeric(df[v1.name], errors="coerce").dropna()
    y = pd.to_numeric(df[v2.name], errors="coerce").dropna()
    common = x.index.intersection(y.index)
    slope, intercept, r, p, se = linregress(x[common], y[common])
    no, title = tc.next(f"{v1.label} → {v2.label} Basit Doğrusal Regresyon")
    rows = [
        ["β (eğim)", fmt_num(slope), "SE", fmt_num(se)],
        ["Kesme (intercept)", fmt_num(intercept), "R²", fmt_r(r ** 2)],
        ["p", fmt_p_display(p), "n", str(len(common))],
    ]
    return make_result(
        "regression", no, title,
        ["Parametre", "Değer", "Parametre", "Değer"],
        rows, "Not. * p < .05.",
        slope=round(float(slope), 3), intercept=round(float(intercept), 3),
        r_squared=round(float(r ** 2), 3), p=round(float(p), 3), se=round(float(se), 3),
        significant=bool(p < 0.05), predictor=v1.label, outcome=v2.label,
    )

def cronbach_analysis(df: pd.DataFrame, columns: List[str], table_no: Optional[int] = None) -> Optional[dict]:
    items_df = df[columns].apply(pd.to_numeric, errors="coerce").dropna()
    k = len(columns)
    if k < 2 or len(items_df) < 3:
        return None
    item_var = items_df.var(axis=0, ddof=1).sum()
    total_var = items_df.sum(axis=1).var(ddof=1)
    if total_var == 0:
        return None
    alpha = float((k / (k - 1)) * (1 - item_var / total_var))

    if alpha >= 0.90:
        interp = "Yüksek güvenilirlik"
    elif alpha >= 0.70:
        interp = "İyi güvenilirlik"
    elif alpha >= 0.60:
        interp = "Kabul edilebilir"
    else:
        interp = "Düşük güvenilirlik"

    if table_no is None:
        tc = TableCounter()
        table_no, title = tc.next("Ölçek Güvenilirlik Analizi (Cronbach α)")
    else:
        title = f"Tablo {table_no}. Ölçek Güvenilirlik Analizi (Cronbach α)"

    return make_result(
        "cronbach", table_no, title,
        ["Madde Sayısı", "Geçerli n", "Cronbach α", "Değerlendirme"],
        [[k, len(items_df), f"{alpha:.3f}", interp]],
        "Not. α = Cronbach alfa iç tutarlılık katsayısı. Kabul edilebilir sınır: α ≥ .70.",
        items=columns, n_items=k, n=int(len(items_df)),
        alpha=round(alpha, 3), interpretation=interp,
    )

def paired_ttest(df: pd.DataFrame, col1: str, col2: str) -> dict:
    s1 = pd.to_numeric(df[col1], errors="coerce")
    s2 = pd.to_numeric(df[col2], errors="coerce")
    mask = s1.notna() & s2.notna()
    s1, s2 = s1[mask], s2[mask]
    diff = (s1 - s2).tolist()
    t, p = ttest_rel(s1, s2)
    sd_d = float(np.std(diff, ddof=1))
    d = float(np.mean(diff) / sd_d) if sd_d > 0 else 0.0
    n = len(s1)
    return {
        "type": "paired_ttest",
        "title": f"{col1} × {col2} - Bağımlı Örneklem t-Testi",
        "headers": ["Değişken", "n", "Ort.1", "Ort.2", "Ort. Fark", "SS Fark", "t", "df", "p", "Cohen's d"],
        "rows": [[f"{col1}–{col2}", str(n), fmt_num(s1.mean()), fmt_num(s2.mean()),
                  fmt_num(np.mean(diff)), fmt_num(np.std(diff, ddof=1)),
                  fmt_num(t), str(n - 1), fmt_p_display(p), fmt_num(d)]],
        "note": "Not. * p < .05.",
        "var1": col1, "var2": col2, "n": n,
        "mean1": round(float(s1.mean()), 2), "mean2": round(float(s2.mean()), 2),
        "mean_diff": round(float(np.mean(diff)), 2),
        "sd_diff": round(float(np.std(diff, ddof=1)), 2),
        "t": round(float(t), 3), "p": round(float(p), 3),
        "cohens_d": round(d, 3), "significant": bool(p < 0.05),
    }

def detect_scale_groups(columns: List[str]) -> Dict[str, List[str]]:
    pattern = re.compile(r"^([a-zA-Z]+)(\d+)$", re.IGNORECASE)
    groups: Dict[str, List[str]] = {}
    for col in columns:
        m = pattern.match(col)
        if m:
            groups.setdefault(m.group(1).lower(), []).append(col)
    return {
        p: sorted(c, key=lambda x: int(re.search(r"\d+", x).group()))
        for p, c in groups.items() if len(c) >= 2
    }

# ── Ana Analiz Akışı ──────────────────────────────────────────────────────────

PRIMARY_GROUPING_KEYS = ("bolum", "cinsiyet", "yas")


def is_primary_grouping(var: Variable) -> bool:
    name = var.name.lower()
    return any(key in name for key in PRIMARY_GROUPING_KEYS)


DEMO_LABEL_KEYWORDS = re.compile(
    r"\b(yaş|yas|boy|kilo|ağırlık|agirlik|beden|bmi|bki|vki|"
    r"height|weight|age)\b", re.I
)

SCALE_SCORE_RE = re.compile(
    r"_toplam$|_puan$|_skor$|_score$|_total$|_sum$", re.I
)


def _is_demographic_continuous(v: Variable) -> bool:
    """Sürekli ama demografik olan değişkenleri etiket/kolon adından tespit et."""
    text = f"{v.name} {v.label or ''}".lower()
    return bool(DEMO_LABEL_KEYWORDS.search(text))


def generate_plan(df: pd.DataFrame, variables: List[Variable]) -> List[dict]:
    active = [v for v in variables if v.included]
    cat_vars = [v for v in active if v.type == "categorical"]
    cont_vars = [v for v in active if v.type == "continuous"]
    grouping_cat = [v for v in cat_vars if v.role == "grouping"]
    outcome_cat = [v for v in cat_vars if v.role == "outcome"]
    grouping_cont = [v for v in cont_vars if v.role == "grouping"]
    outcome_cont = [v for v in cont_vars if v.role == "outcome"]
    all_cont = outcome_cont + grouping_cont

    tests: List[dict] = []

    if all_cont:
        tests.append({
            "id": "descriptive",
            "type": "descriptive",
            "label": "Tanımlayıcı İstatistikler",
            "detail": f"{len(all_cont)} sürekli değişken: {', '.join(v.label for v in all_cont[:3])}{'...' if len(all_cont) > 3 else ''}",
            "recommended": True,
            "count": 1,
        })

    if all_cont:
        tests.append({
            "id": "normality",
            "type": "normality",
            "label": "Normallik Testi",
            "detail": f"{len(all_cont)} değişken için Shapiro-Wilk / Kolmogorov-Smirnov",
            "recommended": True,
            "count": 1,
        })

    freq_vars = grouping_cat + outcome_cat
    if freq_vars:
        tests.append({
            "id": "frequency",
            "type": "frequency",
            "label": "Frekans Tabloları",
            "detail": f"{len(freq_vars)} kategorik değişken: {', '.join(v.label for v in freq_vars[:3])}{'...' if len(freq_vars) > 3 else ''}",
            "recommended": True,
            "count": len(freq_vars),
        })

    for cv in grouping_cat:
        if cv.name not in df.columns:
            continue
        kare_pairs = [
            f"{cv.label} × {ov.label}"
            for ov in outcome_cat
            if ov.name in df.columns
            and not SCALE_SCORE_RE.search(ov.name)
        ]
        if not kare_pairs:
            continue
        tests.append({
            "id": f"chi_square_{cv.name}",
            "type": "chi_square",
            "label": f"Ki-Kare — {cv.label}",
            "detail": "; ".join(kare_pairs[:3]) + ("..." if len(kare_pairs) > 3 else ""),
            "recommended": is_primary_grouping(cv),
            "count": len(kare_pairs),
        })

    cont_targets = [
        v for v in outcome_cont + grouping_cont
        if v.name in df.columns
        and is_numeric_continuous(df, v, {})
        and not _is_demographic_continuous(v)
    ]
    for cv in grouping_cat:
        if cv.name not in df.columns:
            continue
        n_groups = df[cv.name].dropna().nunique()
        if n_groups < 2:
            continue
        test_name = "t-Testi" if n_groups == 2 else "ANOVA"
        pairs = [f"{cv.label} × {sv.label}" for sv in cont_targets]
        if pairs:
            tests.append({
                "id": f"ttest_anova_{cv.name}",
                "type": "ttest_anova",
                "label": f"{cv.label} için {test_name}",
                "detail": "; ".join(pairs[:3]) + ("..." if len(pairs) > 3 else ""),
                "recommended": is_primary_grouping(cv),
                "count": len(pairs),
            })

    corr_vars = [v for v in outcome_cont if is_numeric_continuous(df, v, {})]
    if len(corr_vars) >= 2:
        tests.append({
            "id": "correlation",
            "type": "correlation",
            "label": "Korelasyon Matrisi",
            "detail": f"{len(corr_vars)} değişken: {', '.join(v.label for v in corr_vars)}",
            "recommended": True,
            "count": 1,
        })

    if len(corr_vars) >= 2:
        reg_count = len(corr_vars) * (len(corr_vars) - 1) // 2
        tests.append({
            "id": "regression",
            "type": "regression",
            "label": "Basit Doğrusal Regresyon",
            "detail": f"{reg_count} çift kombinasyon",
            "recommended": False,
            "count": reg_count,
        })

    return tests


PLAN_SYSTEM = """
Sen akademik araştırma metodolojisi uzmanısın.
Verilen değişkenlere göre istatistiksel analiz planı öner.

KURALLAR:
- Normallik sağlanmışsa parametrik, sağlanmamışsa non-parametrik test öner
- 2 grup → t-testi veya Mann-Whitney U
- 3+ grup → ANOVA veya Kruskal-Wallis
- İki kategorik → Ki-kare
- İki sürekli → Korelasyon (Pearson/Spearman)
- Aynı gruplandırma değişkeni için tüm outcome'ları BİRLEŞTİR
  (3 ölçek için cinsiyet → tek tablo, 3 ayrı değil)
- Demografik sürekli değişkenler (yaş, boy, kilo) için
  gruplandırma testleri YAPMA, sadece tanımlayıcıda göster
- Araştırma konusuyla zayıf ilişkili testleri "optional" işaretle

SADECE JSON döndür:
{
  "tests": [
    {
      "id": "ttest_cinsiyet",
      "type": "ttest_anova",
      "label": "Ölçek Puanlarının Cinsiyete Göre Karşılaştırılması",
      "variables": ["OYS_TOPLAM", "NEQ_TOPLAM", "SBITO_TOPLAM"],
      "grouping": "dbf_cinsiyet",
      "test_name": "Bağımsız Örneklem t-Testi",
      "reason": "2 grup, normallik sağlandı",
      "recommended": true,
      "count": 1
    }
  ],
  "notes": "Kısa genel metodoloji notu"
}
"""


def _normalize_ai_plan_tests(tests: List[dict]) -> List[dict]:
    normalized = []
    for t in tests:
        item = dict(t)
        if "detail" not in item:
            parts = []
            if item.get("test_name"):
                parts.append(str(item["test_name"]))
            if item.get("reason"):
                parts.append(str(item["reason"]))
            if item.get("variables"):
                parts.append(", ".join(str(v) for v in item["variables"]))
            item["detail"] = " — ".join(parts) if parts else str(item.get("label", ""))
        if "count" not in item:
            item["count"] = len(item.get("variables") or []) or 1
        if "recommended" not in item:
            item["recommended"] = True
        normalized.append(item)
    return normalized


async def generate_plan_ai(
    variables: List[Variable],
    research_topic: str,
    norm_results: Optional[dict] = None,
) -> Optional[List[dict]]:
    if not ANTHROPIC_API_KEY:
        return None

    grouping_vars = [v for v in variables if v.included and v.role == "grouping"]
    outcome_vars = [v for v in variables if v.included and v.role == "outcome"]

    var_summary = []
    for v in grouping_vars + outcome_vars:
        normal = None
        if norm_results and v.name in norm_results:
            normal = norm_results[v.name].get("normal")
        var_summary.append({
            "name": v.name,
            "label": v.label,
            "type": v.type,
            "role": v.role,
            "normal": normal,
        })

    user_msg = f"""Araştırma konusu: {research_topic or 'Belirtilmemiş'}

Değişkenler:
{json.dumps(var_summary, ensure_ascii=False, indent=2)}

Bu değişkenler için istatistiksel analiz planı oluştur."""

    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        msg = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=1000,
            system=PLAN_SYSTEM,
            messages=[{"role": "user", "content": user_msg}],
        )
        parsed = _parse_llm_json(msg.content[0].text.strip())
        tests = parsed.get("tests", [])
        return _normalize_ai_plan_tests(tests) if tests else None
    except Exception:
        return None


def _test_enabled(test_id: str, enabled: Optional[List[str]]) -> bool:
    if enabled is None:
        return True
    return test_id in enabled


def run_analyze(
    df: pd.DataFrame,
    variables: List[Variable],
    active_types: Optional[List[str]] = None,
    enabled_tests: Optional[List[str]] = None,
    scale_info: Optional[dict] = None,
    missing_codes: Optional[List[str]] = None,
) -> Tuple[List[dict], dict]:
    def enabled(test_key: str) -> bool:
        if enabled_tests is not None:
            return _test_enabled(test_key, enabled_tests)
        if active_types is None:
            return True
        return test_key in active_types

    def enabled_group(group_type: str, item_key: str) -> bool:
        if enabled_tests is not None:
            return (
                _test_enabled(item_key, enabled_tests)
                or _test_enabled(group_type, enabled_tests)
            )
        if active_types is None:
            return True
        return group_type in active_types or item_key in active_types
    tc = TableCounter()
    results: List[dict] = []
    active = [v for v in variables if v.included]

    cat_vars = [v for v in active if v.type == "categorical"]
    cont_vars = [v for v in active if v.type == "continuous"]

    grouping_cat = [v for v in cat_vars if v.role == "grouping"]
    outcome_cat = [v for v in cat_vars if v.role == "outcome"]
    grouping_cont = [v for v in cont_vars if v.role == "grouping"]
    outcome_cont = [v for v in cont_vars if v.role == "outcome"]
    all_cont = cont_vars

    norm_map: Dict[str, dict] = {}
    for v in all_cont:
        if v.name in df.columns:
            try:
                norm_map[v.name] = assess_normality(df[v.name])
            except Exception:
                norm_map[v.name] = {"is_parametric": True}

    meta = {
        "n_total": len(df),
        "intro": build_intro(len(df)),
        "norm_map": {v.name: norm_map[v.name] for v in outcome_cont if v.name in norm_map},
    }

    # 1. Tanımlayıcı
    if enabled("descriptive"):
        try:
            r = table_descriptive(tc, outcome_cont, df, scale_info)
            if r:
                results.append(r)
        except Exception:
            pass

    # 2. Normallik tablosu
    if enabled("normality"):
        try:
            r = table_normality(tc, outcome_cont, df, norm_map)
            if r:
                results.append(r)
        except Exception:
            pass

    # 3. Cronbach
    if enabled("cronbach"):
        for cols in detect_scale_groups(list(df.columns)).values():
            try:
                if all(c in df.columns for c in cols):
                    no, _ = tc.next("Ölçek Güvenilirlik Analizi (Cronbach α)")
                    cb = cronbach_analysis(df, cols, table_no=no)
                    if cb:
                        results.append(cb)
            except Exception:
                pass

    # 4. Frekans
    if not enabled("frequency"):
        pass
    for v in (grouping_cat + outcome_cat if enabled("frequency") else []):
        if v.name in df.columns:
            try:
                results.append(table_frequency(tc, df[v.name], v.label))
            except Exception:
                pass

    # 5. Ki-kare: gruplandırma × sonuç (kategorik)
    for cv in grouping_cat:
        chi_key = f"chi_square_{cv.name}"
        if not enabled_group("chi_square", chi_key):
            continue
        for ov in outcome_cat:
            if cv.name in df.columns and ov.name in df.columns:
                try:
                    results.append(table_chi_square(tc, df, cv, ov, missing_codes))
                except Exception:
                    pass

    # 6. t-test / ANOVA / Mann-Whitney / Kruskal: gruplandırma × (sonuç + ölçüm)
    for cv in grouping_cat:
        ttest_key = f"ttest_anova_{cv.name}"
        if not enabled_group("ttest_anova", ttest_key):
            continue
        for sv in outcome_cont + grouping_cont:
            if cv.name not in df.columns or sv.name not in df.columns:
                continue
            if not is_numeric_continuous(df, sv, norm_map):
                continue
            try:
                n_groups = df[cv.name].dropna().nunique()
                if n_groups == 2:
                    if norm_map.get(sv.name, {}).get("is_parametric", True):
                        results.append(table_ttest(tc, df, cv, sv))
                    else:
                        results.append(table_mann_whitney(tc, df, cv, sv))
                elif n_groups > 2:
                    if norm_map.get(sv.name, {}).get("is_parametric", True):
                        anova_res = table_anova(tc, df, cv, sv)
                        results.append(anova_res)
                        if anova_res.get("significant"):
                            tukey = table_tukey(tc, df, cv, sv)
                            if tukey:
                                results.append(tukey)
                    else:
                        results.append(table_kruskal(tc, df, cv, sv))
            except Exception:
                pass

    # 7. Korelasyon matrisi (kodlanmış kategorikler hariç)
    corr_vars = [v for v in outcome_cont if is_numeric_continuous(df, v, norm_map)]
    if enabled("correlation") and len(corr_vars) >= 2:
        try:
            all_parametric = all(norm_map.get(v.name, {}).get("is_parametric", True) for v in corr_vars)
            r = table_correlation_matrix(tc, corr_vars, df, norm_map, all_parametric)
            if r:
                results.append(r)
        except Exception:
            pass

    # 8. Regresyon: gerçek sürekli sonuç × sonuç (parametrik)
    if enabled("regression"):
        reg_cont = [v for v in outcome_cont if is_numeric_continuous(df, v, norm_map)]
        for i, v1 in enumerate(reg_cont):
            for v2 in reg_cont[i + 1:]:
                if v1.name not in df.columns or v2.name not in df.columns:
                    continue
                if not (norm_map.get(v1.name, {}).get("is_parametric", True) and
                        norm_map.get(v2.name, {}).get("is_parametric", True)):
                    continue
                try:
                    results.append(table_regression(tc, df, v1, v2))
                except Exception:
                    pass

    results = [
        r for r in results
        if r is not None
        and r.get("rows")
        and len(r["rows"]) > 0
        and not all(
            str(cell) in ("—", "nan", "None", "")
            for row in r.get("rows", [])
            for cell in (row[:2] if len(row) >= 2 else row)
        )
    ]

    return results, meta

# ── Word Export (APA 7 görsel tablo kuralları) ────────────────────────────────

APA_BORDER_SZ = 4  # 0.50 pt (Word OOXML: sekizde biri punto)

_WORD_M_MARK = "\x00M\x00"
_WORD_SD_MARK = "\x00SD\x00"
_WORD_STAT_RE = re.compile(
    r"Cohen's d|"
    r"\x00M\x00|\x00SD\x00|"
    r"\bdf\b|"
    r"\bF\b|\bt\b|\bU\b|\bH\b|\bn\b|\bp\b|\bd\b"
)


def _strip_apa_html(text: str) -> str:
    return re.sub(r"</?em>", "", str(text))


def _prepare_word_text(text: str) -> str:
    """Word export: x̄→M, SS→SD; HTML italik etiketlerini kaldır."""
    text = _strip_apa_html(str(text))
    text = text.replace("x̄", _WORD_M_MARK)
    text = re.sub(r"\bSS\b", _WORD_SD_MARK, text)
    return text


def _split_table_title(title: str) -> Tuple[str, str]:
    """'Tablo 1. Açıklama' → ('Tablo 1', 'Açıklama')."""
    title = str(title)
    plain = _strip_apa_html(title)
    m = re.match(r"^(Tablo\s+\d+)\.\s*(.*)$", plain, re.I)
    if not m:
        return title, ""
    cap_match = re.match(r"^Tablo\s+\d+\.\s*(.*)$", title, re.I)
    caption = cap_match.group(1).strip() if cap_match else m.group(2)
    return m.group(1), caption


def _set_cell_border(cell, **edges):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:tcBorders"):
            tcPr.remove(child)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        spec = edges.get(edge)
        if spec is None:
            continue
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), spec.get("val", "nil"))
        if spec.get("val") == "single":
            element.set(qn("w:sz"), str(spec.get("sz", APA_BORDER_SZ)))
            element.set(qn("w:color"), spec.get("color", "000000"))
            element.set(qn("w:space"), "0")
        tcBorders.append(element)
    tcPr.append(tcBorders)


def _clear_table_level_borders(table):
    """Tablo düzeyinde tüm kenarlıkları kapat; çizgiler yalnızca hücre düzeyinde."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for child in list(tblPr):
        if child.tag == qn("w:tblBorders"):
            tblPr.remove(child)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), "nil")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _apply_apa_table_borders(table):
    """
    APA 7 yatay çizgi yerleşimi (tam 3 çizgi):
      1. Başlık satırı üstü  → header hücre top
      2. Başlık / veri ayrımı → header hücre bottom
      3. Son veri satırı altı → son satır hücre bottom
    Dikey çizgi yok. Not metni tablo dışında (ayrı paragraf).
    """
    _clear_table_level_borders(table)

    nil = {"val": "nil"}
    single = {"val": "single", "sz": APA_BORDER_SZ, "color": "000000"}
    n_rows = len(table.rows)
    last_idx = n_rows - 1

    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if r_idx == 0:
                _set_cell_border(cell, top=single, left=nil, right=nil, bottom=single)
            elif r_idx == last_idx:
                _set_cell_border(cell, top=nil, left=nil, right=nil, bottom=single)
            else:
                _set_cell_border(cell, top=nil, left=nil, right=nil, bottom=nil)


def _add_word_runs(paragraph, text: str, bold: bool = False, force_italic: bool = False):
    """Word hücreleri: x̄→M, SS→SD; n,p,df,F,t,d,H,U italik."""
    text = _prepare_word_text(text)
    pos = 0
    for match in _WORD_STAT_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            run.bold = bold
            if force_italic:
                run.italic = True
        token = match.group()
        if token == "Cohen's d":
            r1 = paragraph.add_run("Cohen's ")
            r1.bold = bold
            if force_italic:
                r1.italic = True
            r2 = paragraph.add_run("d")
            r2.italic = True
            r2.bold = bold
        elif token == _WORD_M_MARK:
            run = paragraph.add_run("M")
            run.italic = True
            run.bold = bold
        elif token == _WORD_SD_MARK:
            run = paragraph.add_run("SD")
            run.italic = True
            run.bold = bold
        else:
            run = paragraph.add_run(token)
            run.italic = True
            run.bold = bold
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = bold
        if force_italic:
            run.italic = True


def _add_apa_table_title(doc: Document, title: str):
    num, caption = _split_table_title(title)
    p_num = doc.add_paragraph()
    run = p_num.add_run(num)
    run.bold = True
    if caption:
        p_cap = doc.add_paragraph()
        _add_word_runs(p_cap, caption, force_italic=True)


def _refine_word_header(header: str) -> str:
    """Ki-kare sütunlarından gereksiz 'n (%)' ekini kaldır."""
    h = _strip_apa_html(str(header))
    return re.sub(r"\s+n\s*\(%\)\s*$", "", h, flags=re.I).strip()


def polish_result_for_word(result: dict, label_map: Optional[Dict[str, str]] = None) -> dict:
    """Word export öncesi başlık/sütun/not rafine etme."""
    polished = dict(result)
    resolved = build_resolved_label_map(label_map)

    title = str(polished.get("title", ""))
    num, caption = _split_table_title(title)
    if caption:
        caption = apply_academic_text_rules(substitute_variable_codes(caption, resolved))
        polished["title"] = f"{num}. {caption}"
    elif resolved:
        polished["title"] = apply_academic_text_rules(substitute_variable_codes(title, resolved))

    headers = [_refine_word_header(h) for h in polished.get("headers", [])]
    if resolved:
        headers = [apply_academic_text_rules(substitute_variable_codes(h, resolved)) for h in headers]
    polished["headers"] = headers

    note = str(polished.get("note", ""))
    if resolved:
        note = apply_academic_text_rules(substitute_variable_codes(note, resolved))
    polished["note"] = note
    return polished


def _add_apa_note(paragraph, note_text: str):
    """Not. yalnızca italik; gövde düz, p/F/t/n/d istatistik sembolleri italik."""
    note_text = str(note_text).strip()
    if note_text.startswith("Not."):
        r = paragraph.add_run("Not.")
        r.italic = True
        _add_word_runs(paragraph, note_text[4:].lstrip())
    elif note_text.startswith("Note."):
        r = paragraph.add_run("Note.")
        r.italic = True
        _add_word_runs(paragraph, note_text[5:].lstrip())
    else:
        _add_word_runs(paragraph, note_text)


def add_apa_table(doc: Document, result: dict, label_map: Optional[Dict[str, str]] = None):
    result = polish_result_for_word(result, label_map)
    _add_apa_table_title(doc, result.get("title", ""))

    headers = result.get("headers", [])
    rows = result.get("rows", [])
    if not headers or not rows:
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Normal"
    except Exception:
        pass

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        _add_word_runs(cell.paragraphs[0], h, bold=True)

    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            if j < len(table.rows[i + 1].cells):
                cell = table.rows[i + 1].cells[j]
                cell.text = ""
                _add_word_runs(cell.paragraphs[0], str(val))

    _apply_apa_table_borders(table)

    # Not: alt çizginin altında, tablo gövdesinin dışında
    note_text = result.get("note", "")
    if note_text:
        note_p = doc.add_paragraph()
        note_p.paragraph_format.space_before = Pt(6)
        _add_apa_note(note_p, note_text)

def build_word_document(
    results: List[dict],
    bulgular: Optional[Dict[str, str]] = None,
    intro: str = "",
    label_map: Optional[Dict[str, str]] = None,
) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    doc.add_heading("BULGULAR", level=1)
    if intro:
        doc.add_paragraph(intro)
    doc.add_paragraph()

    for i, result in enumerate(results):
        add_apa_table(doc, result, label_map)
        doc.add_paragraph()
        key = str(i)
        if bulgular and key in bulgular and bulgular[key]:
            bulgu_p = doc.add_paragraph(bulgular[key])
            bulgu_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bulgu_p.paragraph_format.first_line_indent = Pt(0)
            for run in bulgu_p.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
        doc.add_paragraph()
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# ── Endpoints ─────────────────────────────────────────────────────────────────

def _prepare_analysis_df(
    df: pd.DataFrame,
    variables: List[Variable],
    missing_codes: Optional[List[str]] = None,
) -> pd.DataFrame:
    df = normalize_continuous_columns(df, variables)
    df = normalize_categorical_columns(df, variables)
    df = apply_missing_codes(df, variables, missing_codes)
    df = sanitize_invalid_categorical_codes(df, variables, missing_codes)
    for v in variables:
        if re.match(r"^vki$", v.name, re.I) and v.name in df.columns:
            max_val = df[v.name].max()
            if pd.notna(max_val) and max_val > 1000:
                df[v.name] = pd.to_numeric(df[v.name], errors="coerce")
    return df


@app.post("/convert-spss-table")
@app.post("/import-spss-tables")
async def import_spss_tables_endpoint(req: SpssTableRequest):
    """Ham SPSS yapıştırmasını APA 7 tablolarına + Bulgular metnine dönüştürür."""
    method = "pandas"
    pandas_err = None
    try:
        results, meta = convert_spss_to_apa_results(req.content)
    except Exception as err:
        pandas_err = err
        if not ANTHROPIC_API_KEY:
            raise HTTPException(status_code=400, detail=str(pandas_err))
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        msg = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=4000,
            system=SPSS_CONVERT_SYSTEM,
            messages=[{"role": "user", "content": req.content}],
        )
        results = _markdown_tables_to_apa_results(msg.content[0].text.strip())
        if not results:
            raise HTTPException(status_code=400, detail=str(pandas_err))
        meta = {"source": "spss", "intro": "", "table_count": len(results), "ai_fallback": True}
        method = "ai"

    bulgular: Dict[str, str] = {}
    if req.auto_bulgu:
        for i, result in enumerate(results):
            try:
                text = _generate_bulgu_text(result)
                if text:
                    bulgular[str(i)] = text
            except Exception:
                pass

    return sanitize({
        "results": results,
        "bulgular": bulgular,
        "meta": meta,
        "method": method,
        "pandas_error": str(pandas_err) if pandas_err else None,
    })


@app.post("/plan")
async def generate_analysis_plan(req: PlanRequest):
    rows = [r.values for r in req.data]
    df = pd.DataFrame(rows)
    variables = normalize_variable_labels(req.variables)
    df = _prepare_analysis_df(df, variables, req.missing_codes)

    rule_based = generate_plan(df, variables)

    if req.use_ai and req.research_topic:
        ai_tests = await generate_plan_ai(variables, req.research_topic)
        if ai_tests:
            ai_ids = {t.get("id") for t in ai_tests}
            merged = list(ai_tests)
            for rb_test in rule_based:
                if rb_test["id"] not in ai_ids:
                    rb_copy = dict(rb_test)
                    rb_copy["recommended"] = False
                    merged.append(rb_copy)
            return sanitize({"tests": merged, "source": "ai"})

    return sanitize({"tests": rule_based, "source": "rules"})


@app.post("/analyze")
async def analyze(req: AnalysisRequest):
    rows = [r.values for r in req.data]
    df = pd.DataFrame(rows)
    variables = normalize_variable_labels(req.variables)
    variables = apply_scale_info_to_variables(variables, req.scale_info)
    df = _prepare_analysis_df(df, variables, req.missing_codes)
    missing_data = missing_data_report(df, variables)
    results, meta = run_analyze(
        df,
        variables,
        req.active_types,
        req.enabled_tests,
        req.scale_info,
        req.missing_codes,
    )
    return sanitize({"results": results, "missing_data": missing_data, "meta": meta})


@app.post("/analyze/cronbach")
async def analyze_cronbach(req: CronbachRequest):
    import traceback
    try:
        if len(req.columns) < 2:
            raise HTTPException(status_code=400, detail="En az 2 sütun gerekli")
        df = pd.DataFrame([r.values for r in req.data])
        missing = [c for c in req.columns if c not in df.columns]
        if missing:
            raise HTTPException(status_code=400, detail=f"Sütunlar bulunamadı: {missing}")
        for col in req.columns:
            df[col] = pd.to_numeric(df[col].astype(str).str.replace(",", ".", regex=False), errors="coerce")
        result = cronbach_analysis(df, req.columns)
        if result is None:
            raise HTTPException(status_code=400, detail="Cronbach alfa hesaplanamadı (yetersiz veri veya varyans)")
        return sanitize({"result": result})
    except HTTPException:
        raise
    except Exception as e:
        print(traceback.format_exc())
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/detect-scales")
async def detect_scales(req: DetectScalesRequest):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY ayarlanmamış")

    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

    ITEM_PATTERN = re.compile(r"^[a-zA-Z]+_\d+(_ters|_T)?$", re.IGNORECASE)
    item_cols = [c for c in req.columns if ITEM_PATTERN.match(c)]

    if len(item_cols) < 2:
        return {"scales": []}

    prefix_groups: Dict[str, list] = defaultdict(list)
    for col in item_cols:
        prefix = re.match(r"^([a-zA-Z]+)_", col)
        if prefix:
            prefix_groups[prefix.group(1)].append(col)

    valid_groups = {k: v for k, v in prefix_groups.items() if len(v) >= 3}

    if not valid_groups:
        return {"scales": []}

    group_info = "\n".join([
        f"- {prefix}: {', '.join(sorted(cols))}"
        for prefix, cols in valid_groups.items()
    ])

    msg = client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=800,
        system="""Sen akademik ölçek analizi uzmanısın. Verilen madde gruplarını ölçeklere dönüştür.

KURALLAR:
- Her prefix grubu ayrı bir ölçektir
- _ters veya _T ile biten madde varsa o maddenin ters versiyonunu kullan, orijinalini KULLANMA
- Orijinal madde ile _ters versiyonu AYNI ANDA listede olmamalı
- Her ölçeğe Türkçe anlamlı isim ver (oys → OYŞTÖ, neq → GYA veya NEQ, sbito → SBİTO)
- Tüm gruplar için ölçek oluştur, hiçbirini atlama

SADECE JSON döndür:
{
  "scales": [
    {"name": "OYŞTÖ", "items": ["oys_1", "oys_2", "oys_3", "oys_4_ters", "oys_5"]},
    {"name": "GYA", "items": ["neq_1_ters", "neq_2", "neq_3", "neq_4_ters"]}
  ]
}""",
        messages=[{
            "role": "user",
            "content": f"Şu madde gruplarını ölçeklere dönüştür:\n{group_info}",
        }],
    )

    text = msg.content[0].text.strip()
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        try:
            result = json.loads(match.group())
            if not result.get("scales"):
                raise ValueError("Boş sonuç")
            return result
        except Exception:
            pass

    scales = []
    scale_names = {"oys": "OYŞTÖ", "neq": "GYA", "sbito": "SBİTO"}
    for prefix, cols in valid_groups.items():
        ters_cols = {
            c.replace("_ters", "").replace("_T", "")
            for c in cols
            if "_ters" in c.lower() or c.endswith("_T")
        }
        final_items = [c for c in cols if c not in ters_cols]
        name = scale_names.get(prefix.lower(), prefix.upper())
        scales.append({"name": name, "items": final_items})

    return {"scales": scales}


@app.post("/analyze/cronbach-batch")
async def analyze_cronbach_batch(req: CronbachBatchRequest):
    df = pd.DataFrame([r.values for r in req.data])
    results = []
    tc = TableCounter()

    for scale in req.scales:
        name = scale.get("name", "Ölçek")
        items = scale.get("items", [])

        valid_items = []
        for col in items:
            if col in df.columns:
                df[col] = pd.to_numeric(
                    df[col].astype(str).str.replace(",", ".", regex=False),
                    errors="coerce",
                )
                valid_items.append(col)

        if len(valid_items) < 2:
            continue

        try:
            items_df = df[valid_items].dropna()
            k = len(valid_items)
            if len(items_df) < 3:
                continue

            item_vars = items_df.var(axis=0, ddof=1).sum()
            total_var = items_df.sum(axis=1).var(ddof=1)

            if total_var == 0:
                continue

            alpha = float((k / (k - 1)) * (1 - item_vars / total_var))

            if alpha >= 0.90:
                interp = "Yüksek"
            elif alpha >= 0.70:
                interp = "İyi"
            elif alpha >= 0.60:
                interp = "Kabul Edilebilir"
            else:
                interp = "Düşük"

            tno, title = tc.next(f"Ölçek Güvenilirlik Analizi — {name}")

            results.append({
                "type": "cronbach",
                "table_number": tno,
                "title": title,
                "headers": ["Ölçek", "Madde Sayısı", "Geçerli n", "Cronbach α", "Değerlendirme"],
                "rows": [[name, k, len(items_df), f"{alpha:.3f}", interp]],
                "note": "Not. α = Cronbach alfa iç tutarlılık katsayısı. Kabul edilebilir sınır: α ≥ .70.",
                "significant": None,
            })
        except Exception:
            continue

    return sanitize({"results": results})


@app.post("/analyze/paired")
async def analyze_paired(req: PairedRequest):
    df = pd.DataFrame([r.values for r in req.data])
    for col in (req.col1, req.col2):
        if col not in df.columns:
            raise HTTPException(status_code=400, detail=f"Sütun bulunamadı: {col}")
        df[col] = pd.to_numeric(df[col].astype(str).str.replace(",", ".", regex=False), errors="coerce")
    try:
        result = paired_ttest(df, req.col1, req.col2)
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))
    return {"result": result}


CLASSIFY_SYSTEM = """
Sen akademik araştırma veri analisti yardımcısısın.
Sütun adı, kullanıcı etiketi ve örnek değerlere bakarak
her sütunu sınıflandır.

━━━ TYPE KURALI ━━━

continuous:
- Sayısal, geniş aralıklı (ölçek puanı, yaş, boy, kilo, VKİ vb.)
- Örnek değerler birbirinden farklı, aralık geniş

categorical:
- Metin değerli (Erkek/Kadın, Evet/Hayır, bölüm adları vb.)
- Sayısal ama ≤8 farklı benzersiz değer (kod sistemi)

exclude:
- Sadece kimlik/sıra kolonları: id, no, sira, anket_no, serial
- Ölçek maddeleri: harf_rakam pattern (oys_1, bdi_3, sf36_12, sbito_6_ters)

━━━ ROLE KURALI ━━━

grouping (bağımsız/demografik değişken):
- Katılımcıları gruplara ayıran kategorik değişkenler
- Cinsiyet, bölüm, eğitim düzeyi, medeni durum, gelir,
  meslek, yaşadığı yer, çalışma durumu gibi
- Evet/Hayır tipi demografik sorular:
  sigara kullanımı, alkol kullanımı, kronik hastalık,
  ilaç kullanımı, spor yapma, diyet uygulama gibi
- KURAL: categorical + demografik anlam = grouping

outcome (bağımlı/sonuç değişkeni):
- Ölçek toplam/alt boyut puanları (continuous)
- Türetilmiş kategoriler: risk grubu, BKİ kategorisi,
  yaş grubu, puan kategorisi, sınıflandırma (categorical)
- Ham antropometrik ölçümler: yaş, boy, kilo, vücut ağırlığı,
  BKİ/VKİ ham değeri (continuous) — bunlar tanımlayıcı
  istatistikte raporlanır, gruplandırma için kullanılmaz
- KURAL: continuous ölçüm = outcome
- KURAL: _toplam/_total/_score/_puan/_skor ile biten = outcome
- KURAL: _grubu/_grup/_group/_binary/_kategori/_sinif/_level
  ile biten = outcome (categorical)

━━━ RECOMMENDED KURALI ━━━

true:
- Tüm outcome değişkenler
- Ana gruplandırma değişkenleri (cinsiyet, bölüm/departman/fakülte)
- Araştırma konusuyla doğrudan ilişkili demografikler

false:
- İkincil demografikler (medeni durum, gelir, konut durumu)
- Ham antropometrik ölçümler (boy, kilo) — VKİ zaten hesaplanmışsa
- Araştırma sorusuyla zayıf ilişkili değişkenler

━━━ ÖNEMLİ ━━━
- Hiçbir kolon adını hardcode olarak tanıma
- Karar her zaman: kolon adı + etiket + örnek değer kombinasyonuna göre
- Aynı kavramın hem ham (vki=20.5) hem kategorik (vki_kategori=Normal)
  versiyonu varsa: ham → outcome/continuous, kategorik → outcome/categorical
- Etiket varsa kolon adından daha güvenilir — etiketi öncelikle kullan

━━━ REASON YAZMA KURALLARI ━━━

Her değişken için reason yazarken şu ilişkileri tespit et:

1. BOY + KİLO + BKİ/VKİ İLİŞKİSİ:
   Kolon listesinde hem boy/height hem kilo/weight hem de
   bki/vki/bmi kolonu varsa:
   - boy ve kilo kolonları için reason:
     "BKİ/VKİ kolonu mevcut — boy ve kilo seçmeniz gerekmeyebilir"

2. HAM SÜREKLI + KATEGORİK VERSİYON:
   Aynı kavramın hem sürekli (yas, vki) hem kategorik versiyonu
   (yas_grubu, vki_kategori) varsa:
   - Sürekli versiyon için reason:
     "Kategorik versiyonu mevcut — ikisi birden seçilmesi gerekmez,
      biri yeterli"

3. ÖLÇEK TOPLAM + ALT BOYUT:
   Hem toplam puan (_toplam) hem alt boyut kolonları varsa:
   - Alt boyutlar için reason:
     "Toplam puan analizi yeterliyse alt boyutlar opsiyonel"

4. İKİNCİL DEMOGRAFİK:
   Araştırma konusuyla zayıf ilişkili demografikler için reason'a
   neden ikincil olduğunu yaz:
   "İkincil demografik — araştırma sorusuyla doğrudan ilişkisi zayıf"

5. TEMEL DEMOGRAFİK:
   Cinsiyet, bölüm/departman gibi temel gruplandırma değişkenleri için:
   "Temel demografik gruplandırma değişkeni"

6. ÖLÇEK PUANI:
   Toplam puan kolonları için ölçeğin ne ölçtüğünü belirt:
   Etiket varsa etiketten, yoksa kolon adından çıkar:
   "Online yemek siparişi tutumunu ölçen ana bağımlı değişken"

Reason maksimum 10 kelime, Türkçe, öz ve net olsun.

SADECE JSON döndür:
{
  "variables": {
    "kolon_adi": {
      "type": "categorical|continuous|exclude",
      "role": "grouping|outcome|exclude",
      "recommended": true|false,
      "reason": "kısa Türkçe açıklama"
    }
  }
}
"""


@app.post("/classify")
async def classify_columns(req: ClassifyRequest):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY ortam değişkeni ayarlanmamış")

    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    col_info = []
    for col in req.columns:
        samples = req.samples.get(col, [])
        label = (req.labels or {}).get(col, "")
        label_str = f", etiket='{label}'" if label and label != col else ""
        col_info.append(f"- {col}{label_str}: örnek={samples}")

    topic_str = f"\nAraştırma konusu: {req.research_topic}" if req.research_topic else ""
    user_msg = f"Sütunları sınıflandır:{topic_str}\n\n" + "\n".join(col_info)

    msg = client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=1500,
        system=CLASSIFY_SYSTEM,
        messages=[{"role": "user", "content": user_msg}],
    )

    parsed = _parse_llm_json(msg.content[0].text.strip())
    variables = parsed.get("variables", {})

    categorical: List[str] = []
    continuous: List[str] = []
    exclude: List[str] = []
    recommendations: Dict[str, dict] = {}

    for col, info in variables.items():
        if col not in req.columns:
            continue
        t = info.get("type", "exclude")
        role = info.get("role", "exclude")
        rec = info.get("recommended", False)
        reason = info.get("reason", "")

        if t == "exclude" or role == "exclude":
            exclude.append(col)
        elif t == "categorical":
            categorical.append(col)
        else:
            continuous.append(col)

        recommendations[col] = {
            "status": "recommended" if rec else "optional",
            "role": role,
            "reason": reason,
        }

    for col in req.columns:
        if col not in variables:
            exclude.append(col)

    return {
        "categorical": categorical,
        "continuous": continuous,
        "exclude": exclude,
        "recommendations": recommendations,
    }


@app.post("/ai/bulgu")
async def ai_bulgu(req: BulguRequest):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY ortam değişkeni ayarlanmamış")
    text = _generate_bulgu_text(
        req.result,
        req.research_topic,
        req.label_map,
        req.approved_cutoffs,
        req.scale_info,
        req.pdf_context,
    )
    return {"bulgu": text}


@app.post("/export/word")
async def export_word(req: WordExportRequest):
    try:
        doc_bytes = build_word_document(
            req.results, req.bulgular, req.intro or "", req.label_map
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Word dosyası oluşturulamadı: {str(e)}")
    return StreamingResponse(
        io.BytesIO(doc_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=statai_bulgular.docx"},
    )


# ── Ölçek → Kolon Eşleştirme ───────────────────────────────────────────────────

_TR_ASCII = str.maketrans({
    "ş": "s", "Ş": "s",
    "ı": "i", "İ": "i",
    "ğ": "g", "Ğ": "g",
    "ü": "u", "Ü": "u",
    "ö": "o", "Ö": "o",
    "ç": "c", "Ç": "c",
})
_TOTAL_MARKERS = ("toplam", "total", "puan", "skor", "score")
_ITEM_COL_RE = re.compile(r"(?:^|_)\d+(?:_ters|_t)?$", re.I)


def normalize_for_match(text: str) -> str:
    """Eşleştirme için string normalize et."""
    if not text:
        return ""
    s = str(text).translate(_TR_ASCII)
    s = s.lower()
    s = re.sub(r"[\s\-._]+", " ", s)
    words = []
    for word in s.split():
        word = re.sub(r"\d+$", "", word)
        if word:
            words.append(word)
    s = " ".join(words)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _keyword_prefixes(keyword: str, min_len: int = 3) -> List[str]:
    if " " in keyword or len(keyword) < min_len:
        return []
    return [keyword[:i] for i in range(min_len, len(keyword) + 1)]


def extract_scale_keywords(scale_name: str) -> List[str]:
    """Kullanıcının yazdığı ölçek adından eşleştirme keyword'leri üret."""
    keywords: List[str] = []
    raw = (scale_name or "").strip()
    if not raw:
        return []

    raw_no_paren = re.sub(r"\([^)]*\)", " ", raw)
    normalized_full = normalize_for_match(raw_no_paren)
    if normalized_full and " " not in normalized_full:
        keywords.append(normalized_full)

    paren_match = re.search(r"\(([^)]+)\)", raw)
    if paren_match:
        inner = normalize_for_match(paren_match.group(1))
        if inner:
            keywords.append(inner)

    words_raw = re.sub(r"\([^)]*\)", "", raw)
    words_raw = re.sub(r"[^\w\sğüşıöçĞÜŞİÖÇ]", " ", words_raw, flags=re.I)
    word_tokens = [w for w in words_raw.split() if w]
    if word_tokens:
        acronym = "".join(normalize_for_match(w)[:1] for w in word_tokens if w)
        if acronym:
            keywords.append(acronym)

    for word in word_tokens:
        norm_word = normalize_for_match(word)
        if len(norm_word) >= 3:
            keywords.append(norm_word)

    if word_tokens:
        first = normalize_for_match(word_tokens[0])
        if first:
            keywords.append(first)

    expanded: List[str] = []
    for kw in keywords:
        expanded.append(kw)
        expanded.extend(_keyword_prefixes(kw))

    seen = set()
    result: List[str] = []
    for kw in expanded:
        kw = kw.strip()
        if kw and kw not in seen:
            seen.add(kw)
            result.append(kw)
    return result


def _column_stem(col: str) -> str:
    base = re.split(r"[_\-\s]", col, maxsplit=1)[0]
    return normalize_for_match(base)


def _keyword_matches_column(kw: str, col: str) -> bool:
    norm_col = normalize_for_match(col)
    stem = _column_stem(col)
    if not kw or not stem:
        return False
    if stem == kw:
        return True
    if kw.startswith(stem) and len(kw) - len(stem) <= 2:
        return True
    if stem.startswith(kw) and len(stem) - len(kw) <= 1:
        return True
    if len(kw) >= 4 and kw in norm_col:
        return True
    return False


def _column_matches_keywords(col: str, keywords: List[str]) -> bool:
    return any(_keyword_matches_column(kw, col) for kw in keywords)


def _best_keyword_len(col: str, keywords: List[str]) -> int:
    matched = [len(kw) for kw in keywords if _keyword_matches_column(kw, col)]
    return max(matched) if matched else 0


def _classify_matched_column(col: str) -> str:
    norm = normalize_for_match(col)
    if any(marker in norm for marker in _TOTAL_MARKERS):
        return "total"
    if _ITEM_COL_RE.search(col):
        return "item"
    return "subscale"


def _match_confidence(
    total_columns: List[str],
    item_columns: List[str],
    subscale_columns: List[str],
) -> str:
    if total_columns:
        return "high"
    if item_columns:
        return "medium"
    if subscale_columns:
        return "low"
    return "low"


def match_scale_to_columns(scale_name: str, all_columns: List[str]) -> dict:
    """Tek ölçeği kolon listesiyle eşleştir."""
    keywords = extract_scale_keywords(scale_name)
    matched_columns: List[str] = []

    for col in all_columns:
        if _column_matches_keywords(col, keywords):
            matched_columns.append(col)

    item_columns = [c for c in matched_columns if _classify_matched_column(c) == "item"]
    total_columns = [c for c in matched_columns if _classify_matched_column(c) == "total"]
    subscale_columns = [
        c for c in matched_columns
        if c not in item_columns and c not in total_columns
    ]

    return {
        "scale_name": scale_name,
        "keywords_used": keywords,
        "matched_columns": matched_columns,
        "item_columns": item_columns,
        "total_columns": total_columns,
        "subscale_columns": subscale_columns,
        "confidence": _match_confidence(total_columns, item_columns, subscale_columns),
    }


def match_all_scales(scale_names: List[str], all_columns: List[str]) -> dict:
    """Tüm ölçekleri eşleştir; her kolon en fazla bir ölçeğe gitsin."""
    raw_matches = [
        match_scale_to_columns(name.strip(), all_columns)
        for name in scale_names
        if name and name.strip()
    ]

    col_owner: Dict[str, Tuple[int, int]] = {}
    for scale_idx, match in enumerate(raw_matches):
        for col in match["matched_columns"]:
            kw_len = _best_keyword_len(col, match["keywords_used"])
            if col not in col_owner or kw_len > col_owner[col][1]:
                col_owner[col] = (scale_idx, kw_len)

    final_matches: List[dict] = []
    for scale_idx, match in enumerate(raw_matches):
        owned = [
            col for col, (owner_idx, _) in col_owner.items()
            if owner_idx == scale_idx
        ]
        item_columns = [c for c in owned if _classify_matched_column(c) == "item"]
        total_columns = [c for c in owned if _classify_matched_column(c) == "total"]
        subscale_columns = [
            c for c in owned
            if c not in item_columns and c not in total_columns
        ]
        final_matches.append({
            "scale_name": match["scale_name"],
            "keywords_used": match["keywords_used"],
            "matched_columns": owned,
            "item_columns": item_columns,
            "total_columns": total_columns,
            "subscale_columns": subscale_columns,
            "confidence": _match_confidence(total_columns, item_columns, subscale_columns),
        })

    assigned = set(col_owner.keys())
    unmatched_columns = [c for c in all_columns if c not in assigned]

    return {
        "matches": final_matches,
        "unmatched_columns": unmatched_columns,
    }


@app.post("/match-scales")
def match_scales_endpoint(req: ScaleMatchRequest):
    """Saf string matching, AI yok, dış kaynak yok."""
    return match_all_scales(req.scale_names, req.column_names)


# ── PDF Bağlam Çıkarma ────────────────────────────────────────────────────────

_CONTEXT_MAX_CHARS = 1500


def extract_full_text(file_base64: str, file_type: str) -> Tuple[str, int]:
    """PDF veya DOCX'ten tüm metni çek."""
    try:
        raw = base64.b64decode(file_base64)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Base64 decode hatası: {e}")

    if file_type == "pdf":
        try:
            reader = PdfReader(io.BytesIO(raw))
            pages = [(p.extract_text() or "").strip() for p in reader.pages]
            text = "\n\n".join(p for p in pages if p)
            return text, len(reader.pages)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"PDF okunamadı: {e}")

    if file_type == "docx":
        try:
            doc = Document(io.BytesIO(raw))
            paragraphs = [(p.text or "").strip() for p in doc.paragraphs if (p.text or "").strip()]
            text = "\n".join(paragraphs)
            page_count = max(1, (len(paragraphs) + 4) // 5)
            return text, page_count
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"DOCX okunamadı: {e}")

    raise HTTPException(status_code=400, detail="file_type 'pdf' veya 'docx' olmalı")


def _split_paragraphs(full_text: str) -> List[str]:
    if "\n\n" in full_text:
        return [p.strip() for p in full_text.split("\n\n") if p.strip()]
    return [p.strip() for p in full_text.split("\n") if p.strip()]


def _label_search_keywords(label: str) -> List[str]:
    keywords = extract_scale_keywords(label)
    norm_label = normalize_for_match(label)
    if norm_label and " " not in norm_label:
        keywords.insert(0, norm_label)
    for word in re.sub(r"\([^)]*\)", "", label).split():
        norm_word = normalize_for_match(word)
        if len(norm_word) >= 3:
            keywords.append(norm_word)
    seen = set()
    result: List[str] = []
    for kw in keywords:
        if kw and kw not in seen:
            seen.add(kw)
            result.append(kw)
    return result


def find_relevant_paragraphs(full_text: str, label: str, window: int = 3) -> str:
    """Etiket kelimelerine göre ilgili paragrafları bul."""
    if not full_text or not label:
        return ""

    paragraphs = _split_paragraphs(full_text)
    if not paragraphs:
        return ""

    keywords = _label_search_keywords(label)
    if not keywords:
        return ""

    matched_indices: List[int] = []
    for i, para in enumerate(paragraphs):
        norm_para = normalize_for_match(para)
        if any(kw in norm_para for kw in keywords if len(kw) >= 3):
            matched_indices.append(i)

    if not matched_indices:
        return ""

    selected: List[int] = []
    seen_idx = set()
    for idx in matched_indices:
        for j in range(max(0, idx - window), min(len(paragraphs), idx + window + 1)):
            if j not in seen_idx:
                seen_idx.add(j)
                selected.append(j)

    combined = "\n\n".join(paragraphs[i] for i in selected)
    if len(combined) > _CONTEXT_MAX_CHARS:
        combined = combined[:_CONTEXT_MAX_CHARS]
    return combined


def build_label_context_map(
    full_text: str,
    column_labels: List[ColumnLabel],
) -> Dict[str, str]:
    """Her kolon etiketi için ilgili paragrafları çıkar."""
    result: Dict[str, str] = {}
    for cl in column_labels:
        result[cl.column] = find_relevant_paragraphs(full_text, cl.label)
    return result


@app.post("/extract-context")
async def extract_context(req: ExtractContextRequest):
    """PDF/DOCX'ten her etiket için ilgili paragrafları çek. AI yok."""
    file_type = (req.file_type or "").lower().strip()
    if file_type not in ("pdf", "docx"):
        raise HTTPException(status_code=400, detail="file_type 'pdf' veya 'docx' olmalı")
    if not req.column_labels:
        return {"context_map": {}, "page_count": 0, "matched_count": 0}

    full_text, page_count = extract_full_text(req.file_base64, file_type)
    if not full_text.strip():
        raise HTTPException(status_code=400, detail="Dosyadan metin çıkarılamadı")

    context_map = build_label_context_map(full_text, req.column_labels)
    return {
        "context_map": context_map,
        "page_count": page_count,
        "matched_count": sum(1 for v in context_map.values() if v),
    }


# ── Etik Kurul Raporu ─────────────────────────────────────────────────────────

ETHICS_KEYWORDS = [
    # Türkçe
    "ölçek", "olcek", "anket", "cronbach", "alfa", "güvenilirlik",
    "geçerlilik", "gecerlilik", "madde", "alt boyut", "likert", "puan",
    "veri toplama", "ölçüm", "olcum", "kesim noktası",
    # İngilizce
    "alpha", "scale", "questionnaire", "inventory", "reliability",
    "validity", "subscale", "cutoff", "instrument",
    # Yaygın ölçek kısaltmaları (büyük/küçük harf fark etmez)
    "neq", "oysto", "sbito", "gya", "bdi", "bai", "phq", "gad",
]
ETHICS_MAX_FILTERED_CHARS = 24000
ETHICS_FALLBACK_CHARS = 16000

ETHICS_SYSTEM_PROMPT = (
    "Sen bir akademik araştırma asistanısın. Verilen etik kurul raporu metninden "
    "ölçek bilgilerini çıkar. SADECE JSON döndür, başka hiçbir şey yazma."
)


def _section_matches_keywords(text: str) -> bool:
    lower = text.lower()
    normalized = lower.translate(_TR_ASCII)
    normalized_keywords = [kw.translate(_TR_ASCII) for kw in ETHICS_KEYWORDS]
    return (
        any(kw in lower for kw in ETHICS_KEYWORDS) or
        any(kw in normalized for kw in normalized_keywords)
    )


def _extract_pdf_sections(file_bytes: bytes) -> Tuple[List[str], int]:
    reader = PdfReader(io.BytesIO(file_bytes))
    sections = []
    for page in reader.pages:
        try:
            text = (page.extract_text() or "").strip()
        except Exception:
            text = ""
        if len(text) > 50:
            sections.append(text)
    return sections, len(reader.pages)


def _extract_docx_sections(file_bytes: bytes) -> Tuple[List[str], int]:
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [(p.text or "").strip() for p in doc.paragraphs if (p.text or "").strip()]
    if not paragraphs:
        return [], 1
    chunk_size = 5
    sections = []
    for i in range(0, len(paragraphs), chunk_size):
        sections.append("\n".join(paragraphs[i:i + chunk_size]))
    return sections, max(1, len(sections))


def _filter_ethics_sections(sections: List[str]) -> Tuple[str, int, int]:
    total = len(sections)
    matched = [s for s in sections if _section_matches_keywords(s)]
    filtered_count = len(matched)

    if matched:
        combined = "\n\n".join(matched)
        if len(combined) > ETHICS_MAX_FILTERED_CHARS:
            combined = combined[:ETHICS_MAX_FILTERED_CHARS]
        return combined, filtered_count, total

    full_text = "\n\n".join(sections)
    if len(full_text) <= ETHICS_FALLBACK_CHARS:
        return full_text, 0, total

    chunk = ETHICS_FALLBACK_CHARS // 3
    start = full_text[:chunk]
    mid_start = len(full_text) // 2 - chunk // 2
    middle = full_text[mid_start:mid_start + chunk]
    end = full_text[-chunk:]
    combined = f"{start}\n\n[...]\n\n{middle}\n\n[...]\n\n{end}"
    return combined, 0, total


def _build_ethics_user_prompt(filtered_text: str) -> str:
    return f"""Aşağıdaki metinden araştırmada kullanılan TÜM ölçekleri çıkar.

Her ölçek için şunları bul (bulamazsan null bırak):
- name: Ölçeğin tam adı (Türkçe ve/veya İngilizce)
- short_name: Kısaltma (örn: BDI, OYŞTÖ, NEQ)
- item_count: Madde/soru sayısı (integer)
- subscales: Alt boyutlar listesi (string array)
- likert_range: Likert aralığı (örn: "1-5", "0-4")
- min_score / max_score: Teorik min-max puan
- cutoffs: Kesim noktaları (örn: [{{"label": "Normal", "range": "0-9", "score": 9}}])
- cronbach: Raporlanan Cronbach alpha değeri (float)
- developer: Ölçeği geliştiren kişi/kurum
- adaptation: Türkçe uyarlamacısı
- citation: Kaynak/referans metni

Metin:
{filtered_text}

Yanıt formatı:
{{
  "scales": [...],
  "research_topic": "araştırmanın konusu (1-2 cümle)",
  "filtered_page_count": kaç sayfa/bölüm analiz edildi (integer),
  "total_page_count": toplam sayfa sayısı (integer)
}}"""


@app.post("/import-ethics-report")
async def import_ethics_report(
    file: UploadFile = File(...),
    research_topic: Optional[str] = Form(None),
):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY ayarlanmamış")

    filename = (file.filename or "").lower()
    if not (filename.endswith(".pdf") or filename.endswith(".docx")):
        raise HTTPException(status_code=400, detail="Sadece PDF veya DOCX dosyaları desteklenir")

    try:
        file_bytes = await file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Dosya okunamadı: {e}")

    if not file_bytes:
        raise HTTPException(status_code=400, detail="Dosya boş")

    try:
        if filename.endswith(".pdf"):
            sections, total_pages = _extract_pdf_sections(file_bytes)
        else:
            sections, total_pages = _extract_docx_sections(file_bytes)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Metin çıkarılamadı: {e}")

    if not sections:
        raise HTTPException(status_code=400, detail="Dosyadan metin çıkarılamadı")

    filtered_text, filtered_pages, total_pages = _filter_ethics_sections(sections)
    chars_sent = len(filtered_text)

    user_prompt = _build_ethics_user_prompt(filtered_text)
    if research_topic:
        user_prompt = f"Araştırma konusu (kullanıcı): {research_topic}\n\n{user_prompt}"

    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        msg = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=4000,
            system=ETHICS_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": user_prompt}],
        )
        parsed = _parse_llm_json(msg.content[0].text.strip())
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI analizi başarısız: {e}")

    scales = parsed.get("scales") or []
    topic = parsed.get("research_topic") or research_topic or ""
    ai_filtered = parsed.get("filtered_page_count")
    ai_total = parsed.get("total_page_count")
    effective_filtered = filtered_pages if filtered_pages > 0 else total_pages

    return sanitize({
        "scales": scales,
        "research_topic": topic,
        "token_info": {
            "filtered_pages": ai_filtered if ai_filtered is not None else effective_filtered,
            "total_pages": ai_total if ai_total is not None else total_pages,
            "chars_sent": chars_sent,
        },
    })


def _cell_str(val) -> str:
    if val is None or (isinstance(val, float) and math.isnan(val)):
        return ""
    return str(val).strip()


def _parse_excel_rows(rows: List[list]) -> Tuple[List[dict], List[str], Dict[str, str]]:
    if not rows:
        raise HTTPException(status_code=400, detail="Boş dosya")

    header_row = [_cell_str(c) for c in rows[0]]
    labels: Dict[str, str] = {}
    data_start = 1

    if len(rows) > 1:
        second_row = [_cell_str(c) for c in rows[1]]
        is_label_row = (
            all(
                not v.replace(".", "").replace("-", "").isdigit()
                for v in second_row if v
            )
            and second_row != header_row
            and any(
                v and h and v.lower() != h.lower()
                for v, h in zip(second_row, header_row)
            )
        )
        if is_label_row:
            labels = {h: v for h, v in zip(header_row, second_row) if v and h}
            data_start = 2

    records = []
    for row in rows[data_start:]:
        if all(_cell_str(c) == "" for c in row):
            continue
        records.append({
            header_row[i]: row[i] if i < len(row) else ""
            for i in range(len(header_row))
            if header_row[i]
        })

    return records, header_row, labels


@app.post("/read-file")
async def read_file(file: UploadFile = File(...)):
    filename = (file.filename or "").lower()
    file_bytes = await file.read()

    if not file_bytes:
        raise HTTPException(status_code=400, detail="Dosya boş")

    labels: Dict[str, str] = {}
    value_labels: Dict[str, Any] = {}
    records: List[dict] = []
    columns: List[str] = []
    source = "excel"

    if filename.endswith(".sav"):
        try:
            import os
            import tempfile
            import pyreadstat

            with tempfile.NamedTemporaryFile(suffix=".sav", delete=False) as f:
                f.write(file_bytes)
                tmp_path = f.name
            try:
                df, meta = pyreadstat.read_sav(tmp_path)
            finally:
                os.unlink(tmp_path)

            column_names = meta.column_names or list(df.columns)
            column_label_list = meta.column_labels or []
            for i, col in enumerate(column_names):
                if i < len(column_label_list):
                    lbl = column_label_list[i]
                    if lbl and str(lbl).strip():
                        labels[str(col)] = str(lbl).strip()

            value_labels = meta.variable_value_labels or {}
            records = df.replace({np.nan: None}).to_dict(orient="records")
            columns = [str(c) for c in df.columns]
            source = "spss"
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"SAV okunamadı: {e}")

    elif filename.endswith(".xlsx"):
        try:
            from openpyxl import load_workbook

            wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
            ws = wb.active
            rows = list(ws.iter_rows(values_only=True))
            records, columns, labels = _parse_excel_rows(rows)
            wb.close()
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Excel okunamadı: {e}")

    elif filename.endswith(".xls"):
        try:
            df_raw = pd.read_excel(io.BytesIO(file_bytes), header=None)
            rows = df_raw.values.tolist()
            records, columns, labels = _parse_excel_rows(rows)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Excel okunamadı: {e}")

    elif filename.endswith(".csv"):
        raise HTTPException(status_code=400, detail="CSV için frontend okuma kullanın")

    else:
        raise HTTPException(status_code=400, detail="Desteklenmeyen format")

    if not records:
        raise HTTPException(status_code=400, detail="Dosya boş görünüyor")

    return sanitize({
        "data": records,
        "labels": labels,
        "value_labels": value_labels,
        "columns": columns,
        "row_count": len(records),
        "source": source,
        "labels_found": len(labels),
    })


@app.get("/")
def root():
    return {"status": "ok", "app": "StatAI"}


def sanitize(obj):
    if isinstance(obj, (float, np.floating)):
        if math.isnan(obj) or math.isinf(obj):
            return None
    if isinstance(obj, dict):
        return {k: sanitize(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [sanitize(i) for i in obj]
    return obj
