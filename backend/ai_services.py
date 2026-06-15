"""LLM servisleri."""
import base64
import io
import json
import re
from collections import defaultdict
from typing import List, Optional, Dict, Any, Tuple
import anthropic
from docx import Document
from fastapi import HTTPException
from pypdf import PdfReader
from config import ANTHROPIC_API_KEY, ANTHROPIC_MODEL, BULGU_MODEL
from data_cleaning import apply_scale_item_resolution
from data_profile import profile_from_samples, profile_json
from llm_router import (
    _parse_json_object,
    claude_decide,
    format_enrichment_block,
    gemini_enrich_profile,
    gemini_json_task,
    has_claude,
    has_gemini_enrich,
    merge_meta,
)
from schemas import Variable, ColumnLabel
from formatting import substitute_variable_codes, apply_academic_text_rules
from constants import (
    _TR_ASCII, _TOTAL_MARKERS, _ITEM_COL_RE, _CONTEXT_MAX_CHARS,
    ETHICS_KEYWORDS, ETHICS_MAX_FILTERED_CHARS, ETHICS_FALLBACK_CHARS,
    ETHICS_SYSTEM_PROMPT, BULGU_SUMMARY_SYSTEM,
)
from yorum_motoru import build_llm_rules_block
from yorum_motoru import (
    build_summaries_from_results,
    generate_template_summary,
    validate_summary_text,
)
from spss_import import convert_spss_to_apa_results, _markdown_tables_to_apa_results
from stat_tests import (
    generate_plan,
    _normalize_ai_plan_tests,
    _normalize_ai_plan_ids,
)
from schemas import ClassifyRequest, DetectScalesRequest, SpssTableRequest
from utils import sanitize

SPSS_CONVERT_SYSTEM = """Sen SPSS çıktı dönüştürme uzmanısın. Ham metin, HTML veya kopyala-yapıştır formatındaki SPSS analiz çıktılarını (çapraz tablolar, ANOVA, Ki-Kare, t-Testi vb.) sıfır veri kaybıyla kusursuz Markdown tablosuna dönüştür.

KATI KURALLAR:
1. BİRLEŞTİRİLMİŞ HÜCRELER: χ², p, F değerlerini TÜM satırlara forward fill ile çoğalt.
2. HAYALET SÜTUNLAR: Boş sütunları ve sütun kaymalarını düzelt; header ile satır sütun sayısı eşit olsun.
3. KAYIP VERİ: İsimsiz ama frekanslı satırları "Kayıp Veri" olarak adlandır.
4. POST-HOC: (I) grubu birleştirilmişse altındaki (J) satırlarına (I) adını yaz.

SADECE temizlenmiş Markdown tablolarını döndür. Yorum yapma."""

def _parse_llm_json(text: str) -> dict:
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except json.JSONDecodeError:
            pass
    return {}

RESULTS_SYSTEM = """Sen biyoistatistiksel tabloları APA 7 tez formatında Bulgular düz yazısına dönüştüren bir asistansın. Yalnızca tablodaki verilere dayanarak 2-4 cümle Türkçe bulgu yaz; Tartışma yazma; madde işareti/başlık kullanma.

GENEL: Geçmiş zaman (-miştir/-mıştır). Ham kod (OYS_TOPLAM vb.) yasak; sözlükteki akademik adları kullan. kilo→Vücut ağırlığı, boy→Boy uzunluğu, VKI→BKİ. Tablodaki sayıları aynen aktar; yuvarlama yapma. Semboller: x̄, SS, n, %, p, t, F, χ², η², r, U, H, df.

1) TANIMLAYICI (descriptive): "[Ölçek] ortalama puanı x̄ ± SS (n = ..., medyan = ...) olarak hesaplanmıştır." Teorik aralık varsa: "ölçeğin teorik puan aralığı ... – ... arasındadır."

2) NORMALLİK (normality): p < .05 → "dağılımın normal dağılımdan anlamlı sapma gösterdiği saptanmıştır (D = ..., p < .05)." p ≥ .05 → "değişkenin normal dağılım gösterdiği belirlenmiştir (D = ..., p = ...)." Çarpıklık/basıklık verilirken: "Çarpıklık/Std. Hata ve Basıklık/Std. Hata değerlerinin ±2.0 sınırları içinde kalması sebebiyle parametrik testlerin uygulanmasına karar verilmiştir."

3) FREKANS/Kİ-KARE (frequency, chi_square, fisher_exact): Frekansta en yüksek n'li kategori önce. Ki-kare/Fisher p ≥ .05 → anlamlı ilişki/fark yok. p < .05 → anlamlı ilişki var; en baskın yüzdeyi vurgula. Fisher testi kullanıldıysa OR ve p değerini yaz. Beklenen frekans uyarısı varsa nota yansıt. Kayıp Veri kategorisini yorumlama.

4) ANOVA + TUKEY (anova, tukey): "[Bağımlı] toplam puanlarının [Bağımsız]'e göre tek yönlü ANOVA sonucunda gruplar arasında anlamlı fark saptanmıştır/saptanmamıştır [F = ..., p = ...; η² = ...]." Tukey post-hoc anlamlıysa x̄ BÜYÜK olan grup yüksek yazılır; küçük ortalama asla yüksek raporlanamaz.

5) t-TEST / NON-PARAMETRİK (ttest, mann_whitney, kruskal_wallis, dunn, paired_ttest, paired_wilcoxon): t-testi: "[Değişken] puanlarının ... karşılaştırılması amacıyla bağımsız örneklem t-testi sonucunda ... [t(df) = ..., p = ...; Cohen's d = ...]." Mann-Whitney: medyan, U, z, p ve r = |z|/√n (küçük r < .30, orta .30–.50, büyük > .50). Kruskal-Wallis anlamlıysa Dunn post-hoc: anlamlı çiftlerde medyanı YÜKSEK olan grubu belirt. paired_wilcoxon: medyanlar, z, p, r.

6) KORELASYON (correlation, correlation_matrix): r pozitif/negatif yön + anlamlılık (r = ..., p ...). Büyüklük: zayıf (r < .30), orta (.30–.70), güçlü (r > .70).

7) ÇOKLU REGRESYON (multiple_regression): Anlamlı yordayıcıları B, β ve p ile raporla; model R², düzeltilmiş R², F ve model p değerini yaz. VIF > 10 uyarısı varsa nota yansıt.

8) KESİM NOKTASI: Verilmişse ortalama/frekans yorumuna doğal biçimde göm."""

def _bulgu_context_blob(result: dict) -> str:
    chunks = [
        str(result.get("title", "")),
        str(result.get("type", "")),
        str(result.get("var1", "")),
        str(result.get("var2", "")),
        str(result.get("variable", "")),
    ]
    chunks.extend(str(h) for h in result.get("headers", []))
    for row in result.get("rows", []):
        chunks.extend(str(c) for c in row)
    chunks.append(str(result.get("note", "")))
    return " ".join(chunks).lower()

def _compact_bulgu_result(result: dict) -> dict:
    compact = {
        "type": result.get("type"),
        "table_number": result.get("table_number"),
        "title": result.get("title"),
        "headers": result.get("headers"),
        "rows": result.get("rows"),
        "note": result.get("note"),
    }
    for key in (
        "chi2", "p", "dof", "F", "eta2", "t", "df", "r", "d", "U", "H",
        "significant", "var1", "var2", "variable", "n",
    ):
        if key in result and result[key] is not None:
            compact[key] = result[key]
    return compact

def _filter_label_map(label_map: Optional[Dict[str, str]], blob: str) -> Optional[Dict[str, str]]:
    if not label_map:
        return None
    matched = {
        code: label
        for code, label in label_map.items()
        if code.lower() in blob or (label and label.lower() in blob)
    }
    return matched or None

def _filter_scale_info(scale_info: Optional[dict], blob: str) -> Optional[dict]:
    if not scale_info:
        return None
    matched = {}
    for key, info in scale_info.items():
        if not isinstance(info, dict):
            continue
        full = str(info.get("full_name") or "")
        if (
            key.lower() in blob
            or full.lower() in blob
            or any(tok in blob for tok in key.lower().split() if len(tok) > 2)
        ):
            matched[key] = {
                k: info[k]
                for k in ("full_name", "item_count", "min_score", "max_score", "cutoff", "likert")
                if k in info
            }
    return matched or None

def _filter_cutoffs(cutoffs: Optional[List[dict]], blob: str) -> Optional[List[dict]]:
    if not cutoffs:
        return None
    matched = [
        c for c in cutoffs
        if str(c.get("code", "")).lower() in blob
        or str(c.get("label", "")).lower() in blob
        or str(c.get("scale_name", "")).lower() in blob
    ]
    return matched or None

def build_bulgu_user_message(
    result: dict,
    research_topic: Optional[str] = None,
    label_map: Optional[Dict[str, str]] = None,
    approved_cutoffs: Optional[List[dict]] = None,
    scale_info: Optional[dict] = None,
    pdf_context: Optional[str] = None,
) -> str:
    blob = _bulgu_context_blob(result)
    parts = []
    if research_topic:
        parts.append(f"Konu: {research_topic}")
    labels = _filter_label_map(label_map, blob)
    if labels:
        parts.append("Etiketler: " + json.dumps(labels, ensure_ascii=False))
    cutoffs = _filter_cutoffs(approved_cutoffs, blob)
    if cutoffs:
        parts.append("Kesim: " + json.dumps(cutoffs, ensure_ascii=False))
    scales = _filter_scale_info(scale_info, blob)
    if scales:
        parts.append("Ölçek: " + json.dumps(scales, ensure_ascii=False))
    parts.append("Tablo: " + json.dumps(_compact_bulgu_result(result), ensure_ascii=False))
    rules = build_llm_rules_block()
    if rules:
        parts.append(rules)
    message = "\n".join(parts)
    if pdf_context:
        message += f"""

--- ÖLÇEK KAYNAK BİLGİSİ (etik kurul / araştırma raporundan) ---
{pdf_context}
--- Bu bilgiyi bulgularda ölçeği tanıtırken ve yorumlarken kullan.
Birebir kopyalama, kendi cümlelerinle akademik dille yeniden yaz. ---"""
    return message

def _empty_llm_meta() -> dict:
    return {"llm_calls": 0, "approx_input_tokens": 0, "approx_output_tokens": 0}


def _llm_meta_from_usage(usage) -> dict:
    if not usage:
        return {"llm_calls": 1, "approx_input_tokens": 0, "approx_output_tokens": 0}
    return {
        "llm_calls": 1,
        "approx_input_tokens": int(usage.input_tokens or 0),
        "approx_output_tokens": int(usage.output_tokens or 0),
    }


def _generate_bulgu_text_llm(
    result: dict,
    research_topic: Optional[str] = None,
    label_map: Optional[Dict[str, str]] = None,
    approved_cutoffs: Optional[List[dict]] = None,
    scale_info: Optional[dict] = None,
    pdf_context: Optional[str] = None,
) -> Tuple[str, dict]:
    if not ANTHROPIC_API_KEY:
        return "", _empty_llm_meta()
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    msg = client.messages.create(
        model=BULGU_MODEL,
        max_tokens=600,
        system=RESULTS_SYSTEM,
        messages=[{
            "role": "user",
            "content": build_bulgu_user_message(
                result, research_topic, label_map, approved_cutoffs, scale_info, pdf_context
            ),
        }],
    )
    return msg.content[0].text.strip(), _llm_meta_from_usage(msg.usage)


def generate_bulgu(
    result: dict,
    research_topic: Optional[str] = None,
    label_map: Optional[Dict[str, str]] = None,
    approved_cutoffs: Optional[List[dict]] = None,
    scale_info: Optional[dict] = None,
    pdf_context: Optional[str] = None,
    force_llm: bool = False,
    all_results: Optional[List[dict]] = None,
) -> Tuple[str, str, dict]:
    """(bulgu_metni, kaynak, meta) döndürür. kaynak: template | llm."""
    if not force_llm and has_bulgu_template(result):
        template_text = generate_bulgu_from_template(result, label_map, all_results)
        if template_text:
            return template_text, "template", _empty_llm_meta()
    text, meta = _generate_bulgu_text_llm(
        result, research_topic, label_map, approved_cutoffs, scale_info, pdf_context,
    )
    if text and label_map:
        text = substitute_variable_codes(text, label_map)
        text = apply_academic_text_rules(text)
    return text, "llm", meta


def generate_bulgu_summary(
    summaries: List[dict],
    research_topic: Optional[str] = None,
    hypotheses: Optional[List[dict]] = None,
) -> Tuple[str, dict]:
    template_text = generate_template_summary(summaries, hypotheses)
    if not summaries:
        return "", _empty_llm_meta()
    if not ANTHROPIC_API_KEY:
        return template_text, {**_empty_llm_meta(), "source": "template"}

    payload = {
        "konu": research_topic or "",
        "ozetler": summaries,
    }
    if hypotheses:
        payload["hipotezler"] = [
            {
                "id": h.get("id"),
                "label": h.get("label"),
            }
            for h in hypotheses
            if h.get("id")
        ]
    user_msg = json.dumps(payload, ensure_ascii=False)
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    msg = client.messages.create(
        model=BULGU_MODEL,
        max_tokens=800,
        system=BULGU_SUMMARY_SYSTEM,
        messages=[{"role": "user", "content": user_msg}],
    )
    llm_text = msg.content[0].text.strip()
    meta = _llm_meta_from_usage(msg.usage)
    meta["source"] = "llm"
    issues = validate_summary_text(llm_text, summaries)
    if issues:
        meta["summary_validation"] = issues
        meta["template_fallback"] = True
        return template_text, meta
    return llm_text, meta


def compact_summaries_from_results(
    results: List[dict],
    bulgular: Optional[Dict[str, str]] = None,
) -> List[dict]:
    return build_summaries_from_results(results, bulgular)

PLAN_SYSTEM = """
Sen akademik araştırma metodolojisi uzmanısın.
Verilen değişkenlere göre EN ANLAMLI analiz planını öner.
Amacın kullanıcıya gereksiz tablolar değil, araştırma sorusunu
yanıtlayan özlü ve güçlü bir analiz sunmak.

━━━ GENEL KURALLAR ━━━

1. TABLO SAYISINI MİNİMİZE ET
   - Aynı gruplandırma değişkeni için tüm outcome'ları
     TEK tabloda birleştir (3 ölçek × cinsiyet = 1 tablo)
   - Gereksiz testleri recommended:false yap

2. HANGİ TESTİ NE ZAMAN ÖNER ━━━
   
   KESINLIKLE ÖNERİLEN (recommended: true):
   - Tanımlayıcı istatistikler (her zaman)
   - Normallik testi (her zaman)
   - Ana demografikler için frekans (cinsiyet, bölüm/departman)
   - Ana gruplandırma × tüm outcome'lar (ANOVA/t-test)
   - Ölçek puanları arası korelasyon
   - Araştırma hipoteziyle doğrudan ilgili testler

   OPSİYONEL (recommended: false):
   - İkincil demografikler (medeni durum, gelir, kronik hastalık vb.)
   - Birbirine çok benzer testler (sigara ve alkol aynı dağılımdaysa)
   - Zaten kategorik versiyonu olan sürekli değişkenin testi
     (vki varken vki_kategori de varsa birini seç)

   ÖNERME (listeye ekleme):
   - Sürekli demografik değişkenler için ANOVA/t-test
     (yaş, boy, kilo → sadece tanımlayıcıda göster)
   - Korelasyon matrisine yaş/boy/kilo/vki ekleme
     (sadece ölçek puanları arası korelasyon yeterli)
   - Regresyon: sadece korelasyonda anlamlı bulunan çiftler için
     (r > .20 ve p < .05 ise önerilebilir, aksi halde önerme)

3. REASON YAZI ━━━
   Her test için neden önerildiğini/önerilmediğini yaz:
   - "Ana bağımsız değişken, araştırma sorusuyla doğrudan ilgili"
   - "İkincil demografik, araştırma hipoteziyle zayıf ilişki"
   - "Sürekli demografik, tanımlayıcıda yeterli"
   - "Normallik sağlandı, Pearson korelasyon uygun"

4. FREKANS TABLOLARI ━━━
   Sadece şunlar için frekans tablosu öner:
   - Ana demografikler (cinsiyet, bölüm)
   - Kategorik outcome'lar (risk grupları, BKİ kategorisi)
   İkincil demografikler için frekans recommended:false olsun

SADECE JSON döndür:
{
  "tests": [
    {
      "id": "descriptive",
      "type": "descriptive",
      "label": "Tanımlayıcı İstatistikler",
      "variables": ["OYS_TOPLAM", "NEQ_TOPLAM"],
      "reason": "Her araştırmada zorunlu",
      "recommended": true,
      "count": 1
    }
  ],
  "notes": "Genel metodoloji notu"
}
"""

CLASSIFY_SYSTEM = """
Sen deneyimli bir akademik istatistikçisin.
Araştırma verilerini analiz ederken önce SPSS metadata'sını,
sonra etiketleri, sonra veri dağılımını dikkate alırsın.

━━━ KARAR HİYERARŞİSİ ━━━

1. SPSS variable_measure varsa ÖNCE ONU KULLAN:
   nominal  → type: categorical
   ordinal  → type: categorical (Likert maddesi veya sıralı kategori)
   scale    → type: continuous

2. Etiket varsa etiket bilgisini kullan:
   "Gender", "Cinsiyet" → categorical + grouping
   "Age", "Yaş" → continuous + grouping (ama ANOVA hedefi değil)
   "Optimism Scale", "Ölçek Toplam" → continuous + outcome
   "Risk Group", "Kategori" → categorical + outcome

3. Değer aralığından tahmin et:
   unique ≤ 8 veya değerler {0,1} veya {1,2} → categorical
   aralık > 20 veya unique/n > 0.3 → continuous

━━━ TYPE KURALLARI ━━━

continuous: Sayısal, geniş aralık, hesaplama anlamlı
  → ölçek toplamı, yaş, boy, kilo, VKİ, sigara sayısı

categorical: Grup/sınıf, hesaplama anlamsız
  → cinsiyet, bölüm, evet/hayır, risk grubu, eğitim düzeyi

exclude: Analiz dışı
  → ID/sıra no, madde kolonları (op1, pss3, ani2_r),
    log/sqrt dönüşümleri (LG10_, LOG_, SQRT_)

━━━ ROLE KURALLARI ━━━

grouping (bağımsız/demografik):
  Katılımcıları gruplara ayırır, t-test/ANOVA/ki-kare için bağımsız değişken.

  KESİNLİKLE GROUPING:
  - Cinsiyet/gender (nominal, 2 kategori)
  - Bölüm/departman/fakülte (nominal, az kategori)
  - Eğitim düzeyi/medeni durum/gelir grubu (ordinal/nominal)
  - Evet/Hayır demografikler: sigara, alkol, kronik hastalık, ilaç
  - Stres kaynağı, meslek, şehir/bölge
  - Yaş: continuous + grouping (t-test/ANOVA bağımsız değişkeni olabilir)
    AMA: kategorik yaş grubu varsa (agegp) → yaş ham değeri outcome'a geç

outcome (bağımlı/sonuç):
  Analiz edilen, ölçülen sonuç değişkeni.

  KESİNLİKLE OUTCOME:
  - Ölçek toplam puanları (toptim, tpstress, tslfest, OYS_TOPLAM)
  - Risk grupları, BKİ kategorisi, puan kategorisi (_grubu, _kategori, _binary)
  - Antropometrik ölçümler (boy, kilo, VKİ ham değeri, smokenum)
    → sadece tanımlayıcıda gösterilir, t-test/ANOVA hedefi DEĞİL

  KURAL — hem ham hem kategorik varsa:
  age=24 + agegp3=1,2,3 → age=outcome/continuous, agegp3=grouping/categorical
  vki=22.5 + VKI_Kategori=Normal → vki=outcome/continuous, VKI_Kategori=outcome/categorical

━━━ DİSİPLİN KONVANSİYONLARI ━━━

Psikoloji: _r/_rev suffix → ters kodlu madde → exclude
           t prefix (toptim, tpstress) → ölçek toplamı → outcome/continuous

Sağlık: SBP, DBP, HR, BMI, VKİ → sağlık ölçümü → outcome/continuous
        smokenum, drinknum, freq → miktar → outcome/continuous

Eğitim: _score, _puan, _total → outcome/continuous
        school, class, grade, sinif → grouping/categorical

Genel: _toplam, _total, _sum, _score → outcome/continuous
       _grup, _grubu, _category, _binary → outcome/categorical
       id, no, anket_no → exclude

━━━ RECOMMENDED KURALI ━━━
- Tüm outcome/continuous ölçek puanları → true
- Ana demografikler (cinsiyet, bölüm) → true
- İkincil demografikler (medeni, gelir, kronik) → false
- Ham antropometrik (boy, kilo, smokenum) → false
- Kategorik outcome (risk grubu, BKİ kategorisi) → true

SADECE JSON döndür:
{
  "variables": {
    "col_name": {
      "type": "categorical|continuous|exclude",
      "role": "grouping|outcome|exclude",
      "recommended": true|false,
      "reason": "max 10 kelime Türkçe açıklama"
    }
  }
}
"""

def normalize_for_match(text: str) -> str:
    """Eşleştirme için string normalize et."""
    if not text:
        return ""
    s = str(text).translate(_TR_ASCII)
    s = s.lower()
    s = re.sub(r"[\s\-._]+", " ", s)
    words = []
    for word in s.split():
        word = re.sub(r"\d+$", "", word)
        if word:
            words.append(word)
    s = " ".join(words)
    s = re.sub(r"\s+", " ", s).strip()
    return s

def _keyword_prefixes(keyword: str, min_len: int = 3) -> List[str]:
    if " " in keyword or len(keyword) < min_len:
        return []
    return [keyword[:i] for i in range(min_len, len(keyword) + 1)]

def extract_scale_keywords(scale_name: str) -> List[str]:
    """Kullanıcının yazdığı ölçek adından eşleştirme keyword'leri üret."""
    keywords: List[str] = []
    raw = (scale_name or "").strip()
    if not raw:
        return []

    raw_no_paren = re.sub(r"\([^)]*\)", " ", raw)
    normalized_full = normalize_for_match(raw_no_paren)
    if normalized_full and " " not in normalized_full:
        keywords.append(normalized_full)

    paren_match = re.search(r"\(([^)]+)\)", raw)
    if paren_match:
        inner = normalize_for_match(paren_match.group(1))
        if inner:
            keywords.append(inner)

    words_raw = re.sub(r"\([^)]*\)", "", raw)
    words_raw = re.sub(r"[^\w\sğüşıöçĞÜŞİÖÇ]", " ", words_raw, flags=re.I)
    word_tokens = [w for w in words_raw.split() if w]
    if word_tokens:
        acronym = "".join(normalize_for_match(w)[:1] for w in word_tokens if w)
        if acronym:
            keywords.append(acronym)

    for word in word_tokens:
        norm_word = normalize_for_match(word)
        if len(norm_word) >= 3:
            keywords.append(norm_word)

    if word_tokens:
        first = normalize_for_match(word_tokens[0])
        if first:
            keywords.append(first)

    expanded: List[str] = []
    for kw in keywords:
        expanded.append(kw)
        expanded.extend(_keyword_prefixes(kw))

    seen = set()
    result: List[str] = []
    for kw in expanded:
        kw = kw.strip()
        if kw and kw not in seen:
            seen.add(kw)
            result.append(kw)
    return result

def _column_stem(col: str) -> str:
    base = re.split(r"[_\-\s]", col, maxsplit=1)[0]
    return normalize_for_match(base)

def _keyword_matches_column(kw: str, col: str) -> bool:
    norm_col = normalize_for_match(col)
    stem = _column_stem(col)
    if not kw or not stem:
        return False
    if stem == kw:
        return True
    if kw.startswith(stem) and len(kw) - len(stem) <= 2:
        return True
    if stem.startswith(kw) and len(stem) - len(kw) <= 1:
        return True
    if len(kw) >= 4 and kw in norm_col:
        return True
    return False

def _column_matches_keywords(col: str, keywords: List[str]) -> bool:
    return any(_keyword_matches_column(kw, col) for kw in keywords)

def _best_keyword_len(col: str, keywords: List[str]) -> int:
    matched = [len(kw) for kw in keywords if _keyword_matches_column(kw, col)]
    return max(matched) if matched else 0

def _classify_matched_column(col: str) -> str:
    norm = normalize_for_match(col)
    if any(marker in norm for marker in _TOTAL_MARKERS):
        return "total"
    if _ITEM_COL_RE.search(col):
        return "item"
    return "subscale"

def _match_confidence(
    total_columns: List[str],
    item_columns: List[str],
    subscale_columns: List[str],
) -> str:
    if total_columns:
        return "high"
    if item_columns:
        return "medium"
    if subscale_columns:
        return "low"
    return "low"

def match_scale_to_columns(
    scale_name: str,
    all_columns: List[str],
    aliases: Optional[List[str]] = None,
) -> dict:
    """Tek ölçeği kolon listesiyle eşleştir. aliases: registry'den alternatif isimler."""
    all_names = [scale_name]
    if aliases:
        all_names.extend(aliases)

    all_keywords: List[str] = []
    seen_kw: set = set()
    for name in all_names:
        for kw in extract_scale_keywords(name):
            if kw not in seen_kw:
                seen_kw.add(kw)
                all_keywords.append(kw)

    matched_columns: List[str] = []
    for col in all_columns:
        if _column_matches_keywords(col, all_keywords):
            matched_columns.append(col)

    item_columns = [c for c in matched_columns if _classify_matched_column(c) == "item"]
    total_columns = [c for c in matched_columns if _classify_matched_column(c) == "total"]
    subscale_columns = [
        c for c in matched_columns
        if c not in item_columns and c not in total_columns
    ]
    all_items = item_columns
    resolved = apply_scale_item_resolution(all_items)

    return {
        "scale_name": scale_name,
        "keywords_used": all_keywords,
        "matched_columns": matched_columns,
        "item_columns": resolved["items"],
        "cronbach_items": resolved["cronbach_items"],
        "item_count": resolved["item_count"],
        "all_item_columns": resolved["all_items"],
        "total_columns": total_columns,
        "subscale_columns": subscale_columns,
        "confidence": _match_confidence(total_columns, resolved["items"], subscale_columns),
    }

def match_all_scales(scale_names: List[str], all_columns: List[str]) -> dict:
    """Tüm ölçekleri eşleştir; her kolon en fazla bir ölçeğe gitsin."""
    from scale_registry import get_scale_info, resolve_scale_id

    raw_matches: List[dict] = []
    for name in scale_names:
        if not name or not name.strip():
            continue
        stripped = name.strip()
        aliases: List[str] = []
        sid = resolve_scale_id(stripped)
        if sid:
            entry = get_scale_info(sid)
            if entry:
                reg_names = entry.get("names") or []
                norm = stripped.lower()
                aliases = [n for n in reg_names if n.strip().lower() != norm]
        raw_matches.append(match_scale_to_columns(stripped, all_columns, aliases=aliases or None))

    col_owner: Dict[str, Tuple[int, int]] = {}
    for scale_idx, match in enumerate(raw_matches):
        for col in match["matched_columns"]:
            kw_len = _best_keyword_len(col, match["keywords_used"])
            if col not in col_owner or kw_len > col_owner[col][1]:
                col_owner[col] = (scale_idx, kw_len)

    final_matches: List[dict] = []
    for scale_idx, match in enumerate(raw_matches):
        owned = [
            col for col, (owner_idx, _) in col_owner.items()
            if owner_idx == scale_idx
        ]
        item_columns = [c for c in owned if _classify_matched_column(c) == "item"]
        total_columns = [c for c in owned if _classify_matched_column(c) == "total"]
        subscale_columns = [
            c for c in owned
            if c not in item_columns and c not in total_columns
        ]
        resolved = apply_scale_item_resolution(item_columns)
        final_matches.append({
            "scale_name": match["scale_name"],
            "keywords_used": match["keywords_used"],
            "matched_columns": owned,
            "item_columns": resolved["items"],
            "cronbach_items": resolved["cronbach_items"],
            "item_count": resolved["item_count"],
            "all_item_columns": resolved["all_items"],
            "total_columns": total_columns,
            "subscale_columns": subscale_columns,
            "confidence": _match_confidence(total_columns, resolved["items"], subscale_columns),
        })

    assigned = set(col_owner.keys())
    unmatched_columns = [c for c in all_columns if c not in assigned]

    return {
        "matches": final_matches,
        "unmatched_columns": unmatched_columns,
    }

def extract_full_text(file_base64: str, file_type: str) -> Tuple[str, int]:
    """PDF veya DOCX'ten tüm metni çek."""
    try:
        raw = base64.b64decode(file_base64)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Base64 decode hatası: {e}")

    if file_type == "pdf":
        try:
            reader = PdfReader(io.BytesIO(raw))
            pages = [(p.extract_text() or "").strip() for p in reader.pages]
            text = "\n\n".join(p for p in pages if p)
            return text, len(reader.pages)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"PDF okunamadı: {e}")

    if file_type == "docx":
        try:
            doc = Document(io.BytesIO(raw))
            paragraphs = [(p.text or "").strip() for p in doc.paragraphs if (p.text or "").strip()]
            text = "\n".join(paragraphs)
            page_count = max(1, (len(paragraphs) + 4) // 5)
            return text, page_count
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"DOCX okunamadı: {e}")

    raise HTTPException(status_code=400, detail="file_type 'pdf' veya 'docx' olmalı")

def _split_paragraphs(full_text: str) -> List[str]:
    if "\n\n" in full_text:
        return [p.strip() for p in full_text.split("\n\n") if p.strip()]
    return [p.strip() for p in full_text.split("\n") if p.strip()]

def _label_search_keywords(label: str) -> List[str]:
    keywords = extract_scale_keywords(label)
    norm_label = normalize_for_match(label)
    if norm_label and " " not in norm_label:
        keywords.insert(0, norm_label)
    for word in re.sub(r"\([^)]*\)", "", label).split():
        norm_word = normalize_for_match(word)
        if len(norm_word) >= 3:
            keywords.append(norm_word)
    seen = set()
    result: List[str] = []
    for kw in keywords:
        if kw and kw not in seen:
            seen.add(kw)
            result.append(kw)
    return result

def find_relevant_paragraphs(full_text: str, label: str, window: int = 3) -> str:
    """Etiket kelimelerine göre ilgili paragrafları bul."""
    if not full_text or not label:
        return ""

    paragraphs = _split_paragraphs(full_text)
    if not paragraphs:
        return ""

    keywords = _label_search_keywords(label)
    if not keywords:
        return ""

    matched_indices: List[int] = []
    for i, para in enumerate(paragraphs):
        norm_para = normalize_for_match(para)
        if any(kw in norm_para for kw in keywords if len(kw) >= 3):
            matched_indices.append(i)

    if not matched_indices:
        return ""

    selected: List[int] = []
    seen_idx = set()
    for idx in matched_indices:
        for j in range(max(0, idx - window), min(len(paragraphs), idx + window + 1)):
            if j not in seen_idx:
                seen_idx.add(j)
                selected.append(j)

    combined = "\n\n".join(paragraphs[i] for i in selected)
    if len(combined) > _CONTEXT_MAX_CHARS:
        combined = combined[:_CONTEXT_MAX_CHARS]
    return combined

def build_label_context_map(
    full_text: str,
    column_labels: List[ColumnLabel],
) -> Dict[str, str]:
    """Her kolon etiketi için ilgili paragrafları çıkar."""
    result: Dict[str, str] = {}
    for cl in column_labels:
        result[cl.column] = find_relevant_paragraphs(full_text, cl.label)
    return result

def _section_matches_keywords(text: str) -> bool:
    lower = text.lower()
    normalized = lower.translate(_TR_ASCII)
    normalized_keywords = [kw.translate(_TR_ASCII) for kw in ETHICS_KEYWORDS]
    return (
        any(kw in lower for kw in ETHICS_KEYWORDS) or
        any(kw in normalized for kw in normalized_keywords)
    )

def _extract_pdf_sections(file_bytes: bytes) -> Tuple[List[str], int]:
    reader = PdfReader(io.BytesIO(file_bytes))
    sections = []
    for page in reader.pages:
        try:
            text = (page.extract_text() or "").strip()
        except Exception:
            text = ""
        if len(text) > 50:
            sections.append(text)
    return sections, len(reader.pages)

def _extract_docx_sections(file_bytes: bytes) -> Tuple[List[str], int]:
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [(p.text or "").strip() for p in doc.paragraphs if (p.text or "").strip()]
    if not paragraphs:
        return [], 1
    chunk_size = 5
    sections = []
    for i in range(0, len(paragraphs), chunk_size):
        sections.append("\n".join(paragraphs[i:i + chunk_size]))
    return sections, max(1, len(sections))

def _filter_ethics_sections(sections: List[str]) -> Tuple[str, int, int]:
    total = len(sections)
    matched = [s for s in sections if _section_matches_keywords(s)]
    filtered_count = len(matched)

    if matched:
        combined = "\n\n".join(matched)
        if len(combined) > ETHICS_MAX_FILTERED_CHARS:
            combined = combined[:ETHICS_MAX_FILTERED_CHARS]
        return combined, filtered_count, total

    full_text = "\n\n".join(sections)
    if len(full_text) <= ETHICS_FALLBACK_CHARS:
        return full_text, 0, total

    chunk = ETHICS_FALLBACK_CHARS // 3
    start = full_text[:chunk]
    mid_start = len(full_text) // 2 - chunk // 2
    middle = full_text[mid_start:mid_start + chunk]
    end = full_text[-chunk:]
    combined = f"{start}\n\n[...]\n\n{middle}\n\n[...]\n\n{end}"
    return combined, 0, total

def _build_ethics_user_prompt(filtered_text: str) -> str:
    return f"""Aşağıdaki metinden araştırmada kullanılan TÜM ölçekleri çıkar.

Her ölçek için şunları bul (bulamazsan null bırak):
- name: Ölçeğin tam adı (Türkçe ve/veya İngilizce)
- short_name: Kısaltma (örn: BDI, OYŞTÖ, NEQ)
- item_count: Madde/soru sayısı (integer)
- subscales: Alt boyutlar listesi (string array)
- likert_range: Likert aralığı (örn: "1-5", "0-4")
- min_score / max_score: Teorik min-max puan
- cutoffs: Kesim noktaları (örn: [{{"label": "Normal", "range": "0-9", "score": 9}}])
- cronbach: Raporlanan Cronbach alpha değeri (float)
- developer: Ölçeği geliştiren kişi/kurum
- adaptation: Türkçe uyarlamacısı
- citation: Kaynak/referans metni

Metin:
{filtered_text}

Yanıt formatı:
{{
  "scales": [...],
  "research_topic": "araştırmanın konusu (1-2 cümle)",
  "filtered_page_count": kaç sayfa/bölüm analiz edildi (integer),
  "total_page_count": toplam sayfa sayısı (integer)
}}"""


async def generate_plan_ai(
    variables: List[Variable],
    research_topic: str,
    norm_results: Optional[dict] = None,
) -> Optional[List[dict]]:
    if not ANTHROPIC_API_KEY:
        return None

    grouping_vars = [v for v in variables if v.included and v.role == "grouping"]
    outcome_vars = [v for v in variables if v.included and v.role == "outcome"]

    var_summary = []
    for v in grouping_vars + outcome_vars:
        normal = None
        if norm_results and v.name in norm_results:
            normal = norm_results[v.name].get("normal")
        var_summary.append({
            "name": v.name,
            "label": v.label,
            "type": v.type,
            "role": v.role,
            "normal": normal,
        })

    user_msg = f"""Araştırma konusu: {research_topic or 'Belirtilmemiş'}

Değişkenler:
{json.dumps(var_summary, ensure_ascii=False, indent=2)}

Bu değişkenler için istatistiksel analiz planı oluştur."""

    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        msg = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=1000,
            system=PLAN_SYSTEM,
            messages=[{"role": "user", "content": user_msg}],
        )
        parsed = _parse_llm_json(msg.content[0].text.strip())
        tests = parsed.get("tests", [])
        return _normalize_ai_plan_tests(tests) if tests else None
    except Exception:
        return None


TOPLAM_RE = re.compile(r"_(TOPLAM|TOTAL|PUAN|SCORE|SUM)$", re.I)


def _locked_outcome_column_names(columns: List[str]) -> List[str]:
    return [c for c in columns if TOPLAM_RE.search(c)]


def _apply_toplam_outcome_to_variables(variables: Optional[List[Variable]]) -> None:
    """_TOPLAM suffix'li değişkenler her zaman outcome/continuous."""
    if not variables:
        return
    for var in variables:
        if TOPLAM_RE.search(var.name):
            var.type = "continuous"
            var.role = "outcome"


def _apply_toplam_outcome_to_classify(
    columns: List[str],
    categorical: List[str],
    continuous: List[str],
    exclude: List[str],
    recommendations: Dict[str, dict],
) -> Tuple[List[str], List[str], List[str], Dict[str, dict]]:
    cat = list(categorical)
    cont = list(continuous)
    excl = list(exclude)
    recs = dict(recommendations)
    for col in _locked_outcome_column_names(columns):
        if col in excl:
            excl.remove(col)
        if col in cat:
            cat.remove(col)
        if col not in cont:
            cont.append(col)
        prev = recs.get(col, {})
        recs[col] = {
            **prev,
            "role": "outcome",
            "status": prev.get("status") or "recommended",
            "ai_status": "approved",
            "reason": prev.get("reason") or "Ölçek toplam puanı",
        }
    return cat, cont, excl, recs


def run_classify(req: ClassifyRequest) -> dict:
    from ai_pipeline import run_variable_ai_pipeline
    from data_cleaning import normalize_variable_labels, prepare_analysis_df
    import pandas as pd

    df = None
    variables = None
    if req.data:
        rows = [r.values for r in req.data]
        df = pd.DataFrame(rows)
        variables = [
            Variable(
                name=c,
                label=(req.labels or {}).get(c, c),
                type="continuous",
                role="outcome",
                included=True,
            )
            for c in req.columns
        ]
        variables = normalize_variable_labels(variables)
        df = prepare_analysis_df(df, variables, req.missing_codes)
        _apply_toplam_outcome_to_variables(variables)

    result = run_variable_ai_pipeline(req, df=df, variables=variables)

    if result.get("manual_required") and not result.get("categorical") and not result.get("continuous"):
        if not has_claude() and not has_gemini_enrich():
            raise HTTPException(
                status_code=503,
                detail="AI katmanı kullanılamıyor — manuel sınıflandırma kullanın",
            )

    categorical, continuous, exclude, recommendations = _apply_toplam_outcome_to_classify(
        req.columns,
        result.get("categorical") or [],
        result.get("continuous") or [],
        result.get("exclude") or [],
        result.get("recommendations") or {},
    )

    return {
        "categorical": categorical,
        "continuous": continuous,
        "exclude": exclude,
        "recommendations": recommendations,
        "derived": result.get("derived") or [],
        "derived_approved": result.get("derived_approved") or [],
        "derived_review": result.get("derived_review") or [],
        "scales": result.get("scales") or [],
        "research_context": result.get("research_context") or {},
        "llm_meta": result.get("llm_meta") or {},
        "data_profile": result.get("data_profile"),
        "gemini_analysis": result.get("gemini_analysis"),
        "derivative_decision": result.get("derivative_decision"),
        "manual_required": result.get("manual_required", False),
    }


def _enrich_scale_with_registry(scale: dict) -> dict:
    """Registry'den reverse_items ve scale_range ekle."""
    from scale_registry import load_registry

    sid = scale.get("registry_id") or scale.get("id")
    if not sid:
        return scale

    entry = next((s for s in load_registry() if s.get("id") == sid), None)
    if not entry:
        return scale

    result = dict(scale)
    if "reverse_items" not in result or not result["reverse_items"]:
        result["reverse_items"] = entry.get("reverse_items") or []
    if "scale_range" not in result or not result.get("scale_range"):
        sr = entry.get("scale_range")
        result["scale_range"] = list(sr) if sr else [0, 4]
    return result


def _build_detect_scales_response(
    scales: List[dict],
    registry_matches: List[dict],
) -> dict:
    from scale_registry import get_cutoff, resolve_scale_id, validate_turkish

    registry_ids = {m["scale"]["id"] for m in registry_matches}
    registry_matched = []
    for m in registry_matches:
        if m.get("confidence") != "high":
            continue
        sc = m["scale"]
        registry_matched.append({
            "id": sc.get("id"),
            "name": (sc.get("names") or [sc.get("id", "")])[0],
            "confidence": m.get("confidence", "high"),
            "cols": m.get("matched_cols") or [],
            "reverse_items": sc.get("reverse_items") or [],
            "scale_range": sc.get("scale_range") or [0, 4],
        })

    registry_unmatched = []
    cutoffs: Dict[str, dict] = {}
    enriched_scales = [_enrich_scale_with_registry(s) for s in scales]
    for scale in enriched_scales:
        sid = scale.get("id") or scale.get("registry_id") or resolve_scale_id(scale.get("name", ""))
        if sid:
            scale.setdefault("id", sid)
            scale["turkish_valid"] = validate_turkish(sid)
            cutoff = get_cutoff(sid)
            if cutoff:
                cutoffs[sid] = cutoff
        in_registry = sid and sid in registry_ids
        source = scale.get("source", "")
        if not in_registry or source == "gemini":
            registry_unmatched.append({
                "name": scale.get("name"),
                "items": scale.get("items") or [],
                "id": sid,
            })

    return {
        "scales": enriched_scales,
        "registry_matched": registry_matched,
        "registry_unmatched": registry_unmatched,
        "cutoffs": cutoffs,
    }


def run_detect_scales(req: DetectScalesRequest) -> dict:
    from karar_verici import apply_reverse_item_decisions
    from veri_analisti import run_scale_detection, scales_from_registry_matches

    item_pattern = re.compile(
        r"^(?:recoded_|rev_|inv_)?[a-zA-Z]+_\d+"
        r"(?:_reversed|_recoded|_inverted|_ters|_rev|_rc|_inv|_t|_r)?$",
        re.IGNORECASE,
    )
    item_cols = [c for c in req.columns if item_pattern.match(c)]

    if len(item_cols) < 2:
        return {
            "scales": [],
            "registry_matched": [],
            "registry_unmatched": [],
            "cutoffs": {},
        }

    prefix_groups: Dict[str, list] = defaultdict(list)
    for col in item_cols:
        prefix = re.match(r"^([a-zA-Z]+)_", col)
        if prefix:
            prefix_groups[prefix.group(1)].append(col)

    valid_groups = {k: v for k, v in prefix_groups.items() if len(v) >= 3}
    col_labels = getattr(req, "labels", None) or {}

    scales, registry_matches, llm_meta = run_scale_detection(
        item_cols,
        col_labels=col_labels,
        prefix_groups=valid_groups,
        document_context=getattr(req, "document_context", None),
    )

    if not scales and valid_groups:
        registry_matches = registry_matches or []
        from scale_registry import match_scale
        if not registry_matches:
            registry_matches = match_scale(item_cols, col_labels)
        scales = scales_from_registry_matches(registry_matches, min_confidence="medium")

    if not scales and valid_groups:
        fallback_names = {"oys": "OYŞTÖ", "neq": "GYA", "sbito": "SBİTO"}
        for prefix, cols in valid_groups.items():
            resolved = apply_scale_item_resolution(cols)
            scales.append({
                "name": fallback_names.get(prefix.lower(), prefix.upper()),
                "items": resolved["items"],
                "cronbach_items": resolved["cronbach_items"],
                "item_count": resolved["item_count"],
                "source": "prefix_fallback",
            })

    scales = apply_reverse_item_decisions(scales)
    out = _build_detect_scales_response(scales, registry_matches)
    out["llm_meta"] = llm_meta
    return out


def import_spss_tables_service(req: SpssTableRequest) -> dict:
    method = "pandas"
    pandas_err = None
    try:
        results, meta = convert_spss_to_apa_results(req.content)
    except Exception as err:
        pandas_err = err
        if not ANTHROPIC_API_KEY:
            raise HTTPException(status_code=400, detail=str(pandas_err))
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        msg = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=4000,
            system=SPSS_CONVERT_SYSTEM,
            messages=[{"role": "user", "content": req.content}],
        )
        results = _markdown_tables_to_apa_results(msg.content[0].text.strip())
        if not results:
            raise HTTPException(status_code=400, detail=str(pandas_err))
        meta = {"source": "spss", "intro": "", "table_count": len(results), "ai_fallback": True}
        method = "ai"

    bulgular: Dict[str, str] = {}
    if req.auto_bulgu:
        for i, result in enumerate(results):
            try:
                text, _, _ = generate_bulgu(result, all_results=results)
                if text:
                    bulgular[str(i)] = text
            except Exception:
                pass

    return sanitize({
        "results": results,
        "bulgular": bulgular,
        "meta": meta,
        "method": method,
        "pandas_error": str(pandas_err) if pandas_err else None,
    })


def run_import_ethics_report(
    file_bytes: bytes,
    filename: str,
    research_topic: Optional[str] = None,
) -> dict:
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY ayarlanmamış")

    if not (filename.endswith(".pdf") or filename.endswith(".docx")):
        raise HTTPException(status_code=400, detail="Sadece PDF veya DOCX dosyaları desteklenir")

    if not file_bytes:
        raise HTTPException(status_code=400, detail="Dosya boş")

    try:
        if filename.endswith(".pdf"):
            sections, total_pages = _extract_pdf_sections(file_bytes)
        else:
            sections, total_pages = _extract_docx_sections(file_bytes)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Metin çıkarılamadı: {e}")

    if not sections:
        raise HTTPException(status_code=400, detail="Dosyadan metin çıkarılamadı")

    filtered_text, filtered_pages, total_pages = _filter_ethics_sections(sections)
    chars_sent = len(filtered_text)

    user_prompt = _build_ethics_user_prompt(filtered_text)
    if research_topic:
        user_prompt = f"Araştırma konusu (kullanıcı): {research_topic}\n\n{user_prompt}"

    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        msg = client.messages.create(
            model=ANTHROPIC_MODEL,
            max_tokens=4000,
            system=ETHICS_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": user_prompt}],
        )
        parsed = _parse_llm_json(msg.content[0].text.strip())
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI analizi başarısız: {e}")

    scales = parsed.get("scales") or []
    topic = parsed.get("research_topic") or research_topic or ""
    ai_filtered = parsed.get("filtered_page_count")
    ai_total = parsed.get("total_page_count")
    effective_filtered = filtered_pages if filtered_pages > 0 else total_pages

    return sanitize({
        "scales": scales,
        "research_topic": topic,
        "token_info": {
            "filtered_pages": ai_filtered if ai_filtered is not None else effective_filtered,
            "total_pages": ai_total if ai_total is not None else total_pages,
            "chars_sent": chars_sent,
        },
    })


LABEL_SYSTEM = """Sen akademik araştırma veri seti uzmanısın.
Verilen kolon adları için kısa, anlaşılır Türkçe değişken isimleri üret.

Kurallar:
- Kısaltmaları çöz: OYS → Online Yemek Sipariş, NEQ → Night Eating Questionnaire (Gece Yeme Anketi), vb.
- _TOPLAM → "Toplam Puanı", _GRUP → "Grubu", _BINARY → "(İkili)" ekle
- Demografikler: cinsiyet → "Cinsiyet", bolum → "Bölüm"
- Belge bağlamını kullan (ölçek adları, araştırma konusu)
- Etiket max 6-8 kelime olsun
- Çözemezsen kolon adının düzeltilmiş halini yaz

SADECE JSON döndür:
{
  "labels": {
    "KOLON_ADI": "Türkçe Etiket"
  }
}"""


def generate_labels_ai(
    columns: list[str],
    scale_names: list[str] | None = None,
    research_topic: str = "",
) -> tuple[dict[str, str], dict]:
    """Gemini Flash ile kalan boş etiketleri doldurur. Yoksa Haiku."""
    context_parts = []
    if scale_names:
        context_parts.append(f"Ölçekler: {', '.join(scale_names)}")
    if research_topic:
        context_parts.append(f"Araştırma konusu: {research_topic[:300]}")

    user = f"Kolon adları: {', '.join(columns)}"
    if context_parts:
        user += "\n" + "\n".join(context_parts)

    if has_gemini_enrich():
        text, meta = gemini_json_task(LABEL_SYSTEM, user, max_tokens=800)
        if text:
            parsed = _parse_json_object(text)
            return parsed.get("labels", {}), meta

    if has_claude():
        text, meta = claude_decide(LABEL_SYSTEM, user, max_tokens=800)
        parsed = _parse_json_object(text)
        return parsed.get("labels", {}), meta

    return {}, {}

