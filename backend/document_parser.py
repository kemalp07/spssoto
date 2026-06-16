"""Word (.docx) belge ayrıştırıcı — anket formu ve etik kurul raporu."""
from __future__ import annotations

import io
import re
from typing import Any, Dict, List, Optional, Tuple

from docx import Document

from schemas import Variable

_ITEM_RE = re.compile(
    r"^\s*(?:"
    r"(?P<num>\d+)\s*[.)]\s*|"
    r"(?:madde|soru)\s*(?P<num_m>\d+)\s*[.:)]?\s*|"
    r"(?P<letter>[a-zA-Z])\s*[.)]\s*"
    r")(?P<text>.+)$",
    re.IGNORECASE,
)
_REVERSE_HINT = re.compile(r"\(R\)|\(T\)|\bters\b", re.IGNORECASE)
_SECTION_HEADER = re.compile(
    r"^(?:bölüm|section|ölçek|anket|form)\s*[\d.:)\-]*\s*(.+)$",
    re.IGNORECASE,
)
_LIKERT_LINE = re.compile(
    r"(hiç\s+katılmıyorum|katılmıyorum|kararsız|katılıyorum|tamamen\s+katılıyorum)",
    re.IGNORECASE,
)
_AIM_HEADER = re.compile(
    r"^\s*(?:araştırmanın\s+)?(?:amaç|amacı|purpose|hedef)(?:\s*[:.)])?\s*(.*)$",
    re.IGNORECASE,
)
_HYPOTHESIS_RE = re.compile(
    r"^\s*(?:H|AS|HS|AH)\s*(\d+)\s*[:.)]\s*(.+)$",
    re.IGNORECASE,
)
_HYPOTHESIS_NUMBERED = re.compile(
    r"^\s*(?:Soru|Hipotez|Araştırma\s+Sorusu)\s*(\d+)\s*[:.)]\s*(.+)$",
    re.IGNORECASE,
)
_HYPOTHESIS_INLINE = re.compile(
    r"\b(hipotez|araştırma\s+sorusu|araştırma\s+problemi)\s*[:.)]?\s*(.+)$",
    re.IGNORECASE,
)
_HYPOTHESIS_SECTION_HEADER = re.compile(
    r"^\s*(?:araştırma\s+sorular[ıi]|hipotezler|araştırma\s+problemler[ıi]|alt\s+problemler)"
    r"\s*[:.]?\s*$",
    re.IGNORECASE,
)
_BULLET_ITEM = re.compile(
    r"^\s*[-–—•*]\s*(.+)$",
)
_NUMBERED_ITEM = re.compile(
    r"^\s*(\d+)\s*[.):\-]\s*(.+)$",
)
_NUMBERED_LOOSE = re.compile(
    r"^\s*(\d{1,3})\s+([\dA-Za-zÇçĞğİıÖöŞşÜü].{4,})$",
)
_SAMPLE_N = re.compile(
    r"(\d{2,5})\s*(?:katılımcı|öğrenci|kişi|denek|gönüllü)",
    re.IGNORECASE,
)
_QUOTED_NAME = re.compile(r'"([^"]{3,80})"|\'([^\']{3,80})\'')
_CAPS_NAME = re.compile(r"\b([A-ZÇĞİÖŞÜ]{3,20})\b")
_DATE_RE = re.compile(
    r"\b(\d{1,2}[./]\d{1,2}[./]\d{2,4}|\d{4}-\d{2}-\d{2})\b",
)
_INSTITUTION_KW = re.compile(
    r"(üniversite|fakülte|enstitü|hastane|etik\s+kurul)",
    re.IGNORECASE,
)
_INSTITUTION_EXCLUDE = re.compile(
    r"amaç|amacı|hipotez|araştırma\s+sorusu|purpose",
    re.IGNORECASE,
)
_METHODOLOGY_KW = re.compile(
    r"\b(evren|örneklem|yöntem|veri\s+topla|katılımcı|gönüllü|"
    r"etik\s+onay|bilgilendirilmiş\s+onam|materyal|araç\s+ve\s+yöntem)\b",
    re.IGNORECASE,
)
_RESEARCH_Q_HINT = re.compile(
    r"(?:\?"
    r"|\bfark\b|\bfarkl"
    r"|\bilişki\b|\betki\b|\bdüzey\b|\boran\b"
    r"|\bkarşılaştır|\bincelen|\bbelirlen|\bsaptan|\bdeğerlendir"
    r"|\byordam|\btahmin|\baçıklar)\b",
    re.IGNORECASE,
)
_LIKERT_ANCHOR = re.compile(
    r"^(?:hiç|tamamen|katıl|kararsız|never|always|strongly|never|rarely|often)",
    re.IGNORECASE,
)
_MAX_HYPOTHESES = 8


def _docx_paragraphs(file_bytes: bytes) -> List[str]:
    return [text for text, _ in _docx_paragraph_entries(file_bytes)]


def _docx_paragraph_entries(file_bytes: bytes) -> List[Tuple[str, Optional[str]]]:
    """Paragraf metni ve Word stil adı (varsa)."""
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception:
        return []
    entries: List[Tuple[str, Optional[str]]] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style_name = para.style.name if para.style is not None else None
        entries.append((text, style_name))
    for table in doc.tables:
        for row in table.rows:
            cells = [((c.text or "").strip()) for c in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                entries.append((row_text, None))
    return entries


def _is_heading_style(style_name: Optional[str]) -> bool:
    if not style_name:
        return False
    name = style_name.strip()
    if name.startswith("Heading"):
        return True
    if name.startswith("Başlık"):
        return True
    return False


def _parse_item_line(line: str) -> Optional[Tuple[int, str, bool]]:
    line = line.strip()
    if not line:
        return None

    num_raw: Optional[str] = None
    text = ""

    m = _ITEM_RE.match(line)
    if m:
        num_raw = m.group("num") or m.group("num_m") or m.group("letter")
        text = (m.group("text") or "").strip()
    else:
        m = _NUMBERED_ITEM.match(line)
        if m:
            num_raw, text = m.group(1), (m.group(2) or "").strip()
        else:
            m = _NUMBERED_LOOSE.match(line)
            if m:
                num_raw, text = m.group(1), (m.group(2) or "").strip()
            elif "|" in line:
                cells = [c.strip() for c in line.split("|") if c.strip()]
                if len(cells) >= 2 and cells[0].isdigit():
                    num_raw, text = cells[0], cells[1]

    if not text:
        return None

    try:
        if num_raw and str(num_raw).isdigit():
            no = int(num_raw)
        elif num_raw:
            no = ord(str(num_raw).upper()[0]) - 64
        else:
            no = 0
    except (TypeError, ValueError):
        no = 0

    reverse = bool(_REVERSE_HINT.search(text))
    text = _REVERSE_HINT.sub("", text).strip(" -–—")
    if not text or len(text) < 8:
        return None
    if _LIKERT_ANCHOR.match(text):
        return None
    return no, text, reverse


def _is_plausible_hypothesis(text: str) -> bool:
    t = (text or "").strip()
    if len(t) < 25:
        return False
    if _METHODOLOGY_KW.search(t):
        return False
    if t.endswith("?"):
        return True
    return bool(_RESEARCH_Q_HINT.search(t))


def _append_hypothesis(found: List[str], seen: set, text: str, *, strict: bool) -> bool:
    text = (text or "").strip().rstrip(".")
    if not text or text in seen:
        return False
    if strict and not _is_plausible_hypothesis(text):
        return False
    if len(text) < 15:
        return False
    seen.add(text)
    found.append(text)
    return True


def _detect_likert_scale(block_text: str) -> str:
    lower = block_text.lower()
    hits = len(_LIKERT_LINE.findall(lower))
    if hits >= 3:
        if re.search(r"\b7\b|yedi\s*li|7\s*li", lower):
            return "likert_7"
        return "likert_5"
    if re.search(r"strongly\s+disagree|never\s+.*\s+always", lower):
        return "other"
    return "other"


def _is_section_header(line: str, style_name: Optional[str] = None) -> bool:
    if _is_heading_style(style_name):
        return True
    s = line.strip()
    if not s or len(s) > 120:
        return False
    if _ITEM_RE.match(s):
        return False
    if s.endswith("?"):
        return False
    if s.endswith(":"):
        return True
    if _SECTION_HEADER.match(s):
        return True
    if s.isupper() and len(s.split()) <= 8:
        return True
    return False


def parse_anket_docx(file_bytes: bytes) -> dict:
    """Anket formu .docx → bölümler, maddeler, ölçek tipi ipuçları."""
    try:
        entries = _docx_paragraph_entries(file_bytes)
        if not entries:
            return {"parse_error": True, "sections": []}

        sections: List[dict] = []
        current: Optional[dict] = None
        block_lines: List[str] = []

        def flush_section() -> None:
            nonlocal current, block_lines
            if current is None:
                return
            current["scale_type"] = _detect_likert_scale("\n".join(block_lines))
            sections.append(current)
            current = None
            block_lines = []

        for line, style_name in entries:
            if _is_section_header(line, style_name):
                flush_section()
                title = line.rstrip(":").strip()
                m = _SECTION_HEADER.match(title)
                if m and m.group(1):
                    title = m.group(1).strip()
                current = {"title": title, "items": [], "scale_type": "other"}
                block_lines = [line]
                continue

            item = _parse_item_line(line)
            if item:
                if current is None:
                    current = {"title": "Genel", "items": [], "scale_type": "other"}
                no, text, reverse = item
                current["items"].append({
                    "no": no,
                    "text": text,
                    "reverse_hint": reverse,
                })
                block_lines.append(line)
                continue

            if current is not None:
                block_lines.append(line)

        flush_section()

        if not sections:
            return {"parse_error": True, "sections": [], "raw_text": ""}
        raw_lines = [line for line, _ in entries]
        return {
            "sections": sections,
            "raw_text": "\n".join(raw_lines)[:8000],
        }
    except Exception:
        return {"parse_error": True, "sections": [], "raw_text": ""}


def anket_text_from_parse(anket: Optional[dict]) -> str:
    """Anket parse sonucundan Gemini'ye gidecek metin."""
    if not anket or anket.get("parse_error"):
        raw = (anket or {}).get("raw_text") or ""
        return raw.strip()[:8000]

    parts: List[str] = []
    for sec in anket.get("sections") or []:
        title = (sec.get("title") or "").strip()
        items = sec.get("items") or []
        item_lines = []
        for item in items:
            no = item.get("no")
            text = (item.get("text") or "").strip()
            if text:
                item_lines.append(f"{no}. {text}" if no is not None else text)
        block = "\n".join([x for x in [title, "\n".join(item_lines)] if x])
        if block.strip():
            parts.append(block)

    structured = "\n\n".join(parts).strip()
    raw = (anket.get("raw_text") or "").strip()
    if len(structured) >= 80:
        return structured[:8000]
    if raw:
        return raw[:8000]
    return structured


def etik_text_from_parse(etik: Optional[dict]) -> str:
    """Etik kurul parse sonucundan Gemini'ye gidecek metin."""
    if not etik:
        return ""
    parts: List[str] = []
    aim = (etik.get("aim") or "").strip()
    if aim:
        parts.append(f"Amaç: {aim}")
    hyps = etik.get("hypotheses") or []
    if hyps:
        parts.append("Araştırma soruları:")
        for h in hyps:
            if isinstance(h, str) and h.strip():
                parts.append(f"- {h.strip()}")
    stat_methods = (etik.get("statistical_methods") or "").strip()
    if stat_methods:
        parts.append(f"\nİstatistiksel yöntemler: {stat_methods}")
    scale_names = etik.get("scale_names") or []
    if scale_names:
        parts.append(f"Kullanılan ölçekler: {', '.join(str(s) for s in scale_names if s)}")
    structured = "\n".join(parts).strip()
    raw = (etik.get("raw_text") or "").strip()
    if structured and raw:
        return f"{structured}\n\n{raw[:6000]}"
    if structured:
        return structured[:8000]
    return raw[:8000]


def _extract_aim(lines: List[str]) -> Optional[str]:
    for i, line in enumerate(lines):
        m = _AIM_HEADER.match(line)
        if m:
            inline = (m.group(1) or "").strip()
            if inline:
                return inline[:2000]
            parts = []
            for follow in lines[i + 1: i + 6]:
                if _HYPOTHESIS_RE.match(follow) or _is_section_header(follow):
                    break
                if follow.strip():
                    parts.append(follow.strip())
            if parts:
                return " ".join(parts)[:2000]
    for line in lines:
        lower = line.lower()
        if "amaç" in lower or "purpose" in lower:
            return line[:2000]
    return None


def _extract_hypotheses(lines: List[str]) -> List[str]:
    """Regex ile hipotez/araştırma sorusu çıkar. Bulamazsa boş liste döner."""
    found: List[str] = []
    seen: set = set()
    in_section = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if len(found) >= _MAX_HYPOTHESES:
            break

        if _HYPOTHESIS_SECTION_HEADER.match(stripped):
            in_section = True
            continue

        m = _HYPOTHESIS_RE.match(stripped)
        if m:
            if _append_hypothesis(found, seen, m.group(2), strict=False):
                in_section = True
            continue

        m = _HYPOTHESIS_NUMBERED.match(stripped)
        if m:
            if _append_hypothesis(found, seen, m.group(2), strict=False):
                in_section = True
            continue

        if in_section:
            m = _BULLET_ITEM.match(stripped)
            if m:
                _append_hypothesis(found, seen, m.group(1), strict=True)
                continue

            m = _NUMBERED_ITEM.match(stripped)
            if m:
                _append_hypothesis(found, seen, m.group(2), strict=True)
                continue

            if _is_section_header(stripped):
                in_section = False
                continue

        m = _HYPOTHESIS_INLINE.search(line)
        if m and len(stripped) >= 30:
            _append_hypothesis(found, seen, m.group(2), strict=True)

    return found


def _extract_hypotheses_ai(full_text: str) -> List[str]:
    """Regex bulamazsa Gemini Flash ile araştırma sorularını çıkar."""
    if len(full_text.strip()) < 50:
        return []

    relevant = full_text
    lower = full_text.lower()
    for keyword in ["araştırma sorular", "hipotez", "araştırmanın amacı", "alt problem"]:
        idx = lower.find(keyword)
        if idx >= 0:
            start = max(0, idx - 500)
            end = min(len(full_text), idx + 3000)
            relevant = full_text[start:end]
            break

    system = """Verilen etik kurul raporu metninden araştırma sorularını veya hipotezleri çıkar.

Kurallar:
- Her araştırma sorusunu/hipotezi ayrı bir item olarak döndür
- Sorunun/hipotezin tam metnini yaz, kısaltma
- Sadece araştırma soruları/hipotezler, başka bilgi ekleme
- Yöntem, örneklem, etik onam gibi metodoloji maddelerini ekleme
- Genelde 3-6 madde olur; emin değilsen boş liste döndür
- En fazla 8 madde

SADECE JSON döndür:
{"hypotheses": ["soru1 metni", "soru2 metni"]}"""

    try:
        from llm_router import (
            _parse_json_object,
            claude_decide,
            gemini_json_task,
            has_claude,
            has_gemini_enrich,
        )

        text = ""
        if has_gemini_enrich():
            text, _ = gemini_json_task(system, relevant[:4000], max_tokens=1000)
        elif has_claude():
            text, _ = claude_decide(system, relevant[:4000], max_tokens=1000)

        if text:
            parsed = _parse_json_object(text)
            hyps = parsed.get("hypotheses", [])
            if isinstance(hyps, list):
                filtered: List[str] = []
                seen: set = set()
                for h in hyps:
                    if not isinstance(h, str):
                        continue
                    text = h.strip()
                    if not text or text in seen:
                        continue
                    if not _is_plausible_hypothesis(text):
                        continue
                    seen.add(text)
                    filtered.append(text)
                    if len(filtered) >= _MAX_HYPOTHESES:
                        break
                return filtered
    except Exception:
        pass

    return []


def _extract_sample_n(text: str) -> Optional[int]:
    m = _SAMPLE_N.search(text)
    if not m:
        return None
    try:
        return int(m.group(1))
    except ValueError:
        return None


def _extract_scale_names(lines: List[str]) -> List[str]:
    names: List[str] = []
    seen = set()
    blob = "\n".join(lines)
    for m in _QUOTED_NAME.finditer(blob):
        for g in m.groups():
            if g and g not in seen:
                seen.add(g)
                names.append(g.strip())
    for line in lines:
        for m in _CAPS_NAME.finditer(line):
            token = m.group(1)
            if token not in seen and len(token) >= 3:
                seen.add(token)
                names.append(token)
    return names[:20]


def _extract_institution(lines: List[str]) -> Optional[str]:
    for line in lines:
        s = line.strip()
        if not s or len(s) > 120:
            continue
        if _INSTITUTION_EXCLUDE.search(s):
            continue
        if _INSTITUTION_KW.search(s):
            return s[:300]
    return None


def _extract_date(text: str) -> str:
    m = _DATE_RE.search(text)
    return m.group(1) if m else ""


def parse_etik_kurul_docx(file_bytes: bytes) -> dict:
    """Etik kurul raporu .docx → amaç, hipotezler, örneklem vb."""
    try:
        lines = _docx_paragraphs(file_bytes)
        if not lines:
            return {"parse_error": True, "raw_text": ""}

        blob = "\n".join(lines)
        raw_text = blob[:4000]
        aim = _extract_aim(lines)
        hypotheses = _extract_hypotheses(lines)

        if not hypotheses:
            hypotheses = _extract_hypotheses_ai(blob)

        n_val = _extract_sample_n(blob)
        scale_names = _extract_scale_names(lines)
        institution = _extract_institution(lines)
        date = _extract_date(blob)

        STAT_KEYWORDS = [
            "istatistiksel yöntem", "shapiro", "t testi", "t-test",
            "anova", "mann-whitney", "kruskal", "pearson", "spearman",
            "korelasyon analizi", "normallik",
        ]
        statistical_methods = ""
        for para in lines:
            text = para.strip()
            if any(k in text.lower() for k in STAT_KEYWORDS) and len(text) > 50:
                statistical_methods = text[:800]
                break

        if not any([aim, hypotheses, n_val, scale_names, institution, date]):
            return {"parse_error": True, "raw_text": raw_text}

        return {
            "aim": aim,
            "hypotheses": hypotheses or None,
            "n": n_val,
            "scale_names": scale_names or None,
            "institution": institution or None,
            "date": date or None,
            "statistical_methods": statistical_methods or None,
            "raw_text": raw_text,
        }
    except Exception:
        return {"parse_error": True, "raw_text": ""}


def extract_scale_test_requirements(
    anket_result: dict,
    registry_matches: List[dict],
) -> dict:
    """
    Anket dosyasından doğrudan test gereksinimlerini çıkar.
    AI'a gerek yok — ölçek yapısı her şeyi söylüyor.
    """
    requirements: dict = {
        "cronbach_scales": [],
        "descriptive_vars": [],
        "correlation_pairs": [],
    }

    for match in registry_matches or []:
        scale = match.get("scale") or {}
        names = scale.get("names") or []
        requirements["cronbach_scales"].append({
            "id": scale.get("id"),
            "name": names[0] if names else scale.get("id"),
            "items": match.get("matched_cols") or [],
            "reverse_items": scale.get("reverse_items") or [],
            "scale_range": scale.get("scale_range") or [0, 4],
            "test": "cronbach",
            "reason": "Ölçek güvenilirliği — anket dosyasından tespit edildi",
        })

    return requirements


def resolve_scale_test_requirements(
    document_context: dict,
    column_names: List[str],
    variables: List[Variable],
) -> dict:
    """document_context içindeki anket + sütun eşleşmesinden gereksinimleri üret."""
    stored = (document_context or {}).get("test_requirements")
    if stored and stored.get("cronbach_scales"):
        return stored

    anket = (document_context or {}).get("anket") or {}
    if anket.get("parse_error"):
        return {"cronbach_scales": [], "descriptive_vars": [], "correlation_pairs": []}

    from scale_registry import match_scale

    labels = {
        v.name: v.label or v.name
        for v in variables
        if v.included
    }
    registry_matches = match_scale(column_names, labels)
    return extract_scale_test_requirements(anket, registry_matches)


def apply_scale_test_requirements(
    candidates: List[dict],
    requirements: dict,
    df: "pd.DataFrame",
) -> List[dict]:
    """Anket/registry gereksinimlerinden eksik Cronbach adaylarını ekle."""
    from test_planner import make_candidate_id

    existing = {
        tuple(sorted(c.get("vars") or []))
        for c in candidates
        if c.get("test") == "cronbach"
    }
    out = list(candidates)
    seq = max(
        (int(str(c.get("seq", "t0"))[1:]) for c in candidates if c.get("seq")),
        default=0,
    )

    for scale in requirements.get("cronbach_scales") or []:
        items = scale.get("items") or []
        if len(items) < 2:
            continue
        if not all(c in df.columns for c in items):
            continue
        key = tuple(sorted(items))
        if key in existing:
            continue
        seq += 1
        out.append({
            "id": make_candidate_id("cronbach", items),
            "seq": f"t{seq}",
            "test": "cronbach",
            "vars": items,
            "auto_flag": "uygun",
            "reason": scale.get("reason", "Ölçek güvenilirliği"),
            "scale_id": scale.get("id"),
        })
        existing.add(key)

    return out
