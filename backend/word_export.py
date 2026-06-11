"""Word APA export."""
import io
import re
from typing import List, Optional, Dict, Tuple
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from constants import APA_BORDER_SZ, _WORD_M_MARK, _WORD_SD_MARK, _WORD_STAT_RE
from formatting import apply_academic_text_rules, substitute_variable_codes, build_resolved_label_map

def _strip_apa_html(text: str) -> str:
    return re.sub(r"</?em>", "", str(text))

def _prepare_word_text(text: str) -> str:
    """Word export: x̄→M, SS→SD; HTML italik etiketlerini kaldır."""
    text = _strip_apa_html(str(text))
    text = text.replace("x̄", _WORD_M_MARK)
    text = re.sub(r"\bSS\b", _WORD_SD_MARK, text)
    return text

def _split_table_title(title: str) -> Tuple[str, str]:
    """'Tablo 1. Açıklama' → ('Tablo 1', 'Açıklama')."""
    title = str(title)
    plain = _strip_apa_html(title)
    m = re.match(r"^(Tablo\s+\d+)\.\s*(.*)$", plain, re.I)
    if not m:
        return title, ""
    cap_match = re.match(r"^Tablo\s+\d+\.\s*(.*)$", title, re.I)
    caption = cap_match.group(1).strip() if cap_match else m.group(2)
    return m.group(1), caption

def _set_cell_border(cell, **edges):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:tcBorders"):
            tcPr.remove(child)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        spec = edges.get(edge)
        if spec is None:
            continue
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), spec.get("val", "nil"))
        if spec.get("val") == "single":
            element.set(qn("w:sz"), str(spec.get("sz", APA_BORDER_SZ)))
            element.set(qn("w:color"), spec.get("color", "000000"))
            element.set(qn("w:space"), "0")
        tcBorders.append(element)
    tcPr.append(tcBorders)

def _clear_table_level_borders(table):
    """Tablo düzeyinde tüm kenarlıkları kapat; çizgiler yalnızca hücre düzeyinde."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for child in list(tblPr):
        if child.tag == qn("w:tblBorders"):
            tblPr.remove(child)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), "nil")
        tblBorders.append(border)
    tblPr.append(tblBorders)

def _apply_apa_table_borders(table):
    """
    APA 7 yatay çizgi yerleşimi (tam 3 çizgi):
      1. Başlık satırı üstü  → header hücre top
      2. Başlık / veri ayrımı → header hücre bottom
      3. Son veri satırı altı → son satır hücre bottom
    Dikey çizgi yok. Not metni tablo dışında (ayrı paragraf).
    """
    _clear_table_level_borders(table)

    nil = {"val": "nil"}
    single = {"val": "single", "sz": APA_BORDER_SZ, "color": "000000"}
    n_rows = len(table.rows)
    last_idx = n_rows - 1

    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if r_idx == 0:
                _set_cell_border(cell, top=single, left=nil, right=nil, bottom=single)
            elif r_idx == last_idx:
                _set_cell_border(cell, top=nil, left=nil, right=nil, bottom=single)
            else:
                _set_cell_border(cell, top=nil, left=nil, right=nil, bottom=nil)

def _add_word_runs(paragraph, text: str, bold: bool = False, force_italic: bool = False):
    """Word hücreleri: x̄→M, SS→SD; n,p,df,F,t,d,H,U italik."""
    text = _prepare_word_text(text)
    pos = 0
    for match in _WORD_STAT_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            run.bold = bold
            if force_italic:
                run.italic = True
        token = match.group()
        if token == "Cohen's d":
            r1 = paragraph.add_run("Cohen's ")
            r1.bold = bold
            if force_italic:
                r1.italic = True
            r2 = paragraph.add_run("d")
            r2.italic = True
            r2.bold = bold
        elif token == _WORD_M_MARK:
            run = paragraph.add_run("M")
            run.italic = True
            run.bold = bold
        elif token == _WORD_SD_MARK:
            run = paragraph.add_run("SD")
            run.italic = True
            run.bold = bold
        else:
            run = paragraph.add_run(token)
            run.italic = True
            run.bold = bold
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = bold
        if force_italic:
            run.italic = True

def _add_apa_table_title(doc: Document, title: str):
    num, caption = _split_table_title(title)
    p_num = doc.add_paragraph()
    run = p_num.add_run(num)
    run.bold = True
    if caption:
        p_cap = doc.add_paragraph()
        _add_word_runs(p_cap, caption, force_italic=True)

def _refine_word_header(header: str) -> str:
    """Ki-kare sütunlarından gereksiz 'n (%)' ekini kaldır."""
    h = _strip_apa_html(str(header))
    return re.sub(r"\s+n\s*\(%\)\s*$", "", h, flags=re.I).strip()

def polish_result_for_word(result: dict, label_map: Optional[Dict[str, str]] = None) -> dict:
    """Word export öncesi başlık/sütun/not rafine etme."""
    polished = dict(result)
    resolved = build_resolved_label_map(label_map)

    title = str(polished.get("title", ""))
    num, caption = _split_table_title(title)
    if caption:
        caption = apply_academic_text_rules(substitute_variable_codes(caption, resolved))
        polished["title"] = f"{num}. {caption}"
    elif resolved:
        polished["title"] = apply_academic_text_rules(substitute_variable_codes(title, resolved))

    headers = [_refine_word_header(h) for h in polished.get("headers", [])]
    if resolved:
        headers = [apply_academic_text_rules(substitute_variable_codes(h, resolved)) for h in headers]
    polished["headers"] = headers

    note = str(polished.get("note", ""))
    if resolved:
        note = apply_academic_text_rules(substitute_variable_codes(note, resolved))
    polished["note"] = note
    return polished

def _add_apa_note(paragraph, note_text: str):
    """Not. yalnızca italik; gövde düz, p/F/t/n/d istatistik sembolleri italik."""
    note_text = str(note_text).strip()
    if note_text.startswith("Not."):
        r = paragraph.add_run("Not.")
        r.italic = True
        _add_word_runs(paragraph, note_text[4:].lstrip())
    elif note_text.startswith("Note."):
        r = paragraph.add_run("Note.")
        r.italic = True
        _add_word_runs(paragraph, note_text[5:].lstrip())
    else:
        _add_word_runs(paragraph, note_text)

def add_apa_table(doc: Document, result: dict, label_map: Optional[Dict[str, str]] = None):
    result = polish_result_for_word(result, label_map)
    _add_apa_table_title(doc, result.get("title", ""))

    headers = result.get("headers", [])
    rows = result.get("rows", [])
    if not headers or not rows:
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Normal"
    except Exception:
        pass

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        _add_word_runs(cell.paragraphs[0], h, bold=True)

    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            if j < len(table.rows[i + 1].cells):
                cell = table.rows[i + 1].cells[j]
                cell.text = ""
                _add_word_runs(cell.paragraphs[0], str(val))

    _apply_apa_table_borders(table)

    # Not: alt çizginin altında, tablo gövdesinin dışında
    note_text = result.get("note", "")
    if note_text:
        note_p = doc.add_paragraph()
        note_p.paragraph_format.space_before = Pt(6)
        _add_apa_note(note_p, note_text)

def _merge_label_maps(
    label_map: Optional[Dict[str, str]],
    custom_labels: Optional[Dict[str, str]],
) -> Optional[Dict[str, str]]:
    merged: Dict[str, str] = {}
    if label_map:
        merged.update(label_map)
    if custom_labels:
        merged.update(custom_labels)
    return merged or None


from hypothesis_engine import SAMPLE_SECTION_TYPES


def _is_sample_section_result(result: dict) -> bool:
    rtype = str(result.get("type") or "")
    if rtype in SAMPLE_SECTION_TYPES:
        return True
    return not result.get("hypothesis_id")


SAYI_EKI = {
    1: "'e", 2: "'ye", 3: "'e", 4: "'e", 5: "'e",
    6: "'ya", 7: "'ye", 8: "'e", 9: "'a", 10: "'a",
}


def _hypothesis_section_title(hypothesis: dict, index: int) -> str:
    eki = SAYI_EKI.get(index, "'e")
    return f"Araştırma Sorusu {index}{eki} İlişkin Bulgular"


def _group_results_for_export(
    results: List[dict],
    hypotheses: Optional[List[dict]] = None,
) -> List[Tuple[str, List[Tuple[int, dict]]]]:
    """(bölüm başlığı, [(orijinal_index, result), ...]) listesi döndürür."""
    hypotheses = hypotheses or []
    sections: List[Tuple[str, List[Tuple[int, dict]]]] = []

    sample_items = [
        (i, r) for i, r in enumerate(results) if _is_sample_section_result(r)
    ]
    if sample_items:
        sections.append(("Örnekleme İlişkin Bulgular", sample_items))

    hyp_ids = [str(h.get("id")) for h in hypotheses if h.get("id")]
    seen = set(hyp_ids)
    for idx, hid in enumerate(hyp_ids, start=1):
        items = [
            (i, r) for i, r in enumerate(results)
            if str(r.get("hypothesis_id") or "") == hid
        ]
        if items:
            hyp = next((h for h in hypotheses if str(h.get("id")) == hid), {})
            sections.append((_hypothesis_section_title(hyp, idx), items))

    other_ids = {
        str(r.get("hypothesis_id"))
        for r in results
        if r.get("hypothesis_id") and str(r.get("hypothesis_id")) not in seen
    }
    for hid in sorted(other_ids):
        items = [
            (i, r) for i, r in enumerate(results)
            if str(r.get("hypothesis_id") or "") == hid
        ]
        if items:
            sections.append((f"{hid} Bulguları", items))

    ungrouped = [
        (i, r) for i, r in enumerate(results)
        if not _is_sample_section_result(r) and not r.get("hypothesis_id")
    ]
    if ungrouped:
        sections.append(("Diğer Analiz Bulguları", ungrouped))

    if not sections:
        return [("", [(i, r) for i, r in enumerate(results)])]
    return sections


def build_word_document(
    results: List[dict],
    bulgular: Optional[Dict[str, str]] = None,
    intro: str = "",
    label_map: Optional[Dict[str, str]] = None,
    custom_labels: Optional[Dict[str, str]] = None,
    custom_titles: Optional[Dict[str, str]] = None,
    hypotheses: Optional[List[dict]] = None,
) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    doc.add_heading("BULGULAR", level=1)
    if intro:
        doc.add_paragraph(intro)
    doc.add_paragraph()

    resolved_labels = _merge_label_maps(label_map, custom_labels)
    sections = _group_results_for_export(results, hypotheses)

    for section_title, items in sections:
        if section_title:
            doc.add_heading(section_title, level=2)
            doc.add_paragraph()
        for i, result in items:
            export_result = dict(result)
            custom_key = str(i)
            if custom_titles and custom_key in custom_titles and custom_titles[custom_key]:
                num, _ = _split_table_title(export_result.get("title", ""))
                export_result["title"] = f"{num}. {custom_titles[custom_key]}"
            add_apa_table(doc, export_result, resolved_labels)
            doc.add_paragraph()
            key = str(i)
            if bulgular and key in bulgular and bulgular[key]:
                bulgu_p = doc.add_paragraph(bulgular[key])
                bulgu_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                bulgu_p.paragraph_format.first_line_indent = Pt(0)
                for run in bulgu_p.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
            doc.add_paragraph()
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(12)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

